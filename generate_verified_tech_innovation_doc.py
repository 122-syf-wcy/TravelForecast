from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path("智教黔行技术创新及对比_真实核对版.docx")


def set_run(run, font="宋体", size=10.5, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level in (1, 2, 3) else "Normal"
    r = p.add_run(text)
    set_run(r, "黑体", 16 if level == 1 else 13 if level == 2 else 11.5, True)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r)
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    set_run(r, size=9.5, bold=header, color="FFFFFF" if header else None)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, "4472C4")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph("")
    return table


def setup():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return doc


def main():
    doc = setup()
    add_heading(doc, "3.2 技术创新", 1)
    add_p(doc, "本系统面向六盘水山地文旅和研学服务场景，在客流预测、AI 导游、多端协同、数据治理和运营辅助展示等方面进行了综合设计。与传统旅游信息系统相比，本系统不是单纯展示景点信息，而是将预测模型、智能问答、数字人交互、可视化分析和业务管理流程整合到统一平台中，形成面向景区运营和游客服务的一体化应用方案。")

    add_heading(doc, "3.2.1 预测模型创新：ARIMA 与 LSTM 双流融合预测框架", 2)
    add_p(doc, "传统方案多采用单一统计模型或简单趋势外推，面对节假日、天气变化和景区客流突增时适应能力不足。本系统采用 ARIMA 与 LSTM 相结合的双流预测思路：ARIMA 侧重捕捉时间序列中的趋势和周期规律，LSTM 侧重学习节假日、天气、周末、海拔等多因素下的非线性变化。")
    add_p(doc, "本系统通过权重融合方式综合两个模型的预测结果，并提供不同权重组合的对比分析能力，使模型能够根据不同景区和不同数据状态进行调整。该设计比单一模型更适合山地景区客流波动明显、旺淡季差异较大的应用场景。")
    add_p(doc, "系统以双流预测、权重融合和模型对比能力作为主要特色，能够根据不同景区的数据状态进行调整，为景区客流研判提供更加灵活的技术支撑。")

    add_heading(doc, "3.2.2 AI 导游创新：RAG 知识检索 + Function Calling + 语音交互", 2)
    add_p(doc, "传统旅游系统中的智能客服通常以固定问答或菜单式咨询为主，难以回答研学知识、景区安全、客流预测和行程建议等综合问题。本系统引入 RAG 检索增强思路，将六盘水景区知识、研学知识、安全提醒、碳足迹说明等内容作为知识来源，为 AI 导游回答提供上下文支撑。")
    add_p(doc, "在交互方式上，系统支持文本对话、流式回复、语音播报和数字人讲解，并能在用户询问客流或天气时联动预测服务和天气信息，形成“知识问答 + 实时工具调用 + 语音讲解”的组合式体验。")
    add_p(doc, "相较于普通客服机器人，本系统的优势在于回答内容更贴近六盘水本地文旅场景，能够围绕景区特点、研学主题、出行安全和客流状态提供更具场景化的服务。")

    add_heading(doc, "3.2.3 架构创新：多服务解耦与统一网关接入", 2)
    add_p(doc, "传统中小型旅游系统多采用单体式架构，业务、管理端、AI 能力和数据分析耦合较重，后续扩展和维护成本较高。本系统采用前后端分离与多服务协同的架构，将业务管理、统一网关、AI 服务、客流预测、数字人交互、Web 端和小程序端进行模块化拆分。")
    add_p(doc, "统一网关负责对外访问入口和服务转发，后端服务负责业务数据和权限处理，预测服务负责客流分析，AI 服务负责智能问答和数字人交互。不同模块职责清晰，便于独立开发、部署和后续扩展。")
    add_p(doc, "本系统可以概括为“微服务化拆分 + 多端接入 + 智能服务独立部署”的工程架构创新，重点解决文旅系统多角色、多终端、多智能服务协同的问题。")

    add_heading(doc, "3.2.4 数据治理创新：真实业务库 + 预测因子 + 离线隐私实验", 2)
    add_p(doc, "本系统围绕景区、用户、客流记录、预测结果、AI 会话、行程规划和订单等核心业务对象建立数据结构，为后续客流分析和智能服务提供数据基础。与只维护景区介绍文本的传统系统相比，本系统具备更完整的业务数据闭环。")
    add_p(doc, "在预测因子方面，系统综合考虑景区容量、海拔、周末、节假日、天气和温度等因素，使预测逻辑更贴近山地旅游场景。对于历史数据不足的情况，系统也保留了降级处理机制，以保证演示和服务接口的连续可用。")
    add_p(doc, "在隐私保护方面，系统进行了差分隐私加噪的离线实验，用于分析隐私预算与数据效用之间的关系，为后续完善客流数据脱敏和隐私保护机制提供参考。")

    add_heading(doc, "3.2.5 可视化应用优化：景区分布地图、客流热力图与政策沙盘", 2)
    add_p(doc, "景区分布地图、客流热力图和政策沙盘属于成熟可视化技术的场景化集成。本系统的价值在于将这些可视化能力与景区运营场景结合，用更直观的方式呈现景区分布、客流密度和政策调整影响。")
    add_p(doc, "在游客侧，系统提供景区地图、景区标记、客流热力展示和实时服务信息，帮助用户更直观地了解景区位置、游览区域和人流状态。在管理侧，系统通过图表和政策参数模拟辅助运营人员观察客流、收入、满意度等指标变化趋势。")
    add_p(doc, "因此，本节更适合概括为“可视化应用优化”或“运营辅助展示优化”。其优势主要体现在文旅业务数据的直观展示、交互体验提升和运营决策辅助。")

    add_heading(doc, "3.2.6 场景创新：六盘水山地研学与文旅业务适配", 2)
    add_p(doc, "系统围绕六盘水山地文旅和研学场景做了真实适配：知识库内容覆盖高海拔安全、防寒防滑、喀斯特地貌、三线建设、碳足迹等主题；数字人能够结合页面和景区内容进行讲解；预测逻辑也考虑了景区容量、节假日、天气和海拔等因素。")
    add_p(doc, "系统还覆盖 Web 管理端、用户端、微信小程序、景区浏览、行程规划、AI 问答、客流预测、订单和个人中心等业务链路。相比只提供景点列表和静态介绍的传统系统，本系统的真实优势在于“预测服务 + AI 导游 + 多端前端 + 数据可视化 + 微服务部署”的组合能力。")

    add_heading(doc, "3.2.7 创新对比总表", 2)
    add_table(doc, ["对比维度", "传统旅游信息系统", "通用智慧旅游平台", "本系统"], [
        ["预测模型", "以静态统计或人工经验判断为主", "通常采用单一趋势预测或简单报表分析", "采用 ARIMA 与 LSTM 双流融合思路，结合天气、节假日、周末、海拔等因素进行客流预测"],
        ["AI 服务", "以固定问答和人工客服为主", "提供普通智能客服或简单导览问答", "结合本地景区知识库、客流预测和天气信息，提供更贴近六盘水文旅场景的 AI 导游服务"],
        ["系统架构", "单体系统或简单前后端分离", "以管理后台和游客端为主，智能服务耦合较重", "采用多服务协同架构，将业务管理、统一网关、预测服务、AI 服务、数字人交互和多端前端进行解耦"],
        ["数据治理", "主要维护景点介绍和基础订单数据", "具备一定统计分析能力，但预测因子较少", "围绕景区、用户、客流、预测、AI 会话和订单等业务对象形成数据闭环，并探索隐私保护和多因素预测"],
        ["可视化展示", "以列表、图片和二维图表为主", "具备常规大屏和图表展示", "集成地图、客流热力图和政策参数模拟，用于提升展示直观性和运营辅助分析能力"],
        ["场景适配", "通用旅游信息展示", "适配一般景区管理需求", "围绕六盘水山地旅游、研学讲解、高海拔安全、避暑滑雪等场景进行内容和功能适配"],
        ["应用价值", "解决基本信息查询问题", "提升景区数字化管理效率", "兼顾游客服务、研学讲解、客流预测和运营辅助，形成较完整的智慧文旅应用闭环"],
    ])

    add_heading(doc, "3.2.8 技术特色总结", 2)
    add_p(doc, "综合来看，本系统的技术特色可以概括为：双流客流预测、RAG 增强 AI 导游、多服务协同架构、山地研学场景适配、地图与图表可视化展示、政策模拟辅助分析。")
    add_p(doc, "这些能力共同服务于六盘水智慧文旅场景，使系统既能面向游客提供智能导览和出行辅助，也能面向景区管理者提供客流研判和运营决策参考。")

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
