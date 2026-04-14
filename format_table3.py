from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

doc = Document('/Users/dongsiwei/TravelForecast/表3_扩充版.docx')
table = doc.tables[0]

# --- 三线表：清除所有边框，只保留顶线、栏目线（表头底部）、底线 ---

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    # 移除旧的 tcBorders
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcBorders)

# 边框样式
thick = {"val": "single", "sz": "12", "color": "000000"}  # 1.5pt 粗线
thin  = {"val": "single", "sz": "8",  "color": "000000"}  # 1pt 细线
none  = {"val": "nil",    "sz": "0",  "color": "000000"}  # 无边框

num_rows = len(table.rows)
num_cols = len(table.columns)

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        borders = {}
        # 左右竖线：全部去掉
        borders["left"] = none
        borders["right"] = none
        
        if i == 0:
            # 表头行：顶线粗 + 底线细
            borders["top"] = thick
            borders["bottom"] = thin
        elif i == num_rows - 1:
            # 最后一行：底线粗
            borders["top"] = none
            borders["bottom"] = thick
        else:
            # 中间行：无上下线
            borders["top"] = none
            borders["bottom"] = none
        
        set_cell_border(cell, **borders)
        
        # 设置字体和对齐（宋体小五 = 9pt，字母/数字用 Times New Roman 小五）
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置段落级别的行距为单倍
            pPr = paragraph._element.get_or_add_pPr()
            spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="240" w:lineRule="auto"/>')
            old_spacing = pPr.find(qn('w:spacing'))
            if old_spacing is not None:
                pPr.remove(old_spacing)
            pPr.append(spacing)
            for run in paragraph.runs:
                text = run.text
                # 西文字体 Times New Roman，中文字体 宋体，均为小五号(9pt)
                run.font.size = Pt(9)
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:cs'), 'Times New Roman')
                if i == 0:
                    run.bold = True

# 合并同类别单元格（类别列和贡献度列）
# 环境数据: rows 1-3
table.cell(1, 0).merge(table.cell(3, 0))
table.cell(1, 5).merge(table.cell(3, 5))
# 游客行为数据: rows 4-7
table.cell(4, 0).merge(table.cell(7, 0))
table.cell(4, 5).merge(table.cell(7, 5))
# 山地特色数据: rows 8-10
table.cell(8, 0).merge(table.cell(10, 0))
table.cell(8, 5).merge(table.cell(10, 5))
# 节庆事件数据: rows 11-12
table.cell(11, 0).merge(table.cell(12, 0))
table.cell(11, 5).merge(table.cell(12, 5))

# 合并后重新设置对齐
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 垂直居中
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            old_vAlign = tcPr.find(qn('w:vAlign'))
            if old_vAlign is not None:
                tcPr.remove(old_vAlign)
            tcPr.append(vAlign)

# 设置列宽
widths = [Cm(2.5), Cm(3.0), Cm(4.5), Cm(3.5), Cm(3.0), Cm(1.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# 添加注释
p = doc.add_paragraph()
run = p.add_run('注：上述数据总量超过28万条，经清洗、对齐与融合后形成统一的多源异构数据集用于模型训练与验证。')
run.font.size = Pt(9)
run.italic = True
rPr = run._element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), '宋体')
rFonts.set(qn('w:cs'), 'Times New Roman')

doc.save('/Users/dongsiwei/TravelForecast/表3_扩充版.docx')
print("三线表已生成：表3_扩充版.docx")
