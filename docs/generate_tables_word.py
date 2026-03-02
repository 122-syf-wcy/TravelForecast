# -*- coding: utf-8 -*-
"""
游韵华章 · Word三线表文档生成脚本
所有数据均来自项目源码，数据来源在注释中标注
生成用于挑战杯/计算机设计大赛的规范三线表文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'charts')
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '项目三线表_挑战杯与计算机设计大赛.docx')

# ============================================================
# 真实数据 (来源标注)
# ============================================================

# 来源: data_generator.py SCENIC_CONFIG
SCENIC_CONFIG = {
    1: {"name": "梅花山风景区",       "base_flow": 3000, "capacity": 8000, "altitude": 2400},
    2: {"name": "玉舍国家森林公园",   "base_flow": 2500, "capacity": 6000, "altitude": 2300},
    3: {"name": "乌蒙大草原",         "base_flow": 3500, "capacity": 10000, "altitude": 2857},
    4: {"name": "水城古镇",           "base_flow": 2000, "capacity": 5000, "altitude": 1800},
    5: {"name": "明湖国家湿地公园",   "base_flow": 1800, "capacity": 4000, "altitude": 1750},
}

# 来源: main.py SCENIC_NAMES
SCENIC_NAMES_API = {
    1: "玉舍雪山", 2: "乌蒙大草原", 3: "明湖国家湿地公园",
    4: "梅花山国际滑雪场", 5: "野玉海山地旅游度假区",
}


def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    """设置单元格边框 (三线表核心函数)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    if top:
        t = parse_xml(f'<w:top {nsdecls("w")} w:val="{top["val"]}" w:sz="{top["sz"]}" w:color="{top["color"]}" w:space="0"/>')
        tcBorders.append(t)
    if bottom:
        b = parse_xml(f'<w:bottom {nsdecls("w")} w:val="{bottom["val"]}" w:sz="{bottom["sz"]}" w:color="{bottom["color"]}" w:space="0"/>')
        tcBorders.append(b)
    if start:
        s = parse_xml(f'<w:start {nsdecls("w")} w:val="{start["val"]}" w:sz="{start["sz"]}" w:color="{start["color"]}" w:space="0"/>')
        tcBorders.append(s)
    if end:
        e = parse_xml(f'<w:end {nsdecls("w")} w:val="{end["val"]}" w:sz="{end["sz"]}" w:color="{end["color"]}" w:space="0"/>')
        tcBorders.append(e)

    tcPr.append(tcBorders)


def make_three_line_table(doc, headers, rows, caption="", col_widths=None):
    """
    创建标准三线表
    三线表规范：顶线(粗) + 栏目线(细) + 底线(粗)，无竖线
    """
    # 表格标题
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.bold = True

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    thick = {"val": "single", "sz": "12", "color": "000000"}  # 1.5pt
    thin  = {"val": "single", "sz": "6",  "color": "000000"}  # 0.75pt
    none  = {"val": "nil",    "sz": "0",  "color": "000000"}

    # 填充表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.bold = True
        # 表头: 顶线粗，底线细，无竖线
        set_cell_border(cell, top=thick, bottom=thin, start=none, end=none)

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            if i == len(rows) - 1:
                # 最后一行: 底线粗
                set_cell_border(cell, top=none, bottom=thick, start=none, end=none)
            else:
                # 中间行: 无边框
                set_cell_border(cell, top=none, bottom=none, start=none, end=none)

    doc.add_paragraph()  # 空行
    return table


def add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_source_note(doc, source_text):
    """添加数据来源注释"""
    p = doc.add_paragraph()
    run = p.add_run(f"数据来源：{source_text}")
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ============================================================
    # 封面标题
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('游韵华章 · 智慧文旅一体化平台')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('项目核心数据表 (三线表)')
    run2.font.size = Pt(14)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('用于挑战杯计划书 / 计算机设计大赛')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ============================================================
    # 表1: 微服务架构配置表
    # ============================================================
    add_heading(doc, '一、微服务架构配置表', level=1)

    make_three_line_table(doc,
        headers=["服务名称", "端口", "技术栈", "核心职责", "限流(req/s)"],
        rows=[
            # 来源: application.yml routes + 各服务README
            ["API网关",       "8888", "Spring Cloud Gateway",  "JWT认证/限流/路由/熔断",     "全局"],
            ["Web业务后端",   "8080", "Spring Boot 3.2 + Java 17", "管理员/用户/商家业务 (56个Controller)", "100/200"],
            ["AI智能服务",    "8081", "Spring Boot 3.2 + DashScope", "AI聊天/行程规划/RAG知识库/研学",   "50/100"],
            ["小程序后端",    "8082", "Spring Boot 3.2 + 微信SDK", "电商/积分/微信支付/独立部署",      "80/150"],
            ["预测服务",      "8001", "FastAPI + TensorFlow",  "ARIMA/LSTM/双流融合预测",    "30/60"],
            ["数字人服务",    "8083", "FastAPI + TTS",         "AI对话/语音合成/场景讲解",    "40/80"],
        ],
        caption="表1  系统微服务架构配置",
        col_widths=[3, 1.5, 4, 5, 2.5],
    )
    add_source_note(doc, "application.yml 路由与限流配置 + 各服务README.md")

    # ============================================================
    # 表2: 预测模型性能对比表
    # ============================================================
    add_heading(doc, '二、预测模型性能对比表', level=1)

    make_three_line_table(doc,
        headers=["模型", "准确率", "置信度", "输入维度", "序列长度", "网络结构", "训练数据量"],
        rows=[
            # 来源: arima_model.py, lstm_model.py, lstm_new.py, dual_stream_model.py
            ["ARIMA",       "82%", "80%", "1维 (客流)", "-",    "SARIMAX(1,1,1)(1,1,1,7)", "90天"],
            ["LSTM (旧版)", "87%", "85%", "1维 (客流)", "14天", "LSTM(50)→DO→LSTM(50)→DO→Dense(25)→Dense(1)", "90天"],
            ["LSTM (多变量)", "-", "-",   "6维 (多特征)", "14天", "LSTM(64)→DO→LSTM(32)→DO→Dense(16)→Dense(1)", "365天"],
            ["Dual-Stream", "92%", "92%", "6维+ARIMA融合", "14天", "ARIMA + 多变量LSTM + 动态权重", "365天"],
        ],
        caption="表2  四种预测模型性能对比",
        col_widths=[2.5, 1.3, 1.3, 2.5, 1.5, 5, 1.8],
    )
    add_source_note(doc, "arima_model.py (line 22-23) / lstm_model.py (line 26-28) / dual_stream_model.py (line 119-120)")

    # ============================================================
    # 表3: 景区基础数据表
    # ============================================================
    add_heading(doc, '三、景区基础数据表', level=1)

    make_three_line_table(doc,
        headers=["景区ID", "景区名称", "日均基础客流", "最大承载量", "海拔(m)", "拥挤阈值"],
        rows=[
            # 来源: data_generator.py SCENIC_CONFIG + arima_model.py _get_congestion_level
            [1, "梅花山风景区",       3000, 8000,  2400, ">6000:极端/>4500:高/>3000:中"],
            [2, "玉舍国家森林公园",   2500, 6000,  2300, "同上"],
            [3, "乌蒙大草原",         3500, 10000, 2857, "同上"],
            [4, "水城古镇",           2000, 5000,  1800, "同上"],
            [5, "明湖国家湿地公园",   1800, 4000,  1750, "同上"],
        ],
        caption="表3  五大景区基础参数配置",
        col_widths=[1.5, 3.5, 2, 2, 1.5, 5],
    )
    add_source_note(doc, "data_generator.py SCENIC_CONFIG (line 1-7) + arima_model.py _get_congestion_level (line 258-267)")

    # ============================================================
    # 表4: 6维特征向量说明表
    # ============================================================
    add_heading(doc, '四、多变量LSTM特征向量说明表', level=1)

    make_three_line_table(doc,
        headers=["维度", "特征名称", "取值范围", "归一化方法", "数据来源"],
        rows=[
            # 来源: lstm_new.py line 26 + data_generator.py get_features() line 284-306
            ["Dim 0", "历史客流 (flow)",     "0 ~ capacity", "MinMaxScaler",     "flow_records表"],
            ["Dim 1", "节庆标记 (is_holiday)", "0 或 1",       "二值编码",          "HOLIDAYS_2025列表"],
            ["Dim 2", "周末标记 (is_weekend)", "0 或 1",       "二值编码",          "weekday() >= 5"],
            ["Dim 3", "天气指数 (weather_idx)", "0.0 ~ 1.0",   "type_index / 3.0", "WEATHER_TYPES字典"],
            ["Dim 4", "温度 (temp_avg)",       "0.0 ~ 1.0",   "temp_avg / 40.0",  "天气API/模拟生成"],
            ["Dim 5", "海拔 (altitude)",       "0.0 ~ 1.0",   "altitude / 3000.0", "SCENIC_CONFIG"],
        ],
        caption="表4  多变量LSTM 6维特征向量",
        col_widths=[1.5, 3.5, 2, 3, 3.5],
    )
    add_source_note(doc, "lstm_new.py (line 25-26, feature_dim=6) + data_generator.py get_features() (line 284-306)")

    # ============================================================
    # 表5: 天气对客流影响因子表
    # ============================================================
    add_heading(doc, '五、客流影响因子表', level=1)

    make_three_line_table(doc,
        headers=["因子类型", "因子名称", "系数/取值", "说明"],
        rows=[
            # 来源: data_generator.py WEATHER_TYPES + calculate_daily_flow
            ["天气", "晴天 (sunny)",    "x1.2, P=0.40", "客流增加20%"],
            ["天气", "多云 (cloudy)",    "x1.0, P=0.30", "无影响 (基准)"],
            ["天气", "阴天 (overcast)",  "x0.9, P=0.15", "客流减少10%"],
            ["天气", "雨天 (rainy)",     "x0.6, P=0.15", "客流减少40%"],
            ["季节", "夏季 (6-8月)",     "x1.3",         "六盘水凉都旅游旺季"],
            ["季节", "春秋 (4-5,9-10月)", "x1.1",        "温和季节"],
            ["季节", "冬季",             "x0.8",         "淡季"],
            ["日期", "周末",             "x1.5",         "weekday() >= 5"],
            ["日期", "节假日",           "x2.0",         "HOLIDAYS_2025列表"],
            ["随机", "随机波动",         "U(0.9, 1.1)",  "random.uniform()"],
        ],
        caption="表5  客流量计算影响因子",
        col_widths=[2, 3.5, 3, 5],
    )
    add_source_note(doc, "data_generator.py WEATHER_TYPES + calculate_daily_flow() (line 123-167)")

    # ============================================================
    # 表6: 小时级客流分布系数表
    # ============================================================
    add_heading(doc, '六、小时级客流分布系数表', level=1)

    make_three_line_table(doc,
        headers=["时段", "8:00", "9:00", "10:00", "11:00", "12:00", "13:00",
                 "14:00", "15:00", "16:00", "17:00", "18:00", "19:00"],
        rows=[
            # 来源: data_generator.py HOURLY_DISTRIBUTION
            ["分布系数", 0.2, 0.4, 0.7, 0.9, 1.0, 0.9, 1.0, 0.95, 0.8, 0.6, 0.4, 0.2],
            ["梅花山 (人/h)", 50, 100, 175, 225, 250, 225, 250, 237, 200, 150, 100, 50],
            ["乌蒙大草原 (人/h)", 58, 116, 204, 262, 291, 262, 291, 277, 233, 175, 116, 58],
        ],
        caption="表6  小时级客流分布系数",
        col_widths=[2.5] + [1.1]*12,
    )
    add_source_note(doc, "data_generator.py HOURLY_DISTRIBUTION + main.py hourly_distribution (line 308-311)")

    # ============================================================
    # 表7: API网关路由与限流配置表
    # ============================================================
    add_heading(doc, '七、API网关路由与限流配置表', level=1)

    make_three_line_table(doc,
        headers=["路由ID", "匹配路径", "转发目标", "限流(req/s)", "突发容量", "熔断等待时间"],
        rows=[
            # 来源: application.yml routes + resilience4j
            ["business-service",     "/api/**",              "http://localhost:8080", 100, 200, "5s"],
            ["ai-service",           "/ai-api/**",           "http://localhost:8081", 50,  100, "10s"],
            ["miniprogram-service",  "/miniprogram-api/**",  "http://localhost:8082", 80,  150, "-"],
            ["prediction-service",   "/prediction-api/**",   "http://localhost:8001", 30,  60,  "15s"],
            ["digital-human-service","/digital-human-api/**","http://localhost:8083", 40,  80,  "-"],
            ["digital-human-ws",     "/ws/**",               "ws://localhost:8083",   "-", "-",  "-"],
        ],
        caption="表7  Spring Cloud Gateway 路由与限流配置",
        col_widths=[3, 3.5, 3, 1.8, 1.5, 2],
    )
    add_source_note(doc, "TravelForecastGateway/src/main/resources/application.yml (line 59-151)")

    # ============================================================
    # 表8: 后端代码规模统计表
    # ============================================================
    add_heading(doc, '八、后端代码规模统计表', level=1)

    make_three_line_table(doc,
        headers=["模块分类", "数量", "占比", "主要功能"],
        rows=[
            # 来源: 文件计数统计
            ["Controller (admin)",    11, "19.6%", "景区管理/用户管理/商家审核/仪表盘"],
            ["Controller (merchant)", 12, "21.4%", "景区运营/数据分析/合同/通知"],
            ["Controller (common)",    6, "10.7%", "高德地图/AI聊天/OSS上传/图片代理"],
            ["Controller (content)",   5,  "8.9%", "公告/轮播图/新闻/内容/案例"],
            ["Controller (scenic)",    6, "10.7%", "景区详情/搜索/收藏/评论"],
            ["Controller (user)",      4,  "7.1%", "个人中心/收藏/浏览历史"],
            ["Controller (其他)",     12, "21.4%", "认证/验证码/预测/订单/统计/系统"],
            ["Service层",            71,  "-",    "业务逻辑层"],
            ["Mapper层",             50,  "-",    "MyBatis-Plus数据访问层"],
            ["Entity层",             69,  "-",    "数据实体/DTO/VO"],
            ["总计Java文件",         299,  "-",    "完整后端代码"],
        ],
        caption="表8  Web业务后端代码规模统计",
        col_widths=[3.5, 1.5, 1.5, 7],
    )
    add_source_note(doc, "TravelForecastBackend/src/main/java/com/travel/ 文件计数统计")

    # ============================================================
    # 表9: 预测服务API接口表
    # ============================================================
    add_heading(doc, '九、预测服务API接口表', level=1)

    make_three_line_table(doc,
        headers=["方法", "接口路径", "参数", "功能说明"],
        rows=[
            # 来源: main.py API endpoints
            ["GET",  "/api/prediction/flow/{scenic_id}",    "model, days, factors",    "单景区未来N天客流预测"],
            ["GET",  "/api/prediction/total",               "days, model",             "全区聚合客流预测"],
            ["GET",  "/api/prediction/hourly/{scenic_id}",  "date, model",             "单景区小时级客流预测"],
            ["GET",  "/api/prediction/hourly/total",        "date, model",             "全区小时级聚合预测"],
            ["POST", "/api/prediction/train/{scenic_id}",   "model",                   "触发模型训练"],
            ["GET",  "/api/prediction/models/info",         "-",                       "获取所有模型信息"],
        ],
        caption="表9  Python预测服务API接口",
        col_widths=[1.5, 5, 3.5, 4],
    )
    add_source_note(doc, "TravelForecast-PythonPredictionService/main.py (line 107-482)")

    # ============================================================
    # 表10: 数据治理技术参数表
    # ============================================================
    add_heading(doc, '十、数据治理技术参数表', level=1)

    make_three_line_table(doc,
        headers=["技术", "参数名", "取值", "作用", "代码位置"],
        rows=[
            # 来源: preprocessing.py
            ["小波去噪", "wavelet",     "db4",       "Daubechies 4小波基",  "preprocessing.py line 27"],
            ["小波去噪", "level",       "2",         "分解层数",            "preprocessing.py line 27"],
            ["小波去噪", "threshold",   "sigma*sqrt(2*log(n))", "通用阈值公式", "preprocessing.py line 46-47"],
            ["小波去噪", "sigma",       "median(|coeffs|)/0.6745", "噪声估计",   "preprocessing.py line 46"],
            ["小波去噪", "mode",        "soft",      "软阈值处理",          "preprocessing.py line 54"],
            ["差分隐私", "epsilon",     "1.0",       "隐私预算",            "preprocessing.py line 73"],
            ["差分隐私", "sensitivity", "1.0",       "敏感度",              "preprocessing.py line 73"],
            ["差分隐私", "mechanism",   "Laplace",   "拉普拉斯噪声",       "preprocessing.py line 82"],
            ["归一化",   "method",      "MinMaxScaler", "最小-最大缩放",    "lstm_new.py line 63-64"],
            ["归一化",   "range",       "(0, 1)",    "输出范围",            "lstm_new.py line 63"],
        ],
        caption="表10  数据治理层技术参数",
        col_widths=[2, 2.5, 3.5, 3, 4],
    )
    add_source_note(doc, "utils/preprocessing.py + models/lstm_new.py")

    # ============================================================
    # 表11: 技术栈依赖版本表
    # ============================================================
    add_heading(doc, '十一、技术栈依赖版本表', level=1)

    make_three_line_table(doc,
        headers=["类别", "技术/框架", "版本", "用途"],
        rows=[
            # 来源: pom.xml + requirements.txt + 各README
            ["后端框架",   "Spring Boot",          "3.2.0",    "Java微服务框架"],
            ["后端语言",   "Java",                 "17",       "主力开发语言"],
            ["网关",       "Spring Cloud Gateway",  "-",       "API统一网关"],
            ["数据库",     "MySQL",                "8.0",      "关系型数据库"],
            ["缓存",       "Redis",                "6.0",      "缓存+限流+会话"],
            ["ORM",        "MyBatis-Plus",          "-",       "数据访问层"],
            ["认证",       "JWT",                   "-",       "Token认证"],
            ["对象存储",   "阿里云OSS",             "-",       "文件/图片存储"],
            ["Python框架", "FastAPI",              "0.115.0",  "预测服务API框架"],
            ["ASGI服务器", "Uvicorn",              "0.30.0",   "高性能Python服务器"],
            ["数据模型",   "Pydantic",             "2.9.0",    "请求/响应验证"],
            ["深度学习",   "TensorFlow (CPU)",     ">=2.15.0", "LSTM模型训练与推理"],
            ["时间序列",   "statsmodels",          ">=0.14.0", "ARIMA/SARIMAX模型"],
            ["数值计算",   "NumPy",                ">=1.26.0", "矩阵运算"],
            ["数据处理",   "Pandas",               ">=2.1.0",  "数据清洗"],
            ["机器学习",   "scikit-learn",         ">=1.3.0",  "数据预处理/MinMaxScaler"],
            ["小波变换",   "PyWavelets",            "-",       "小波去噪 (db4)"],
            ["AI大模型",   "通义千问 (DashScope)",   "-",       "AI聊天/RAG/行程规划"],
            ["地图",       "高德地图API",            "-",       "POI搜索/导航/定位"],
            ["3D渲染",     "Three.js",              "-",       "数字人3D模型"],
            ["小程序",     "UniApp + Vue3",          "-",       "微信小程序前端"],
            ["熔断",       "Resilience4j",           "-",       "服务熔断与降级"],
        ],
        caption="表11  技术栈与依赖版本",
        col_widths=[2, 3.5, 2.5, 5],
    )
    add_source_note(doc, "pom.xml + requirements.txt + 各模块README.md")

    # ============================================================
    # 保存文档
    # ============================================================
    doc.save(OUTPUT_PATH)
    print(f"[OK] Word三线表文档已生成: {OUTPUT_PATH}")
    print(f"     共 11 张三线表")


if __name__ == '__main__':
    print("=" * 60)
    print("  游韵华章 · Word三线表文档生成")
    print("  所有数据来源于项目源码")
    print("=" * 60)
    main()
    print("=" * 60)
