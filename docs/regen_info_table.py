"""
重新生成《智教黔行_作品信息概要表.docx》

依据：截至 2026-04-19 仓库实际代码（commit df09348 之后）
覆盖：刷新创新描述、作品简介、参考文献、相关文件清单、AI 工具说明
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_SUBMISSION = REPO_ROOT / "提交材料-智教黔行/03-设计与开发文档/智教黔行_作品信息概要表.docx"
OUT_ROOT = REPO_ROOT / "智教黔行_作品信息概要表.docx"


def _set_cn_font(run, name: str = "宋体", size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _shade(cell, fill_hex: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    color: RGBColor | None = None,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.35
    for idx, line in enumerate(text.split("\n")):
        if idx == 0:
            run = p.add_run(line)
        else:
            run = p.add_run("\n" + line)
        _set_cn_font(run, name="宋体", size_pt=size, bold=bold)
        if color is not None:
            run.font.color.rgb = color


def _label_cell(cell, text: str) -> None:
    _shade(cell, "EAF1F8")
    _set_cell_text(cell, text, bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for tag in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{tag}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "808080")
        borders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _new_full_width_table(doc: Document, rows: int, cols: int) -> "Table":
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _set_table_borders(tbl)
    return tbl


def _set_global_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)


def build(doc: Document) -> None:
    _set_global_styles(doc)

    # ===== 标题 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国大学生计算机设计大赛")
    _set_cn_font(run, name="黑体", size_pt=18, bold=True)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("作品信息概要表（2025 版 · 智教黔行专用）")
    _set_cn_font(run, name="黑体", size_pt=14, bold=True)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # ===== 作品基本信息 =====
    tbl = _new_full_width_table(doc, rows=2, cols=4)
    widths = [3.6, 4.2, 3.0, 5.6]
    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    _label_cell(tbl.rows[0].cells[0], "作品编号")
    _set_cell_text(tbl.rows[0].cells[1], "【报名后由大赛系统分配】")
    _label_cell(tbl.rows[0].cells[2], "作品名称")
    _set_cell_text(
        tbl.rows[0].cells[3],
        "智教黔行——基于多源数据融合与 LSTM-ARIMA 双流动态权重模型的六盘水山地智慧文旅一体化研学平台",
    )

    _label_cell(tbl.rows[1].cells[0], "作品大类")
    _set_cell_text(tbl.rows[1].cells[1], "软件应用与开发")
    _label_cell(tbl.rows[1].cells[2], "作品小类")
    _set_cell_text(tbl.rows[1].cells[3], "【按报名赛道填写，建议选“人工智能 / 大数据 / 智慧城市”相关方向】")

    doc.add_paragraph()

    # ===== 作品简介 =====
    tbl = _new_full_width_table(doc, rows=1, cols=1)
    _set_cell_text(
        tbl.rows[0].cells[0],
        (
            "【作品简介】（≤100 字）\n"
            "智教黔行是面向六盘水山地旅游的智慧研学一体化平台。系统由 6 个微服务组成，"
            "提供 ARIMA-LSTM 双流融合客流预测、RAG 混合检索 AI 数字人“黔小游”、3D 地形导览、"
            "AI 行程规划、政策沙盒、紧急救援与可观测性栈，覆盖 Web 与微信小程序双端，"
            "已部署上线 https://travel.dongsiwei.com。"
        ),
    )

    doc.add_paragraph()

    # ===== 创新描述 =====
    tbl = _new_full_width_table(doc, rows=1, cols=1)
    _set_cell_text(
        tbl.rows[0].cells[0],
        (
            "【创新描述】（≤100 字）\n"
            "1) 算法创新：双流动态权重融合（21 点网格搜索 α∈[0,1]）+ 海拔归一化特征 + 指标注册表；"
            "2) 工程创新：BM25 + DashScope 1536 维向量混合检索（Redis 轻量向量库 + 优雅降级）；"
            "3) 系统创新：Function Calling + 分句流式 TTS + 三级缓存（28 问预热）+ LGTM 全栈可观测性。"
        ),
    )

    doc.add_paragraph()

    # ===== 特别说明 =====
    tbl = _new_full_width_table(doc, rows=1, cols=1)
    _set_cell_text(
        tbl.rows[0].cells[0],
        (
            "【特别说明】\n"
            "1. 地图说明：3D 地形图基于 Apache ECharts GL 渲染，高程数据来源于公开 DEM 数据集；"
            "系统不涉及国界 / 疆域标注，符合大赛相关合规要求。\n"
            "2. 前期基础：本作品为本次参赛周期内全新设计、独立开发的系统；71 张 MySQL 数据表、"
            "约 60 万行代码（不含 node_modules / target）均由团队自主完成，无前期基础可继承。\n"
            "3. 人工智能辅助工具使用说明：\n"
            "   (1) 产品功能中集成的 AI 服务全部为国内合规平台：阿里云通义千问 DashScope（"
            "https://help.aliyun.com/zh/dashscope/）、阿里云 CosyVoice 语音合成、深度求索 DeepSeek（"
            "https://platform.deepseek.com/）、Microsoft Edge TTS。所有外部 AI 服务均通过官方 API 合规调用，"
            "并受制于网关 RateLimiter 与 Resilience4j 熔断保护。\n"
            "   (2) 开发过程中曾使用国内 AI 辅助编码工具进行代码补全、重构建议与文档草拟，"
            "AI 辅助生成的代码占比约 15%，所有 AI 生成内容均经团队成员逐行审查、修改并经过单元测试 / 接口测试后方可入库。\n"
            "   (3) 项目核心创新——ARIMA-LSTM 双流动态权重模型的算法设计、海拔特征引入、混合检索框架（BM25 + 向量）、"
            "三级缓存机制、可观测性栈集成、微服务架构方案、71 张 MySQL 表的全部业务逻辑与数据库设计，"
            "均为团队成员独立完成，未依赖 AIGC 工具直接生成。"
        ),
    )

    doc.add_paragraph()

    # ===== 作者及分工 =====
    p = doc.add_paragraph()
    run = p.add_run(
        "【作者及其分工比例】（请将“队员1”等替换为真实姓名；分工百分比按真实贡献填写，每行合计 100%）"
    )
    _set_cn_font(run, name="宋体", size_pt=10.5, bold=True)

    members = ["队员1", "队员2", "队员3"]
    work_items = [
        ("组织协调", "[__%]", "[__%]", "[__%]"),
        ("作品创意 / 需求调研", "[__%]", "[__%]", "[__%]"),
        ("竞品分析 / 文献阅读", "[__%]", "[__%]", "[__%]"),
        ("方案设计 / 架构设计", "[__%]", "[__%]", "[__%]"),
        ("算法实现（双流模型 / 混合检索）", "[__%]", "[__%]", "[__%]"),
        ("Java 后端实现（4 个微服务）", "[__%]", "[__%]", "[__%]"),
        ("Python 服务实现（预测 / 数字人）", "[__%]", "[__%]", "[__%]"),
        ("Vue 前端 + 微信小程序", "[__%]", "[__%]", "[__%]"),
        ("部署 / 可观测性 / 安全治理", "[__%]", "[__%]", "[__%]"),
        ("测试与压测", "[__%]", "[__%]", "[__%]"),
        ("文档撰写 / 答辩准备", "[__%]", "[__%]", "[__%]"),
    ]
    cols = 1 + len(members)
    tbl = _new_full_width_table(doc, rows=1 + len(work_items), cols=cols)

    headers = ["项目"] + members
    for c_idx, h in enumerate(headers):
        _label_cell(tbl.rows[0].cells[c_idx], h)
    for c_idx in [0, 1, 2, 3]:
        for row in tbl.rows:
            row.cells[c_idx].width = Cm(4.0 if c_idx == 0 else 4.0)

    for r_idx, item in enumerate(work_items):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(item):
            _set_cell_text(row.cells[c_idx], val, align=WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ===== 指导教师 / 平台 / 工具 / 参考 / 提交内容 =====
    rows = [
        (
            "指导教师作用",
            "■作品创意   ■理论指导   ■技术方案   □实验场地   □硬件资源   □数据提供   □后勤支持   □宣讲通知   ■组织协调   □经费支持",
        ),
        (
            "开发制作平台",
            "■Windows 11   ■Linux（CentOS 7+ / Ubuntu 22.04）   ■macOS（开发机）   □其他",
        ),
        (
            "运行展示平台",
            "■Windows 10/11（Web）   ■Linux（生产 ECS）   ■macOS（Web）   ■Android（小程序）   ■iOS（小程序）   ■其他：微信开发者工具 / 微信客户端",
        ),
        (
            "开发制作工具",
            "IntelliJ IDEA 2024（Java 后端、Spring Boot 3）；Visual Studio Code 1.95 + Cursor（Python / Vue / TypeScript）；HBuilderX（UniApp 微信小程序）；"
            "Maven 3.9 / npm 10 / Vite 5；MySQL Workbench 8.0、Redis Insight；Postman、Apifox；Docker 24 + Docker Compose v2；"
            "Git + GitHub；微信开发者工具；Grafana 11 / Prometheus 2.55 / Tempo / Loki / Promtail",
        ),
        (
            "参考文献、项目或作品（前 3 项）",
            "1) Box G E P, et al. Time Series Analysis: Forecasting and Control, 5th ed, Wiley, 2015.\n"
            "2) Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, NeurIPS, 2020.\n"
            "3) Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond, 2009.",
        ),
        (
            "提交内容",
            "■素材压缩包   ■报告文档   ■演示视频   ■PPT   ■源代码   ■部署文件（Docker / Compose / Nginx / start-all.sh）   ■数据集（experiments/results/*.csv）   ■模型（指标注册表 + 预热缓存）   □作品文件   ■其他：可观测性配置（Grafana 仪表盘 / Prometheus / Tempo / Loki）",
        ),
    ]
    for label, val in rows:
        tbl = _new_full_width_table(doc, rows=1, cols=2)
        for col_idx, w in enumerate([4.0, 12.4]):
            for row in tbl.rows:
                row.cells[col_idx].width = Cm(w)
        _label_cell(tbl.rows[0].cells[0], label)
        _set_cell_text(tbl.rows[0].cells[1], val)
        doc.add_paragraph()

    # ===== 相关文件清单 =====
    p = doc.add_paragraph()
    run = p.add_run("【相关文件清单】（按照大赛要求列出本作品涉及的所有提交文件）")
    _set_cn_font(run, name="宋体", size_pt=10.5, bold=True)

    files = [
        (
            "1",
            "智教黔行_作品信息概要表.pdf",
            "中国大学生计算机设计大赛作品信息概要表（本文件 PDF 版本）",
        ),
        (
            "2",
            "智教黔行_设计和开发文档.pdf",
            "系统需求分析、概要 / 详细设计、测试报告、安装与使用、项目总结（V3.0）",
        ),
        (
            "3",
            "答辩演示文档.pptx / .pdf",
            "答辩 PPT 及其 PDF 版本（17 页）",
        ),
        (
            "4",
            "答辩视频.mp4",
            "答辩录像（≤10 分钟，1080p，≤500MB）",
        ),
        (
            "5",
            "作品演示视频.mp4",
            "智教黔行系统完整功能演示录屏（Web 端 + 商家端 + 管理后台 + 微信小程序）",
        ),
        (
            "6",
            "智教黔行-素材与源码.zip",
            "源代码压缩包：Gateway / Backend / AI-Backend / MP-Backend / Python 预测 / 数字人 / Vue 前端 / 微信小程序 / Docker / Nginx",
        ),
        (
            "7",
            "deploy/docker-compose.yml + travel_prediction_FIXED_20251206.sql",
            "容器化部署配置（含 LGTM 可观测性栈）与数据库 Schema（71 张表 / 约 2541 行 DDL）",
        ),
        (
            "8",
            "deploy/observability/（Grafana 仪表盘 + Prometheus + Tempo + Loki + Promtail）",
            "可观测性栈配置文件，docker compose up 即可一键启动",
        ),
        (
            "9",
            "experiments/results/weight_analysis.csv & privacy_impact.csv",
            "双流权重敏感性 + 差分隐私效用保留实验结果（可由 evaluate_models.py 一键复现）",
        ),
        (
            "10",
            "代码证据清单.md",
            "PPT / 概要表 / 设计文档中每条技术陈述对应的代码文件路径与行号",
        ),
        (
            "11",
            "运行网址 / 体验二维码",
            "Web 端：https://travel.dongsiwei.com/#/landing ；微信小程序 AppID：wx9569d09c12f8de06",
        ),
    ]

    tbl = _new_full_width_table(doc, rows=1 + len(files), cols=4)
    headers = ["序号", "文件名", "描述", "文件状态 / 版权状态"]
    widths = [1.4, 4.6, 7.4, 4.0]
    for c_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[c_idx].width = Cm(w)
    for c_idx, h in enumerate(headers):
        _label_cell(tbl.rows[0].cells[c_idx], h)
    for r_idx, (no, name, desc) in enumerate(files):
        row = tbl.rows[r_idx + 1]
        _set_cell_text(row.cells[0], no, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], name)
        _set_cell_text(row.cells[2], desc)
        _set_cell_text(row.cells[3], "■已上传到网盘 □未上传\n■自制 □开源 □获得授权")

    doc.add_paragraph()

    # ===== 承诺 =====
    tbl = _new_full_width_table(doc, rows=1, cols=1)
    _set_cell_text(
        tbl.rows[0].cells[0],
        (
            "【全体参赛队员郑重承诺】\n"
            "本作品全体参赛队员确认：本表所列内容是正式参赛内容的重要组成部分，已严格按照"
            "本大类参赛作品类别提交要求提交了评审必需的文档、数据等参赛材料，本表内容按照要求如实填写。"
            "如因提交的参赛材料不符合要求或本表填写内容不属实，将自愿承担因此导致奖项等级降低甚至终止本作品参加比赛的责任。\n\n"
            "全体参赛队员签名：（可附授权使用的电子签名图片）\n"
            "[队员1] / [队员2] / [队员3]\n\n"
            "日期：2026 年 4 月       日"
        ),
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "填写说明：① 所有 □ 可根据需要变化为 ■；② “作者及其分工比例”与“相关文件清单”可根据实际增减；"
        "③ 请将本表导出为 PDF 后上传到作品目录的“03设计与开发文档”子目录；"
        "④ 版权状态如有外部授权请注明授权方与来源地址。"
    )
    _set_cn_font(run, name="宋体", size_pt=10)


def main() -> None:
    for path in [OUT_SUBMISSION, OUT_ROOT]:
        doc = Document()
        build(doc)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print(f"[OK] 写入 {path} ({path.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
