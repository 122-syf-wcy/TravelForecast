#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""在已有答辩文档中追加：双流模型详细问题 + 系统整体架构问题"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('/Users/dongsiwei/TravelForecast/docs/挑战杯答辩问题准备.docx')

def add_section(title, questions):
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2A, 0x9D, 0x8F)
    for item in questions:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        p2 = doc.add_paragraph()
        run2 = p2.add_run('【参考答案】')
        run2.bold = True
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x2A, 0x9D, 0x8F)

        p3 = doc.add_paragraph(item['a'])
        p3.paragraph_format.left_indent = Inches(0.3)
        for r in p3.runs:
            r.font.size = Pt(11)

        doc.add_paragraph('')

# =====================================================================
# 新增一：双流融合模型（Dual-Stream）专项问题
# =====================================================================
add_section('十一、双流融合模型（Dual-Stream）专项', [
    {
        'q': '31. 什么是双流模型？你们为什么不直接用混合ARIMA+LSTM，而要专门设计双流结构？',
        'a': '双流模型（DualStreamHybridModel）的核心思想是：客流时间序列同时包含"线性趋势+周期性"（由ARIMA建模）和"非线性模式+多因素交叉影响"（由多变量LSTM建模），两者捕获的信息是互补的，而非重复的。\n\n'
             '与简单混合（直接平均）的区别：\n'
             '① 简单混合固定权重0.5，无法适应不同景区、不同季节的特性\n'
             '② 双流模型通过动态权重搜索，为每个景区自适应确定最优α（ARIMA权重）\n'
             '③ 双流LSTM流引入小波去噪预处理，降低历史数据噪声对非线性建模的干扰\n'
             '④ 支持特征屏蔽（Feature Masking），可分析单一因素对客流的独立影响'
    },
    {
        'q': '32. 双流模型的两个"流"分别是什么？各自负责什么？',
        'a': '流1——线性流（ARIMA流）：\n'
             '  · 模型：SARIMAX(1,1,1)×(1,1,1,7)，引入季节性参数s=7（周周期）\n'
             '  · 负责：捕捉客流的线性趋势（如节假日黄金周整体增长）和周期性规律（工作日vs周末）\n'
             '  · 输入：单变量历史客流序列\n\n'
             '流2——非线性流（LSTM流）：\n'
             '  · 模型：双层LSTM（64→32单元）+ Dropout(0.2) + Dense(16→1)\n'
             '  · 负责：捕捉多因素非线性交叉影响\n'
             '  · 输入：6维特征向量——[历史客流, 节庆指数, 是否周末, 天气类型, 温度, 景区海拔]\n'
             '  · 预处理：对历史客流进行小波去噪（Wavelet Denoising），去除采集噪声\n\n'
             '融合公式：final = α × ARIMA预测 + (1-α) × LSTM预测\n'
             '其中α通过网格搜索在验证集上动态确定（每个景区独立优化）'
    },
    {
        'q': '33. 动态权重是如何搜索的？具体算法是什么？',
        'a': '方法：验证集网格搜索（Grid Search on Validation Set）\n\n'
             '步骤：\n'
             '① 取最近7天历史数据作为验证集\n'
             '② ARIMA和LSTM分别对这7天做回测预测\n'
             '③ 遍历α从0.0到1.0（步长0.05，共21个候选值）\n'
             '④ 对每个α计算融合预测的MSE：MSE = Σ(α×ARIMA_i + (1-α)×LSTM_i - true_i)² / n\n'
             '⑤ 选取MSE最小对应的α作为最优权重\n\n'
             '结果：每个景区有独立的最优权重，如梅花山α可能倾向LSTM（特征丰富），水城古镇α可能倾向ARIMA（季节规律明显）\n\n'
             '优势：无需手工调参，自适应景区特性；每次重新训练时自动更新权重'
    },
    {
        'q': '34. 多变量LSTM中引入"海拔"这个特征有什么意义？',
        'a': '六盘水各景区海拔差异显著（梅花山1800m vs 明湖湿地1700m vs 乌蒙大草原2800m），海拔直接影响：\n'
             '① 气候适宜性：高海拔景区夏季凉爽，对避暑游客吸引力强，旺季集中在6-8月\n'
             '② 天气-客流交互效应：同样的降雨，高海拔景区（乌蒙草原）受影响更大\n'
             '③ 体力消耗：海拔影响游客游览时长和再访意愿\n\n'
             '在模型中，海拔作为静态特征（每个景区固定值）与其他动态特征拼接，让LSTM学习"不同海拔景区对天气/节假日的不同敏感度"，这是单变量模型无法捕捉的'
    },
    {
        'q': '35. 特征屏蔽（Feature Masking）是什么？有什么应用价值？',
        'a': '特征屏蔽是指在预测时将某些特征置为"中性值"（如天气置为0.5均值），通过对比有无该特征时的预测结果差异，量化该因素对客流的独立贡献。\n\n'
             '支持屏蔽的因素：weather（天气）、holiday（节假日/节庆）\n\n'
             'API参数：GET /api/prediction/{id}?factors=weather,holiday\n\n'
             '应用价值：\n'
             '① 景区管理者可了解"下雨天客流减少多少？" → 指导雨天优惠策略\n'
             '② "节假日相比平日增加多少？" → 指导节假日备勤方案\n'
             '③ "去掉所有外部因素，仅靠历史趋势预测是多少？" → 评估自然增长基线\n'
             '④ 这也支撑了前端"政策沙盘"功能——模拟调整某因素后的客流变化'
    },
    {
        'q': '36. 双流模型的准确率92%是怎么来的？比单流高在哪里？',
        'a': '各模型准确率（MAPE指标，1-MAPE）：\n'
             '  · ARIMA单流：82%（擅长线性趋势，遇到节假日突变时误差大）\n'
             '  · LSTM单流：87%（能捕捉非线性，但对短期突变反应滞后）\n'
             '  · 双流融合：92%（两流互补：ARIMA提供稳定基线，LSTM修正非线性残差）\n\n'
             '双流优势体现在：\n'
             '① 节假日前后：ARIMA捕捉"节假日普遍增多"的周期规律，LSTM捕捉"具体增幅受天气影响的非线性"部分\n'
             '② 异常天气：LSTM多变量特征能快速响应天气骤变，ARIMA提供稳定的趋势基准避免过拟合\n'
             '③ 动态权重使模型在不同景区、不同预测窗口下始终保持最优组合'
    },
])

# =====================================================================
# 新增二：系统整体架构专项问题
# =====================================================================
add_section('十二、系统整体架构专项', [
    {
        'q': '37. 请画出（或描述）系统整体架构图，各层之间如何交互？',
        'a': '系统分四层架构：\n\n'
             '【接入层】\n'
             '  · Nginx（80端口）：统一流量入口，静态文件服务、反向代理、SSL终止\n'
             '  · 路由规则：/api/ → 8080主后端，/ai-api/ → 8081 AI后端，/miniprogram-api/ → 8082小程序后端，/prediction-api/ → 8001预测服务，/digital-human-api/ → 8083数字人，/ws/ → WebSocket代理\n\n'
             '【网关层】\n'
             '  · Spring Cloud Gateway（8888）：JWT认证、Redis令牌桶限流、Resilience4j熔断、CORS\n\n'
             '【服务层（5个微服务）】\n'
             '  · 主后端(8080)：Java Spring Boot，负责景区/用户/商户/订单等核心业务\n'
             '  · AI后端(8081)：Java Spring Boot，通义千问对话、RAG知识库、TTS、行程规划\n'
             '  · 小程序后端(8082)：Java Spring Boot，微信登录、文创商城、小程序专属API\n'
             '  · 预测服务(8001)：Python FastAPI，ARIMA+双流模型客流预测\n'
             '  · 数字人后端(8083)：Python FastAPI，WebSocket实时对话、TTS语音合成\n\n'
             '【数据层】\n'
             '  · MySQL 8.0（3306）：核心业务数据，30+张表\n'
             '  · Redis 6.0（6379）：限流计数、API缓存、会话Token、LLM/TTS缓存\n'
             '  · 阿里云OSS：图片/视频/文件存储，通过/api/oss/proxy代理访问\n\n'
             '【客户端层】\n'
             '  · Web前端（Vue3）：游客端+商户端+管理员端，三角色统一代码库\n'
             '  · 微信小程序（Uni-app）：游客移动端，微信OAuth登录+支付'
    },
    {
        'q': '38. 为什么同时有API网关(8888)和Nginx？职责有什么区别？',
        'a': 'Nginx和API网关职责不同、不重复：\n\n'
             'Nginx负责：\n'
             '  ① 静态文件服务（Vue打包后的HTML/CSS/JS）\n'
             '  ② 反向代理（将外网80端口请求分发到内网各服务端口）\n'
             '  ③ SSL/HTTPS终止（生产环境）\n'
             '  ④ 静态资源缓存（30天）\n'
             '  ⑤ 处理跨域（CORS头）\n\n'
             'Spring Cloud Gateway负责：\n'
             '  ① 业务层JWT Token认证（识别用户身份和角色）\n'
             '  ② 基于Redis的API限流（业务级QPS控制）\n'
             '  ③ 服务熔断（某后端异常时快速失败）\n'
             '  ④ 动态路由配置（可运行时修改）\n'
             '  ⑤ 请求日志（业务审计）\n\n'
             '两者形成双层防护：Nginx在网络层拦截非法流量，Gateway在业务层做权限和限流控制'
    },
    {
        'q': '39. 三端（游客/商户/管理员）是如何在同一套前端代码中实现隔离的？',
        'a': '前端通过路由和权限守卫实现三端隔离：\n\n'
             '① 路由命名空间：游客页面 /user/*，商户页面 /business/*，管理员 /admin/*\n'
             '② 路由守卫（beforeEach）：读取用户角色(user.role)，根据角色重定向到对应路由，越权访问自动跳转403\n'
             '③ 状态管理（Pinia）：userStore统一管理登录状态、角色信息、JWT Token\n'
             '④ 布局分离：三端使用不同的Layout组件（侧边栏/导航栏样式不同）\n'
             '⑤ API权限：后端对每个接口通过RoleInterceptor校验角色，前端请求即使越权也会被后端拦截\n\n'
             '角色存储：users表role字段（USER/MERCHANT/ADMIN），JWT payload中携带role信息，无需每次查库'
    },
    {
        'q': '40. JWT Token的设计方案是什么？如何处理Token过期？',
        'a': 'JWT双Token机制：\n\n'
             'AccessToken：\n'
             '  · 有效期：2小时（短期），携带userId、role、email等claims\n'
             '  · 存储：前端localStorage\n'
             '  · 用途：每次API请求Authorization: Bearer {token}\n\n'
             'RefreshToken：\n'
             '  · 有效期：7天（长期），存储于Redis（key=refreshToken，value=userId）\n'
             '  · 用途：AccessToken过期时，用RefreshToken换取新AccessToken\n\n'
             '过期处理流程：\n'
             '  ① Axios响应拦截器捕获401错误\n'
             '  ② 自动调用/api/auth/refresh接口，携带RefreshToken\n'
             '  ③ 成功则更新本地AccessToken，重试原请求\n'
             '  ④ RefreshToken也过期则跳转登录页\n\n'
             '安全措施：RefreshToken存Redis可主动吊销（退出登录时删除Redis记录），防止Token泄露'
    },
    {
        'q': '41. 系统中的缓存策略是怎么设计的？用了哪几层缓存？',
        'a': '四层缓存体系：\n\n'
             '① 浏览器缓存：\n'
             '  · 静态资源（JS/CSS/图片）Nginx配置Cache-Control: max-age=2592000（30天）\n'
             '  · OSS代理图片：Cache-Control: public, max-age=86400（1天）\n\n'
             '② 前端内存缓存（requestCache）：\n'
             '  · cachedRequest工具，60秒内相同请求不重发\n'
             '  · Landing页API数据10分钟内存缓存\n\n'
             '③ Redis服务端缓存：\n'
             '  · 限流计数：滑动窗口计数器\n'
             '  · LLM回答缓存：相同问题直接返回缓存（TTL 24小时）\n'
             '  · TTS音频缓存：相同文本直接返回缓存音频（TTL 24小时）\n'
             '  · 预测结果缓存：相同参数预测结果缓存1小时\n\n'
             '④ 应用内存缓存：\n'
             '  · Python预测服务：训练好的模型权重常驻内存（避免每次请求重新加载）\n'
             '  · ARIMA模型字典：scenic_id → 已训练的SARIMAX模型对象'
    },
    {
        'q': '42. 数据库中最复杂的表关系是什么？如何设计的？',
        'a': '最复杂的是"景区-商户-用户"三方关联体系：\n\n'
             'scenic_spots（景区主表）\n'
             '  ↓ 1:N\n'
             'scenic_sub_spots（子景点）、scenic_images（图片）、scenic_videos（视频）、scenic_facilities（设施）、scenic_activities（活动）\n\n'
             'merchant_profiles（商户）← N:M →scenic_spots（商户可管理多景区，景区可绑定多商户）\n'
             '  · 中间表：merchant_scenics（商户-景区关联）\n\n'
             'users（用户）：\n'
             '  · user_favorites（收藏）→ scenic_spots\n'
             '  · ticket_orders（订单）→ scenic_spots\n'
             '  · reviews（评价）→ scenic_spots\n\n'
             '设计原则：\n'
             '① 景区主表与媒体资源分离（避免大字段拖慢查询）\n'
             '② 订单表独立（高写入频率，避免锁表）\n'
             '③ 操作日志表独立（只追加、不更新，适合日志场景）\n'
             '④ 合理冗余：scenic_spots.rating冗余计算值，避免每次聚合查询'
    },
    {
        'q': '43. 如果让你重新设计这套系统，有什么会做不同？',
        'a': '（这类问题体现反思能力，建议真诚回答，以下供参考）\n\n'
             '① 消息队列：订单创建、通知发送、日志记录等异步场景引入RabbitMQ/Kafka，减少同步等待\n'
             '② 服务注册与发现：引入Nacos/Consul，替代现有静态Nginx配置，支持动态扩缩容\n'
             '③ 分布式链路追踪：引入Sleuth+Zipkin，8个微服务的请求链路难以排查问题\n'
             '④ 数据库分库分表：预留设计，用户行为日志等高增长表提前规划分表策略\n'
             '⑤ 前端状态管理：部分组件数据管理混乱，更严格地划分Pinia store职责\n'
             '⑥ 测试覆盖：补充单元测试和集成测试（目前几乎没有），TDD先写测试再写实现'
    },
])

# =====================================================================
# 新增三：数字人与WebSocket专项
# =====================================================================
add_section('十三、数字人 & WebSocket 专项', [
    {
        'q': '44. 数字人的实时对话是如何通过WebSocket实现的？整个链路描述一下。',
        'a': '完整链路（7步）：\n\n'
             '① 用户在小程序/Web输入文字（或语音→STT转文字）\n'
             '② 前端通过WebSocket发送JSON消息到 ws://39.97.232.141/ws/avatar\n'
             '③ Nginx将WebSocket请求代理到数字人后端8083（proxy_http_version 1.1 + Upgrade头）\n'
             '④ Python FastAPI WebSocket处理器接收消息\n'
             '⑤ 查询Redis LLM缓存 → 命中直接返回；未命中 → 调用通义千问API生成回复文本\n'
             '⑥ 查询Redis TTS缓存 → 命中直接返回音频；未命中 → 调用TTS合成音频（MP3/WAV）\n'
             '⑦ WebSocket分两步推送：先发text_output（文字）→ 再发bytes（音频二进制）→ 再发viseme（口型数据）→ 最后发status:idle\n\n'
             '前端接收后：AudioContext解码音频 → 播放音频 → 同步驱动3D人物口型动画'
    },
    {
        'q': '45. 音频播放时如何实现与3D数字人口型的同步？',
        'a': 'Lipsync口型同步实现：\n\n'
             '① 服务端发送viseme数据：根据文本中每个音素（phoneme）估算对应的口型帧（viseme），附带时间戳\n'
             '② 前端AudioContext播放音频，同时启动口型动画定时器\n'
             '③ 3D模型（Three.js + VRM格式）预定义口型BlendShape（A/I/U/E/O等口型权重）\n'
             '④ 按viseme时间轴逐帧插值更新BlendShape权重\n'
             '⑤ requestAnimationFrame驱动渲染循环，保持音频与动画同步\n\n'
             '降级方案：若WebSocket连接失败，回退到录制好的视频（静态视频+本地TTS）'
    },
])

output_path = '/Users/dongsiwei/TravelForecast/docs/挑战杯答辩问题准备.docx'
doc.save(output_path)
print(f'文档已更新：{output_path}')
print(f'新增问题：双流模型6题 + 系统架构7题 + 数字人2题 = 15题')
print(f'总计问题数：30 + 15 = 45题')
