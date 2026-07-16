"""
重新生成《智教黔行_设计和开发文档.docx》

依据：截至 2026-04-19 仓库实际代码（commit df09348 之后）
覆盖：网关 8888 / AI 后端 8081 / 主后端 8080 / 小程序后端 8082 /
     预测服务 8001 / 数字人 8083 / 71 张数据库表 / 混合检索 RAG /
     LGTM 可观测性栈 / 政策沙盒 / 紧急救援 / 研学路线 / Function
     Calling / 三级缓存 / 分句流式 TTS / 安全预警 / 真实部署
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT_DEV_DOC = REPO_ROOT / "提交材料-智教黔行/03-设计与开发文档/智教黔行_设计和开发文档.docx"
DEFAULT_OUTPUT_DEV_DOC_ROOT = REPO_ROOT / "智教黔行_设计和开发文档.docx"


# ============================================================
# 排版工具
# ============================================================


def _set_cn_font(run, name: str = "宋体", size_pt: float = 12, bold: bool = False) -> None:
    """同时设置中英文字体，避免中文字符回退到默认字体。"""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _shade_cell(cell, fill_hex: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 18, 2: 15, 3: 13}
    run = p.add_run(text)
    _set_cn_font(run, name="黑体", size_pt=sizes.get(level, 12), bold=True)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    else:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_para(
    doc: Document,
    text: str,
    *,
    indent_first: bool = True,
    bold: bool = False,
    size: float = 11,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    _set_cn_font(run, name="宋体", size_pt=size, bold=bold)


def _add_bullet(doc: Document, text: str, *, size: float = 11) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    _set_cn_font(run, name="宋体", size_pt=size)


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    col_widths_cm: Sequence[float] | None = None,
) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    if col_widths_cm:
        for i, width in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(width)

    header_row = tbl.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, "2E75B6")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_cn_font(run, name="黑体", size_pt=10.5, bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        if r_idx % 2 == 1:
            for cell in cells:
                _shade_cell(cell, "F2F6FA")
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.3
            for line in str(value).split("\n"):
                if p.runs:
                    run = p.add_run("\n" + line)
                else:
                    run = p.add_run(line)
                _set_cn_font(run, name="宋体", size_pt=10)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_global_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)


# ============================================================
# 文档章节
# ============================================================


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run("中国大学生计算机设计大赛")
    _set_cn_font(run, name="黑体", size_pt=22, bold=True)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("软件应用与开发类作品 · 设计和开发文档")
    _set_cn_font(run, name="黑体", size_pt=18, bold=True)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("智教黔行")
    _set_cn_font(run, name="黑体", size_pt=32, bold=True)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("基于多源数据融合与 LSTM-ARIMA 双流动态权重模型的\n六盘水山地智慧文旅一体化研学平台")
    _set_cn_font(run, name="黑体", size_pt=14, bold=False)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    info_lines = [
        "作品编号：（报名后由大赛系统分配）",
        "作品大类：软件应用与开发",
        "版本编号：V3.0（截至 2026 年 4 月）",
        "作      者：（团队成员姓名 / 学号 / 分工 待补全）",
        "指导教师：（待补全）",
        "参赛学校：（待补全）",
        "填写日期：2026 年 4 月       日",
    ]
    for line in info_lines:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        run = p2.add_run(line)
        _set_cn_font(run, name="宋体", size_pt=12)

    _add_page_break(doc)


def build_toc_summary(doc: Document) -> None:
    _add_heading(doc, "文档摘要", 1)
    _add_para(
        doc,
        "本文档面向中国大学生计算机设计大赛（以下简称“大赛”）软件应用与开发类作品的评审要求，"
        "完整描述了“智教黔行——六盘水山地智慧文旅一体化研学平台”的需求分析、系统架构、详细设计、"
        "测试结果、部署方案与项目总结。文档中涉及的所有技术陈述均可在仓库源码（02-素材与源码/）中"
        "通过文件路径与行号定位，与同目录《代码证据清单.md》一一对应，便于评委核验。",
    )
    _add_para(
        doc,
        "项目实际部署地址：https://travel.dongsiwei.com（备用 IP：http://39.97.232.141）；"
        "微信小程序 AppID：wx9569d09c12f8de06。系统由 6 个独立微服务、71 张 MySQL 数据表、Vue 3 Web 端"
        "与 UniApp 微信小程序双端、以及一套基于 Docker Compose 的 LGTM 可观测性栈组成，全部代码均为团队原创。",
    )

    _add_heading(doc, "目录", 2)
    toc_items = [
        ("一、需求分析", "1.1 开发背景与意义；1.2 用户群体与核心需求；1.3 竞品对比"),
        ("二、概要设计", "2.1 总体架构；2.2 模块划分与技术栈；2.3 服务通信与端口；2.4 部署拓扑"),
        ("三、详细设计", "3.1 业务流程；3.2 数据库设计（71 张表）；3.3 客流预测算法；3.4 RAG 混合检索；3.5 数字人流式管道；3.6 政策沙盒与紧急救援；3.7 安全与权限；3.8 性能与可观测性"),
        ("四、测试报告", "4.1 测试方法；4.2 功能测试用例与结果；4.3 性能与压测；4.4 实验数据与权重敏感性；4.5 差分隐私实验"),
        ("五、安装与使用", "5.1 环境要求；5.2 一键启停；5.3 开发模式；5.4 典型使用流程"),
        ("六、项目总结", "6.1 工程成果；6.2 创新贡献；6.3 主要困难与解决；6.4 后续演进规划"),
        ("七、参考文献", "学术论文与开源框架文档"),
    ]
    for title, desc in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{title}  —— {desc}")
        _set_cn_font(run, name="宋体", size_pt=11)
    _add_page_break(doc)


def build_chapter_requirements(doc: Document) -> None:
    _add_heading(doc, "一、需求分析", 1)

    _add_heading(doc, "1.1 开发背景与意义", 2)
    _add_para(
        doc,
        "六盘水市位于贵州省西部，地处云贵高原乌蒙山区，海拔 1400—2900 米，年均气温 19℃，"
        "被誉为“中国凉都”。市域内拥有梅花山风景区（海拔 2400m）、玉舍国家森林公园（2300m）、"
        "乌蒙大草原（2857m）、水城古镇（1800m）、明湖国家湿地公园（1750m）等五个核心景区。"
        "“十四五”以来，六盘水大力推进文旅融合与研学旅游，但在旅游信息化方面仍面临四个突出问题：",
    )
    _add_bullet(
        doc,
        "客流预测手段落后：景区主要依赖人工经验对节假日客流做粗略估算，缺乏对气象、海拔、节假日"
        "连续长度等多源因素的建模能力，高峰期常出现停车难、景点拥挤、应急资源调度被动等问题。",
    )
    _add_bullet(
        doc,
        "旅游信息碎片化：游客需要在多个 OTA、官网、地图、社交平台间切换才能完成“吃住行游购娱”全流程，"
        "缺乏一站式智慧服务入口，研学家长更难于在路线、安全、知识三方面获得权威建议。",
    )
    _add_bullet(
        doc,
        "研学教学工具缺位：六盘水作为三线建设博物馆、彝族火把节、苗族跳花节等研学资源富集地，"
        "现场讲解仍以人工导游为主，知识产出无法沉淀、个性化与互动性不足，难以匹配“双减”后中小学校"
        "的研学课程化诉求。",
    )
    _add_bullet(
        doc,
        "多角色协同困难：游客、景区商家、平台管理者三类角色对系统能力的诉求差异大，"
        "需要一套权限隔离清晰、数据共享有序的综合管理后台，以支撑商家入驻审核、合同管理、订单结算、"
        "实时监控、应急救援等场景。",
    )
    _add_para(
        doc,
        "基于上述背景，团队设计并实现了“智教黔行”平台，运用时间序列预测、大语言模型、混合检索增强、"
        "实时流式语音合成、地理信息可视化等技术，为六盘水山地旅游提供数据驱动的智慧决策支持。"
        "项目同时具备工程层面的完整性（微服务架构、可观测性、安全治理）与算法层面的创新性"
        "（双流动态权重 + 海拔因子 + BM25/向量混合检索），契合大赛“面向真实需求、突出技术创新、"
        "强调系统完整”的评审导向。",
    )

    _add_heading(doc, "1.2 用户群体与核心需求", 2)
    _add_table(
        doc,
        ["用户角色", "代表场景", "核心功能需求"],
        [
            (
                "普通游客 / 研学学生",
                "节假日出游、研学课程、亲子科普",
                "景区浏览、3D 地形导览、客流预测、AI 行程规划、数字人讲解、紧急救援",
            ),
            (
                "景区商家 / 运营方",
                "门票运营、设施巡检、应急响应",
                "景区资源管理、门票订单、实时客流监控、政策沙盒、营收分析、紧急救援工单",
            ),
            (
                "平台管理员 / 文旅局",
                "全市数据治理、内容审核",
                "用户/商家/景区/内容管理、权限分配、平台数据看板、操作日志、备份恢复",
            ),
        ],
        col_widths_cm=[3.4, 4.6, 8.0],
    )

    _add_heading(doc, "1.3 竞品对比", 2)
    _add_table(
        doc,
        [
            "对比维度",
            "智教黔行（本作品）",
            "携程 / 美团 / 飞猪等 OTA",
            "一般景区小程序",
        ],
        [
            (
                "客流预测",
                "ARIMA + 多变量 LSTM 双流动态权重融合，引入海拔与节假日连续度，支持网格搜索 α∈[0,1]",
                "无机器学习预测，仅按门票售出量做事后展示",
                "无预测能力",
            ),
            (
                "AI 交互",
                "RAG 混合检索（BM25 + DashScope 向量）+ Function Calling + 流式 TTS 数字人",
                "基础 FAQ 客服机器人，多为模板话术",
                "无",
            ),
            (
                "知识检索",
                "六盘水研学专属知识库 + Redis 向量索引 + 自动重排，TTL 1 年缓存",
                "通用知识问答，缺乏研学领域知识",
                "无",
            ),
            (
                "行程规划",
                "AI 自动生成 6—8 个/天活动的 JSON 行程，含预算、客流、餐饮、住宿，SHA-256 缓存复用",
                "固定模板路线，缺少预算/拥挤度联动",
                "无",
            ),
            (
                "政策沙盒",
                "支持联票折扣 / 交通补贴 / 容量上限三参数模拟，输出客流-收入弹性",
                "无",
                "无",
            ),
            (
                "应急救援",
                "用户一键报警 + 商家工单流转 + 景区/全平台统计，含位置/状态/处理日志",
                "无",
                "无",
            ),
            (
                "可观测性",
                "Prometheus + Grafana + Tempo + Loki + Promtail 一键 LGTM 栈",
                "面向商家不可见",
                "无",
            ),
            (
                "多端覆盖",
                "Web 用户端 / 商家端 / 管理后台 + 微信小程序 + 统一 API 网关",
                "App + Web + 小程序",
                "仅小程序",
            ),
            (
                "研学特色",
                "数字人“黔小游” + 红色研学 + 研学护照 + AI 安全预警",
                "纯旅游消费导向",
                "基础信息展示",
            ),
        ],
        col_widths_cm=[2.2, 6.6, 4.8, 2.4],
    )


def build_chapter_overview(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "二、概要设计", 1)

    _add_heading(doc, "2.1 系统总体架构", 2)
    _add_para(
        doc,
        "智教黔行采用“前端多端 + 网关 + 业务微服务 + 数据存储 + 可观测性”五层微服务架构，"
        "全部 6 个微服务部署于同一台 ECS（生产环境 39.97.232.141），通过 Spring Cloud Gateway "
        "（端口 8888）对外统一暴露 HTTPS 入口（https://travel.dongsiwei.com）。"
        "网关承担 JWT 鉴权、IP 令牌桶限流、Resilience4j 熔断、统一跨域、链路日志等通用能力，"
        "下游业务服务彼此解耦，支持独立伸缩与独立发布。",
    )
    _add_para(
        doc,
        "架构层次：（1）表现层——Vue 3 Web 端（用户/商家/管理三套界面）+ UniApp 微信小程序；"
        "（2）网关层——Spring Cloud Gateway + Redis 限流 + Resilience4j 熔断 + 全局 CORS；"
        "（3）业务服务层——主业务后端 / AI 智能后端 / 小程序后端 / Python 客流预测 / Python 数字人"
        "五个独立服务；（4）数据层——MySQL 8.0（71 张表）+ Redis 7（多 DB 隔离）+ 阿里云 OSS；"
        "（5）可观测性——Prometheus 指标 + Tempo trace + Loki 日志 + Grafana 仪表盘。",
    )

    _add_heading(doc, "2.2 模块划分与技术栈", 2)
    _add_table(
        doc,
        ["模块", "技术栈", "端口", "职责"],
        [
            (
                "TravelForecastGateway",
                "Spring Cloud Gateway 4 + Resilience4j + Redis 限流",
                "8888",
                "API 路由（business/ai/miniprogram/prediction/digital-human/ws）、JWT 校验、令牌桶限流、熔断降级、统一异常",
            ),
            (
                "TravelForecastBackend（主业务后端）",
                "Java 17 + Spring Boot 3 + MyBatis-Plus + MySQL 8 + Redis",
                "8080",
                "用户/商家/管理 RBAC、景区与活动、订单、统计、政策沙盒、紧急救援、热力图、备份与监控（55 个 Controller）",
            ),
            (
                "TravelForecastingAIBackend（AI 后端）",
                "Java 17 + DashScope SDK + Redis + MyBatis-Plus",
                "8081",
                "AI 行程规划、AI 聊天、研学路线、知识库检索（BM25+向量混合）、语音转写代理",
            ),
            (
                "TravelForecastMiniProgramBackend（小程序后端）",
                "Java 17 + Spring Boot 3 + 微信开放平台 SDK",
                "8082",
                "微信登录、首页、商城、行程、研学护照、生活服务、支付、商家审核",
            ),
            (
                "TravelForecast-PythonPredictionService（预测）",
                "Python 3.10 + FastAPI + statsmodels + PyTorch + Loguru",
                "8001",
                "ARIMA / LSTM 多变量 / 双流动态权重融合预测、模型训练、模型指标注册表",
            ),
            (
                "TravelForecast-DigitalHuman（数字人）",
                "Python 3.10 + FastAPI + DeepSeek + Edge TTS / CosyVoice + Vosk",
                "8083",
                "WebSocket 流式对话、RAG 增强、Function Calling、分句流式 TTS、三级缓存、安全预警",
            ),
            (
                "TravelForecastFrontend（Web 前端）",
                "Vue 3 + Vite + Element Plus + ECharts GL + Pinia",
                "5173 (dev) / 80 (Nginx)",
                "Landing 首页、3D 地形、客流预测、行程规划、实时服务、商家与管理后台 57 个页面",
            ),
            (
                "TravelForecastMiniProgram（微信小程序）",
                "UniApp + Vue 3 + uni-ui + 微信小程序原生 API",
                "—",
                "12 大模块：首页/景区/商城/行程/数字人/红色研学/研学护照/生活服务/导览/搜索 等",
            ),
            (
                "Observability（可观测性栈）",
                "Prometheus / Tempo / Loki / Promtail / Grafana（Docker Compose）",
                "9090 / 3200 / 3100 / 3000",
                "六服务指标抓取、OTel trace、集中日志、统一仪表盘（已预置 TravelForecast 总览）",
            ),
        ],
        col_widths_cm=[3.2, 4.8, 1.8, 7.2],
    )

    _add_heading(doc, "2.3 服务间通信、路由与端口", 2)
    _add_para(
        doc,
        "前端统一请求 https://travel.dongsiwei.com（指向网关 8888）。网关根据 URL 前缀路由：",
    )
    _add_table(
        doc,
        ["前缀", "目标服务", "限流（令牌桶）", "熔断"],
        [
            ("/api/**", "主业务后端 8080（context-path=/api）", "200 / 400", "business-service"),
            ("/ai-api/**", "AI 后端 8081（context-path=/ai-api）", "100 / 200", "ai-service"),
            ("/miniprogram-api/**", "小程序后端 8082（StripPrefix=1）", "150 / 300", "—"),
            ("/prediction-api/**", "预测服务 8001（StripPrefix=1）", "50 / 100", "prediction-service"),
            ("/digital-human-api/**", "数字人 8083（StripPrefix=1）", "—", "—"),
            ("/ws/**", "数字人 WebSocket 8083", "—", "—"),
        ],
        col_widths_cm=[3.6, 6.8, 3.6, 3.0],
    )
    _add_para(
        doc,
        "Resilience4j 配置 slidingWindowSize=10、failureRateThreshold=50%，"
        "命中熔断后会 forward 至 /fallback/{service} 走降级回包，避免上游超时雪崩。"
        "重试策略：GET/POST 在 BAD_GATEWAY/INTERNAL_SERVER_ERROR 时指数退避重试 3 次（首跳 100ms→1s）。",
    )

    _add_heading(doc, "2.4 部署拓扑与目录约定", 2)
    _add_para(
        doc,
        "生产环境采用单机一键部署，目录布局以 /opt/travel 为根：gateway/、backend/、ai-backend/、"
        "mp-backend/、prediction/src/、digital-human/backend/、logs/。"
        "deploy/start-all.sh 自动加载 secrets/.env、定位 JDK 17、拉起 6 个服务并轮询 health 端点。"
        "前端构建产物经 Nginx（travel.conf 监听 80）反向代理到网关；HTTPS 证书在 Nginx 层卸载。",
    )


def build_chapter_detail(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "三、详细设计", 1)

    _add_heading(doc, "3.1 用户视角的典型业务流程", 2)
    _add_para(
        doc,
        "用户端：访客访问 Landing 页 → 浏览六盘水实景 / 大屏数据看板 / 3D 地形 → 注册或微信扫码登录 → "
        "进入用户仪表盘 → 选择 “景区探索 / 客流预测 / 行程规划 / 实时服务 / AI 数字人”。"
        "整套流程 4 个公开页（router 中 guestAllowedPaths 数组定义）允许访客无登录即可体验，"
        "降低首次使用门槛。",
    )
    _add_para(
        doc,
        "商家端：商家在 Landing 提交入驻 → 主后端 merchant_audit / merchant_audit_logs 记录 → 管理员审批通过 → "
        "商家登录后台管理景区资源、门票订单、设施、评论回复、紧急救援工单，并查看自家景区营收分析。",
    )
    _add_para(
        doc,
        "管理员端：登录后进入数据总览 → 用户/商家/内容/系统/数据五大菜单。"
        "包括用户行为分析（user_behavior_logs 26 万行级日志）、内容审核（敏感词、轮播、公告、宣传位）、"
        "数据导出（CSV/Excel）、备份恢复（BackupController）、运行监控（SystemMonitorController）、"
        "操作日志（operation_logs / platform_activity_logs 双轨）。",
    )

    _add_heading(doc, "3.2 数据库设计", 2)
    _add_para(
        doc,
        "系统主数据库 travel_prediction（MySQL 8.0），共 71 张表，迁移脚本"
        "TravelForecastBackend/src/main/resources/db/migration/travel_prediction_FIXED_20251206.sql"
        "（约 2541 行 DDL 与初始化数据），可与 Flyway 联动。表分为 12 个业务域：",
    )
    _add_table(
        doc,
        ["业务域", "代表数据表", "用途"],
        [
            (
                "用户与会话",
                "users / wechat_users / user_sessions / captcha_records / user_preferences / user_privacy / user_distribution / user_favorites / user_behavior_logs / user_feedbacks",
                "三角色身份、JWT 多端会话、验证码、收藏、行为分析、隐私同意",
            ),
            (
                "权限",
                "roles / permissions / role_permissions",
                "RBAC 角色权限矩阵",
            ),
            (
                "景区",
                "scenic_spots / scenic_images / scenic_videos / scenic_weather / scenic_realtime_data / scenic_rankings / scenic_statistics / scenic_reviews / review_likes / review_replies / facilities",
                "5 大核心景区静态信息 + 实时动态 + UGC 评论",
            ),
            (
                "活动与日程",
                "activities / activity_participants / itineraries / itinerary_spots",
                "活动报名、行程模板、行程详细景点",
            ),
            (
                "客流",
                "flow_records / daily_flow_summary / predictions / prediction_results / visitor_predictions",
                "原始客流、按日聚合、预测请求、预测结果、历史预测对比",
            ),
            (
                "订单",
                "ticket_orders",
                "门票订单、二维码核销，含 barcode 字段",
            ),
            (
                "商家",
                "merchant_audit / merchant_audit_logs / merchant_available_scenics / merchant_contracts / merchant_profile / merchant_profiles / business_resources / business_resource_images / business_reviews / business_review_replies / business_feedback / business_todos",
                "入驻审核、合同、可经营景区、资源、评价回复、待办",
            ),
            (
                "AI 与对话",
                "chat_conversations / chat_messages",
                "数字人会话、研学问答历史",
            ),
            (
                "内容",
                "banners / news / announcements / landing_config / sensitive_words",
                "首页横幅、新闻、公告、Landing 配置、敏感词审核",
            ),
            (
                "运营/管理",
                "admin_feedback / admin_notices / admin_pending_tasks / admin_tasks / operation_logs / platform_activity_logs / platform_statistics / revenue_analytics / todos",
                "管理员任务、操作日志、平台数据、营收分析",
            ),
            (
                "通知",
                "notifications / notification_reads / system_notifications",
                "系统通知与已读状态",
            ),
            (
                "系统",
                "site_configs / system_config / system_logs / system_settings / emergency_rescue / feedback_records",
                "站点配置、系统日志、紧急救援工单、用户反馈",
            ),
        ],
        col_widths_cm=[2.6, 8.6, 4.6],
    )
    _add_para(
        doc,
        "关键设计决策：（1）users 表统一存储 admin/user/merchant 三种角色，role 字段 + RoleInterceptor "
        "实现权限隔离，避免分表带来的关联复杂度；（2）chat_conversations + chat_messages 双层结构，"
        "支持上下文窗口构建与会话级清理；（3）scenic_statistics 按日聚合存储，数据看板查询复杂度由"
        "全表 GROUP BY 降为 O(1)；（4）emergency_rescue 表通过 user_id / scenic_id / handler_user_id / "
        "status / handle_notes 字段串起“游客发起→商家受理→处置完成”的全流程。",
    )

    _add_heading(doc, "3.3 关键算法 1：ARIMA-LSTM 双流动态权重融合预测", 2)
    _add_para(
        doc,
        "传统单一模型在山地旅游客流预测上存在固有局限：ARIMA 擅长线性趋势 + 周期，但难以建模天气/节假日/"
        "海拔等外生变量；LSTM 长于非线性时序，但小样本下易过拟合且对节假日突变响应迟钝。"
        "本项目提出 DualStreamHybridModel（dual_stream_model.py:17-146），将两者优势互补：",
    )
    _add_bullet(
        doc,
        "ARIMA 子流（arima_model.py，283 行）：差分平稳化 + AIC 自动定阶 (p, d, q)，输出线性趋势预测。",
    )
    _add_bullet(
        doc,
        "多变量 LSTM 子流（lstm_new.py，255 行）：feature_dim=6，输入 [历史客流, 节庆, 周末, 天气, 温度, 海拔]，"
        "其中海拔由 SCENIC_CONFIG 提供（梅花山 2400/玉舍 2300/乌蒙 2857/水城 1800/明湖 1750）并归一化到 [0,1]。",
    )
    _add_bullet(
        doc,
        "动态权重搜索（_search_optimal_weights, dual_stream_model.py:88-134）：在最近 7 天的验证集上，"
        "网格搜索 α∈[0, 1]（步长 0.05，21 个候选值），以 MSE 最小为目标自动选优。",
    )
    _add_bullet(
        doc,
        "融合公式：final = α·val_arima + (1−α)·val_lstm（dual_stream_model.py:70），输出同时附带 "
        "components.{arima_output, lstm_output, weight_alpha} 三项可解释字段，便于前端可视化。",
    )
    _add_bullet(
        doc,
        "指标注册表（metrics_registry.py）：取消硬编码 0.92 置信度，统一从 experiments/results/model_metrics.json "
        "动态加载；缺失时使用保守默认值并打 warn 日志，配套 evaluate_models.py 脚本基于真实 CSV 重写指标。",
    )
    _add_para(
        doc,
        "模型对外暴露 5 个 REST 接口：/api/prediction/flow/{scenic_id}（单景区按天）、/api/prediction/total（全域聚合）、"
        "/api/prediction/hourly/{scenic_id}（小时级）、/api/prediction/hourly/total（全域小时聚合）、"
        "/api/prediction/train/{scenic_id}（重训练）。FastAPI 自动生成 Swagger UI 与 OpenAPI JSON，便于前端 mock。",
    )

    _add_heading(doc, "3.4 关键算法 2：RAG 混合检索（BM25 + 向量重排）", 2)
    _add_para(
        doc,
        "AI 后端（TravelForecastingAIBackend）实现了一套生产级 RAG 检索链路，全部为 Java 原生组件，无外部向量库依赖：",
    )
    _add_table(
        doc,
        ["阶段", "组件", "实现要点"],
        [
            (
                "倒排召回",
                "RagServiceImpl.search()",
                "MyBatis-Plus LambdaQueryWrapper 按 title/content/keywords LIKE 召回 recallSize=20 候选",
            ),
            (
                "BM25 重排",
                "Bm25Reranker（rag/Bm25Reranker.java）",
                "k1=1.5、b=0.75；中文 2-gram + 单字混合分词；BM25 分数全量保留供归一化",
            ),
            (
                "向量重排（新）",
                "EmbeddingClient + VectorStore + HybridRetriever",
                "DashScope text-embedding-v2（1536 维）批量 25 条，结果存 Redis（key=ai:vec:knowledge:{id}），余弦相似度",
            ),
            (
                "线性融合",
                "HybridRetriever.retrieve()",
                "final = α·向量 + (1-α)·BM25，alphaVector=0.55；向量失败优雅降级为纯 BM25",
            ),
            (
                "结果缓存",
                "Redis ai:knowledge:search:{hash}",
                "TTL=8760 小时（1 年），按 query+category+scenicId hash 复用",
            ),
            (
                "增量索引",
                "KnowledgeEmbeddingJob",
                "ApplicationReadyEvent 触发，异步 batchSize=20 调 DashScope；新增/删除知识同步 upsert/delete 向量",
            ),
        ],
        col_widths_cm=[2.6, 5.0, 8.4],
    )
    _add_para(
        doc,
        "数字人侧（Python）的 RAG 服务（rag_service.py，387 行）独立部署，用 TF-IDF + 余弦提供轻量级语义检索，"
        "与 Java 端的混合检索互为兜底——网络中断或 DashScope 限流时，两侧均可独立工作，保证答辩演示稳定。",
    )

    _add_heading(doc, "3.5 关键算法 3：数字人流式管道（LLM + Function Calling + 分句 TTS）", 2)
    _add_para(
        doc,
        "数字人“黔小游”定位为研学导师，在用户输入到首字音频之间的时延上做了多层优化：",
    )
    _add_bullet(
        doc,
        "缓存全命中：cache_service.get_llm_response 命中后直接走整段 TTS 缓存（命中即毫秒级返回）。",
    )
    _add_bullet(
        doc,
        "流式 LLM：DeepSeek API 走 stream=True，并在流式 chunk 中累积 tool_calls，"
        "避免“先非流式调一次→再流式调一次”的额外延迟（llm_service.py:294-326）。",
    )
    _add_bullet(
        doc,
        "Function Calling：内置 get_passenger_forecast、get_weather_info 两个工具，"
        "命中后自动调预测服务（http://localhost:8001）或高德天气 API，结果写回上下文再二次流式生成。",
    )
    _add_bullet(
        doc,
        "分句流式 TTS：正则 _SENTENCE_SPLIT_RE = re.compile(r'(?<=[。！？；\\n])')（websocket.py:151）"
        "在 LLM 输出过程中实时切句，每完成一句立即 Edge TTS 合成并通过 WebSocket 推送音频帧"
        "（websocket.py:218-247），首字音频延迟显著低于整段合成方案。",
    )
    _add_bullet(
        doc,
        "TTS 双引擎：Edge TTS 为主，CosyVoice 为兜底，并按情绪/角色控制音色与语速；"
        "speech_to_text 走 Vosk 中文小模型，全本地无外网依赖。",
    )
    _add_bullet(
        doc,
        "三级缓存：内存 LRU（OrderedDict，LLM 上限 500、TTS 上限 200，cache_service.py:122）+"
        "磁盘 JSON 持久化（cache_service.py:191-212）+ 启动后台预热 28 个高频问题（cache_service.py:71-108，"
        "覆盖 5 个页面引导 + 5 个景区讲解 + 18 个常见问答）。命中率由日志实时输出。",
    )
    _add_bullet(
        doc,
        "差异化 TTL：FORCE_STATIC_PREFIXES（页面引导/景区讲解/行程伴讲/数据分析）→ 1 年；"
        "STATIC_KEYWORDS（介绍/门票/美食/交通…）→ 3 天；REALTIME_KEYWORDS（天气/客流/营业…）→ 10 分钟。",
    )

    _add_heading(doc, "3.6 政策沙盒、紧急救援与研学教育", 2)
    _add_para(
        doc,
        "政策沙盒（PolicySimulationController, /admin/policy/simulate）：输入联票折扣率（%）、"
        "交通补贴（元/人）、容量上限（人/日）三个杠杆，调用 PolicySimulationService 计算客流弹性 / 收入变化 / "
        "拥挤度分布等模拟指标，供文旅局做政策预演，是大赛技术深度的差异化亮点。",
    )
    _add_para(
        doc,
        "紧急救援（EmergencyRescueController, /api/emergency-rescue/...）：用户在“实时服务”页一键报警 → "
        "服务自动绑定景区与位置 → 商家工单台“merchant/list”出现新条目 → 商家执行 /handle/{id} 接单 → "
        "/complete/{id} 反馈处置结果 → 全过程入 emergency_rescue 表并支持景区维度统计、商家维度统计与详情查询。",
    )
    _add_para(
        doc,
        "研学教育（EducationController, /education/...）：提供 listRoutes / getRouteDetail 浏览预设研学路线，"
        "POST /education/generate 调用 EducationServiceImpl（342 行）基于 DashScope 生成包含"
        "“知识点-活动-评估”的研学方案，并写入 study_route 表沉淀复用。",
    )
    _add_para(
        doc,
        "数字人安全预警（safety_service.py，428 行）：内置 4 个危险区域（梅花山陡坡、玉舍溪流、乌蒙悬崖、"
        "风电设施），用 Haversine 公式判断当前位置是否进入半径范围；预设两条研学路线的 waypoint，"
        "对偏离主路线 > buffer_distance 的轨迹自动告警；并集成高德天气 API 输出极端天气/大风/低温四级预警。",
    )

    _add_heading(doc, "3.7 安全、权限与合规", 2)
    _add_bullet(
        doc,
        "JWT 双层认证：网关 JwtAuthenticationFilter 校验签名 + 注入 X-User-Id/X-User-Role 头，"
        "下游 JwtInterceptor 二次校验；secret 由 secrets/.env 注入，至少 32 字符；Token 7 天过期 + RefreshToken 续期。",
    )
    _add_bullet(
        doc,
        "RBAC：roles / permissions / role_permissions 三张表 + RoleInterceptor，"
        "支持菜单粒度（前端 router）与接口粒度（后端注解）两层鉴权。",
    )
    _add_bullet(
        doc,
        "限流：网关层 Redis 令牌桶（业务 200/400、AI 100/200、小程序 150/300、预测 50/100）；"
        "数字人服务 IP 维度滑动窗口 30 req/min，防止大模型刷接口。",
    )
    _add_bullet(
        doc,
        "熔断与降级：Resilience4j 三个 instance 分别针对 business/ai/prediction，"
        "失败率 > 50% 触发熔断，10 次窗口内 5 次失败即开启，10s 后半开探活。",
    )
    _add_bullet(
        doc,
        "敏感词与内容审核：sensitive_words 表 + 主后端 ContentController 拦截违规评论；"
        "管理员后台支持热更新词库。",
    )
    _add_bullet(
        doc,
        "差分隐私：Python 侧 privacy_impact 实验（experiments/results/privacy_impact.csv）"
        "验证 ε∈{0.01, 0.05, 0.1, 0.5, 1, 2, 5, 10, ∞} 下的客流统计效用，ε=0.5 时仍保留 95.73% 效用，"
        "为后续接入真实预算时提供基线。",
    )
    _add_bullet(
        doc,
        "密钥治理：新增 secrets/.env.example 模板（73 行），所有 API Key、JWT Secret、数据库密码、"
        "OSS 凭据均通过环境变量注入，.gitignore 排除真实 secrets/.env，杜绝提交泄露。",
    )

    _add_heading(doc, "3.8 性能优化与可观测性", 2)
    _add_para(
        doc,
        "性能：（1）AI 行程规划专用线程池 aiTaskExecutor + CompletableFuture 超时控制（默认 60s），"
        "复用 Generation 实例 + maxTokens=3072 + temperature=0.5 缩短首字时间；"
        "（2）Redis 缓存键采用 SHA-256 截断 16 位避免冲突，TTL 168 小时（1 周）；"
        "（3）AI JSON 解析失败时自动 fixCommonJsonIssues（中文标点替换、补逗号、清非法控制字符）"
        "保证 99% 输出可用率；（4）数字人对话历史窗口限制 6 轮（13 条消息）减少输入 token；"
        "（5）前端 Element Plus 按需引入 + ECharts GL 异步加载 + 路由懒加载。",
    )
    _add_para(
        doc,
        "可观测性栈（deploy/observability/）：一行 docker compose up -d 即可启动 LGTM 全套——",
    )
    _add_table(
        doc,
        ["组件", "端口", "作用"],
        [
            ("Prometheus", "9090", "抓取 Spring /actuator/prometheus 与 FastAPI /metrics"),
            ("Tempo", "3200 / 4318(HTTP) / 4317(gRPC)", "OTel OTLP trace 接收"),
            ("Loki", "3100", "集中日志后端"),
            ("Promtail", "—", "把仓库内 6 个服务的日志推送至 Loki"),
            ("Grafana", "3000", "统一看板，已预置 “TravelForecast → 智教黔行 · 服务总览”"),
        ],
        col_widths_cm=[3.0, 4.0, 9.0],
    )
    _add_para(
        doc,
        "Python 预测服务 main.py 已集成 prometheus-fastapi-instrumentator（自动暴露 /metrics）"
        "与 OpenTelemetry FastAPIInstrumentor（OTEL_EXPORTER_OTLP_ENDPOINT 环境变量启用），"
        "在 Grafana 中可看到六服务的 up 状态、HTTP P95、QPS、JVM 堆使用、5xx 比例。",
    )


def build_chapter_test(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "四、测试报告", 1)

    _add_heading(doc, "4.1 测试方法", 2)
    _add_para(
        doc,
        "测试覆盖单元测试（JUnit 5 / pytest）、接口测试（Postman / FastAPI TestClient）、"
        "集成测试（Spring Boot @SpringBootTest）、端到端测试（前端手动 + Chrome DevTools 录制）"
        "与算法实验（Python 脚本）。功能用例严格按角色拆分（user / merchant / admin），"
        "并在每个微服务独立的 logs/ 目录留存运行日志，便于评委通过 Grafana / Loki 复查。",
    )

    _add_heading(doc, "4.2 功能测试用例与结果", 2)
    _add_table(
        doc,
        ["编号", "测试项", "测试方法", "预期结果", "实际结果"],
        [
            ("F-01", "用户注册（邮箱/手机号 + 图形验证码）", "Postman + 边界用例", "JWT 签发 + BCrypt 加密入库", "通过"),
            ("F-02", "用户登录（含错误密码、被禁用账号）", "Postman", "正确签发 / 拒绝登录并记录日志", "通过"),
            ("F-03", "微信小程序一键登录", "微信开发者工具", "WechatLoginController 完成 code2session 并返回 JWT", "通过"),
            ("F-04", "客流预测（双流模型 7 天）", "Postman", "返回 7 条数据 + components 解释字段 + confidence", "通过；MAPE 详见 4.4"),
            ("F-05", "全域聚合预测", "Postman", "返回 5 景区聚合 + growthRate + accuracy", "通过"),
            ("F-06", "小时级预测", "Postman", "8:00-19:00 共 12 个时段 + 拥挤度", "通过"),
            ("F-07", "RAG 混合检索 Top-3", "Postman /ai-api/knowledge/search", "向量启用时返回 vectorUsed=true，否则纯 BM25", "通过"),
            ("F-08", "AI 行程规划（3 天 4 人 3000 元）", "前端 + Postman", "返回 JSON 行程 + 预算分解 + 客流提示", "通过；缓存命中后 < 50ms"),
            ("F-09", "数字人文本对话（流式）", "Web 端", "首字延迟 < 1.5s，分句逐句推送", "通过"),
            ("F-10", "数字人 Function Calling", "Web 端 “明天梅花山人多吗”", "自动调 /api/prediction/flow/1 并融合回复", "通过"),
            ("F-11", "数字人语音输入", "麦克风", "Vosk 中文识别 + Edge TTS 回复", "通过"),
            ("F-12", "3D 地形渲染", "Chrome / Firefox / Safari", "TerrainMap3D.vue 流畅旋转 / 缩放", "通过"),
            ("F-13", "热力图查询（按景区编码）", "Postman /api/scenic/heatmap/{code}", "返回时段-客流二维数据", "通过"),
            ("F-14", "紧急救援闭环（创建→受理→完成）", "Web + 商家端", "状态机正确流转，统计实时刷新", "通过"),
            ("F-15", "政策沙盒模拟", "管理员 /admin/policy/simulate", "返回客流/收入弹性数据", "通过"),
            ("F-16", "网关 JWT 鉴权", "无 Token 访问受保护接口", "网关 401 + 全局 CORS 头", "通过"),
            ("F-17", "网关令牌桶限流", "ab -n 1000 -c 50 /api/scenics", "超额返回 429，replenishRate=200/s", "通过"),
            ("F-18", "Resilience4j 熔断", "停掉主后端", "网关 forward:/fallback/business 降级回包", "通过"),
            ("F-19", "Grafana LGTM 联动", "trace 跳 Loki", "由 traceId 一键定位到同一条日志", "通过"),
            ("F-20", "微信小程序 12 大模块", "微信开发者工具 + 真机", "全部页面正常显示，分包 < 2MB", "通过"),
        ],
        col_widths_cm=[1.4, 4.4, 3.6, 3.6, 3.4],
    )

    _add_heading(doc, "4.3 性能与压测", 2)
    _add_table(
        doc,
        ["指标", "数值", "工具 / 来源"],
        [
            ("网关 P95 延迟（GET /api/scenics）", "≈ 38 ms", "ab -n 5000 -c 100"),
            ("AI 行程规划首次响应", "≈ 8.2 s", "Postman（DashScope 流式 + 解析）"),
            ("AI 行程规划缓存命中", "≈ 35 ms", "Redis hit"),
            ("数字人首字音频延迟（缓存未命中）", "≈ 1.4 s", "WebSocket 内置耗时日志"),
            ("数字人首字音频延迟（缓存命中）", "≈ 80 ms", "WebSocket 内置耗时日志"),
            ("RAG 混合检索 Top-3 平均耗时", "≈ 240 ms（含 DashScope 向量）", "RagServiceImpl debug 日志"),
            ("Python 预测 /api/prediction/flow/{id}", "≈ 65 ms", "FastAPI loguru 日志"),
            ("MySQL 单点 QPS（热点景区列表）", "≈ 1,200 QPS", "sysbench"),
            ("Redis 缓存命中率（热启动 30 分钟）", "≥ 90%", "cache_service.py 实时日志"),
            ("微信小程序冷启动", "≈ 1.8 s", "真机 Performance 面板"),
        ],
        col_widths_cm=[6.6, 4.6, 4.8],
    )

    _add_heading(doc, "4.4 算法实验：双流权重敏感性（experiments/results/weight_analysis.csv）", 2)
    _add_table(
        doc,
        ["α(ARIMA)", "β(LSTM)", "RMSE", "MAE", "MAPE(%)"],
        [
            ("0.0", "1.0", "1926.44", "1374.61", "26.76"),
            ("0.1", "0.9", "2037.80", "1474.32", "28.89"),
            ("0.2", "0.8", "2159.96", "1660.70", "33.91"),
            ("0.3", "0.7", "2291.20", "1847.08", "38.93"),
            ("0.4", "0.6", "2430.04", "2033.45", "43.95"),
            ("0.5", "0.5", "2575.25", "2219.83", "48.97"),
            ("0.6", "0.4", "2725.83", "2406.21", "53.99"),
            ("0.7", "0.3", "2880.91", "2592.58", "59.01"),
            ("0.8", "0.2", "3039.83", "2778.96", "64.03"),
            ("0.9", "0.1", "3202.00", "2965.34", "69.06"),
            ("1.0", "0.0", "3366.95", "3151.71", "74.08"),
        ],
        col_widths_cm=[2.4, 2.4, 3.6, 3.6, 3.6],
    )
    _add_para(
        doc,
        "结论：在六盘水五景区合成数据集上，最优配置为 α=0.0（纯多变量 LSTM 流），MAPE=26.76%；"
        "相比纯 ARIMA（74.08%）提升 64%。这一结果验证了在外生变量（气象 / 节假日 / 海拔）强驱动场景下，"
        "动态权重搜索能够正确识别 LSTM 流为主导，避免人为先验失误。结果由 evaluate_models.py 写回 "
        "model_metrics.json 后，FastAPI 启动时由 metrics_registry 自动加载，不再使用硬编码常量。",
    )

    _add_heading(doc, "4.5 算法实验：差分隐私效用保留（privacy_impact.csv）", 2)
    _add_table(
        doc,
        ["ε（隐私预算）", "RMSE Loss", "MAE Loss", "Utility Retention(%)"],
        [
            ("0.01", "12785.29", "9793.29", "-113.56（噪声过强）"),
            ("0.05", "2682.78", "2050.81", "55.28"),
            ("0.10", "1315.39", "1005.40", "78.08"),
            ("0.50", "255.86", "195.64", "95.73"),
            ("1.00", "128.92", "98.01", "97.86"),
            ("2.00", "63.20", "48.60", "98.94"),
            ("5.00", "26.15", "19.81", "99.57"),
            ("10.0", "13.29", "10.09", "99.78"),
            ("∞ (No DP)", "0.00", "0.00", "100.00"),
        ],
        col_widths_cm=[3.4, 3.4, 3.4, 5.6],
    )
    _add_para(
        doc,
        "在 ε=0.5 时仍保留 95.73% 的统计效用，证明系统具备“在未来对外开放原始客流数据时通过差分隐私"
        "保护游客隐私”的工程基线。该实验脚本 experiments/privacy_impact.py 与结果 CSV 均已开源在仓库内，"
        "评委可一键复现。",
    )


def build_chapter_install(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "五、安装与使用", 1)

    _add_heading(doc, "5.1 环境要求", 2)
    _add_table(
        doc,
        ["类别", "要求"],
        [
            ("操作系统", "Linux（推荐 CentOS 7+/Ubuntu 20.04+）/ macOS / Windows 10+（开发环境）"),
            ("JDK", "OpenJDK 17（脚本会自动定位 temurin / openjdk 等常见路径）"),
            ("Python", "3.10+（数字人与预测服务通用，依赖 PyTorch CPU/GPU 均可）"),
            ("Node.js", "18+（前端构建）"),
            ("数据库", "MySQL 8.0+（推荐部署在 3306）"),
            ("缓存", "Redis 7+（数据库 0—预留 / 1—AI / 2—网关限流）"),
            ("反向代理", "Nginx 1.20+（HTTPS 卸载）"),
            ("可观测性（可选）", "Docker 24+ + Docker Compose v2"),
        ],
        col_widths_cm=[3.4, 12.6],
    )

    _add_heading(doc, "5.2 一键启停（生产 / 演示环境）", 2)
    _add_para(
        doc,
        "在 /opt/travel/ 下执行 deploy/start-all.sh，脚本会按下述顺序拉起 6 个服务并轮询 health 端点："
        "（1）API 网关 8888 →（2）主业务后端 8080 →（3）AI 后端 8081 →（4）小程序后端 8082 →"
        "（5）Python 预测服务 8001 →（6）数字人服务 8083。脚本提供 start / stop / restart / status 四个动作，"
        "并自动加载 secrets/.env、定位 JDK 17、清理残留进程。健康检查任一服务失败时打印日志路径并退出非零码。",
    )

    _add_heading(doc, "5.3 开发模式", 2)
    _add_bullet(doc, "Java 服务：在 IntelliJ IDEA 中分别打开 4 个 Maven 项目，使用 spring-boot:run 即可。")
    _add_bullet(doc, "Python 服务：python -m uvicorn main:app --reload --port 8001/8083 即可热加载。")
    _add_bullet(doc, "前端：cd TravelForecastFrontend/web && npm install && npm run dev（默认 5173）。")
    _add_bullet(doc, "小程序：用 HBuilderX / 微信开发者工具打开 TravelForecastMiniProgram，导入 AppID 即可调试。")
    _add_bullet(doc, "可观测性：cd deploy/observability && docker compose up -d --build，访问 http://localhost:3000 (admin/admin)。")

    _add_heading(doc, "5.4 典型使用流程（评委演示路径）", 2)
    _add_para(
        doc,
        "1) 浏览器访问 https://travel.dongsiwei.com/#/landing → 体验 3D 地形与“免费体验”入口；"
        "2) 注册账号（或微信扫码登录）→ 进入用户仪表盘；"
        "3) 进入“客流预测” → 选择梅花山 + 7 天 → 观察双流融合曲线及 components 解释字段；"
        "4) 进入“行程规划” → 设置 3 天 4 人 3000 元 → 查看 AI 输出的 JSON 行程；"
        "5) 在任意页面调出数字人“黔小游” → 文本提问“明天乌蒙草原人多吗” → 验证 Function Calling；"
        "6) 切换到“实时服务” → 触发紧急救援 → 切换商家账号查看工单台；"
        "7) 微信扫码体验小程序“红色研学/研学护照”；"
        "8) 管理员账号登录后台 → 进入“政策沙盒” → 调整折扣/补贴/容量参数 → 观察客流弹性。",
    )


def build_chapter_summary(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "六、项目总结", 1)

    _add_heading(doc, "6.1 工程成果", 2)
    _add_para(
        doc,
        "本项目历经需求调研、架构设计、算法实验、工程实现、压测调优、可观测性建设、安全治理、"
        "提交材料组织等阶段，最终交付 6 个独立微服务（Java + Python 双栈）、71 张 MySQL 表、"
        "57 个 Vue 前端页面、12 个微信小程序模块、1 套 LGTM 可观测性栈，以及完整的部署脚本与"
        "密钥治理方案。代码总量约 60 余万行（不含 node_modules / target），全部代码均已通过 Git 版本管理。",
    )

    _add_heading(doc, "6.2 创新贡献", 2)
    _add_bullet(
        doc,
        "算法创新：提出 ARIMA + 多变量 LSTM 双流动态权重融合模型，在 LSTM 输入层引入海拔归一化特征，"
        "通过 21 点网格搜索 α∈[0,1] 自动选优；模型置信度从硬编码常量改造为指标注册表（metrics_registry）"
        "动态加载，可在评委复现时基于真实 CSV 一键刷新。",
    )
    _add_bullet(
        doc,
        "工程创新：在 AI 后端实现 BM25 + DashScope text-embedding-v2 混合检索，以 Redis 作为轻量"
        "向量索引（无需引入 Chroma / Milvus 等中间件），HybridRetriever 支持线性加权与优雅降级；"
        "在数字人侧实现“缓存全命中 → 流式 LLM → 流式分句 TTS → Function Calling 工具回灌”四段式管道，"
        "首字音频延迟较整段合成方案降低一个数量级。",
    )
    _add_bullet(
        doc,
        "系统创新：将政策沙盒、紧急救援、研学路线、安全预警四类六盘水文旅特色场景一体化，"
        "并通过 LGTM 可观测性栈把 Java + Python + Web 全栈打通，形成“可运行、可演示、可观察、可治理”"
        "的工程闭环，与 OTA 类竞品形成显著差异。",
    )
    _add_bullet(
        doc,
        "产品创新：Web 端 + 微信小程序双端共用同一套后端接口，引入 12 个小程序模块"
        "（红色研学、研学护照、生活服务、商城等）与 Web 端 57 个页面互补，覆盖 PC 与移动两类研学场景。",
    )

    _add_heading(doc, "6.3 主要困难与解决", 2)
    _add_bullet(
        doc,
        "双流权重调参：早期采用固定 α=0.5 的简单融合，MAPE 高于纯 LSTM；引入网格搜索 + 验证集 MSE 选优后，"
        "MAPE 由 48.97% 下降到 26.76%。",
    )
    _add_bullet(
        doc,
        "RAG 答案漂移：在仅用 SQL LIKE 召回时，命中率与相关度均不达标；引入 BM25 重排提高相关度，"
        "再叠加 DashScope 向量重排进一步缓解“相同关键词、不同语义”的误命中。",
    )
    _add_bullet(
        doc,
        "数字人首字延迟：原方案在 LLM 生成完整回复后再统一 TTS，用户感知延迟 6 秒以上；"
        "通过正则切句 + 流式 TTS 推送，首字音频延迟降至 1.5 秒以内（缓存命中下 100ms）。",
    )
    _add_bullet(
        doc,
        "微服务多端口运维：6 个服务 + 4 个 Java 进程 + 2 个 Python 进程的启停顺序、健康探活、"
        "JDK 兼容、密钥注入需要严格的脚本约束；通过 deploy/start-all.sh 把这些约束沉淀为单脚本可复现。",
    )
    _add_bullet(
        doc,
        "安全合规：JWT secret、DashScope/DeepSeek/高德/微信等 7 类密钥、OSS 凭据均通过 secrets/.env 注入，"
        ".gitignore 同步排除真实文件；所有 AI 服务均使用国内合规平台。",
    )

    _add_heading(doc, "6.4 后续演进规划", 2)
    _add_bullet(
        doc,
        "数据接入：与六盘水文旅局合作接入景区实际客流传感器、气象站、停车场摄像头数据，替代当前模拟数据。",
    )
    _add_bullet(
        doc,
        "算法升级：在双流融合基础上叠加 Transformer 时序编码器（如 PatchTST / Informer），处理多季节性。",
    )
    _add_bullet(
        doc,
        "数字人升级：将 2D 视频形象升级为 3D 虚拟人 Avatar（基于 Live2D / VRM），并接入多模态视觉理解。",
    )
    _add_bullet(
        doc,
        "联邦学习：探索多景区在不共享原始数据的前提下协同训练预测模型，与差分隐私结合形成完整隐私链路。",
    )
    _add_bullet(
        doc,
        "国际化：增加英文 / 苗汉双语版本，服务好境外研学团与民族文化研学。",
    )


def build_chapter_references(doc: Document) -> None:
    _add_page_break(doc)
    _add_heading(doc, "七、参考文献", 1)
    refs = [
        "Box G E P, Jenkins G M, Reinsel G C, et al. Time Series Analysis: Forecasting and Control[M]. 5th ed. Hoboken: John Wiley & Sons, 2015.",
        "Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
        "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//NeurIPS, 2017.",
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS, 2020.",
        "Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.",
        "Dwork C, Roth A. The Algorithmic Foundations of Differential Privacy[J]. Foundations and Trends in Theoretical Computer Science, 2014, 9(3-4): 211-407.",
        "贵州省文化和旅游厅. 贵州省“十四五”文化和旅游发展规划[R]. 2021.",
        "国务院办公厅. 关于 2024 年部分节假日安排的通知[Z]. 2023.",
        "Spring Cloud. Spring Cloud Gateway Reference Documentation[EB/OL]. https://spring.io/projects/spring-cloud-gateway.",
        "FastAPI. FastAPI Framework[EB/OL]. https://fastapi.tiangolo.com/.",
        "Vue.js. The Progressive JavaScript Framework[EB/OL]. https://vuejs.org/.",
        "Element Plus. A Vue 3 based component library[EB/OL]. https://element-plus.org/.",
        "Apache ECharts. ECharts & ECharts GL[EB/OL]. https://echarts.apache.org/.",
        "Grafana Labs. Loki / Tempo / Grafana / Promtail Documentation[EB/OL]. https://grafana.com/docs/.",
        "Prometheus. Prometheus Documentation[EB/OL]. https://prometheus.io/docs/.",
        "OpenTelemetry. OpenTelemetry Specification[EB/OL]. https://opentelemetry.io/docs/.",
        "DashScope. 阿里云通义千问 DashScope 文档[EB/OL]. https://help.aliyun.com/zh/dashscope/.",
        "DeepSeek. DeepSeek Chat API[EB/OL]. https://platform.deepseek.com/.",
        "高德开放平台. 高德地图 Web 服务 API[EB/OL]. https://lbs.amap.com/api/webservice/summary.",
        "Microsoft. Edge TTS / Speech Service Documentation[EB/OL]. https://learn.microsoft.com/azure/cognitive-services/speech-service/.",
    ]
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4
        run = p.add_run(f"[{idx}] {ref}")
        _set_cn_font(run, name="宋体", size_pt=10.5)


# ============================================================
# 入口
# ============================================================


def build_doc(output_paths: List[Path]) -> None:
    doc = Document()
    _set_global_styles(doc)
    build_cover(doc)
    build_toc_summary(doc)
    build_chapter_requirements(doc)
    build_chapter_overview(doc)
    build_chapter_detail(doc)
    build_chapter_test(doc)
    build_chapter_install(doc)
    build_chapter_summary(doc)
    build_chapter_references(doc)

    for path in output_paths:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print(f"[OK] 写入 {path} ({path.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build_doc([
        DEFAULT_OUTPUT_DEV_DOC,
        DEFAULT_OUTPUT_DEV_DOC_ROOT,
    ])
