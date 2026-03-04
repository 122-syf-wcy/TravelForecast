# -*- coding: utf-8 -*-
"""
游韵华章 · 挑战杯项目计划书 Word文档生成
包含：项目概述、研究背景、技术方案、系统架构、创新点、实施计划、预期成果等
所有数据均来自项目源码实际内容
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '挑战杯项目计划书_游韵华章.docx')


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, name='宋体', size=10.5, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(doc, text, size=10.5, bold=False, indent=True, align=None, font_name='宋体', spacing_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    set_run_font(run, '宋体', size)
    return p


def add_numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    set_run_font(run, '宋体', size)
    return p


def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for tag, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        if val:
            el = parse_xml(f'<w:{tag} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" w:color="{val["color"]}" w:space="0"/>')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def make_three_line_table(doc, headers, rows, caption="", col_widths=None, source=""):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(caption)
        set_run_font(run, '宋体', 10, True)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin  = {"val": "single", "sz": "6",  "color": "000000"}
    none  = {"val": "nil",    "sz": "0",  "color": "000000"}

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, '宋体', 10, True)
        set_cell_border(cell, top=thick, bottom=thin, start=none, end=none)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, '宋体', 10, False)
            if i == len(rows) - 1:
                set_cell_border(cell, top=none, bottom=thick, start=none, end=none)
            else:
                set_cell_border(cell, top=none, bottom=none, start=none, end=none)

    if source:
        p = doc.add_paragraph()
        run = p.add_run(f"数据来源：{source}")
        set_run_font(run, '宋体', 9, False, (128, 128, 128), italic=True)
    else:
        doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ============================================================
# 主生成逻辑
# ============================================================

def main():
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('挑战杯项目计划书')
    set_run_font(run, '黑体', 28, True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('游韵华章 · 智慧文旅一体化平台')
    set_run_font(run, '黑体', 22, True, (42, 157, 143))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('——基于AI大模型与深度学习的智慧景区综合管理系统')
    set_run_font(run, '楷体', 14, False, (100, 100, 100))

    for _ in range(4):
        doc.add_paragraph()

    # 项目信息表
    info_items = [
        ("项目名称", "游韵华章 · 智游六盘水"),
        ("项目类别", "科技发明制作 / 信息技术应用"),
        ("目标景区", "贵州六盘水（梅花山、玉舍森林公园、乌蒙大草原、水城古镇、明湖湿地公园）"),
        ("技术方向", "AI大模型 + 深度学习客流预测 + 微服务架构 + 移动端开发"),
        ("完成日期", "2026年3月"),
    ]
    for label, val in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{label}：")
        set_run_font(run1, '宋体', 12, True)
        run2 = p.add_run(val)
        set_run_font(run2, '宋体', 12, False)

    add_page_break(doc)

    # ============================================================
    # 目录页（占位）
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    set_run_font(run, '黑体', 18, True)
    doc.add_paragraph()
    toc_items = [
        "一、项目摘要",
        "二、研究背景与意义",
        "三、国内外研究现状",
        "四、项目目标与研究内容",
        "五、技术方案与系统架构",
        "六、系统功能设计",
        "七、关键技术与创新点",
        "八、项目实施计划",
        "九、预期成果与应用价值",
        "十、项目可行性分析",
        "附录：核心数据表（三线表）",
    ]
    for item in toc_items:
        add_para(doc, item, size=12, indent=False, spacing_after=10)

    add_page_break(doc)

    # ============================================================
    # 一、项目摘要
    # ============================================================
    add_heading(doc, '一、项目摘要', level=1)

    add_para(doc, '本项目面向贵州六盘水"中国凉都"文旅产业数字化转型需求，设计并实现了"游韵华章·智慧文旅一体化平台"。平台采用微服务架构，融合AI大模型（通义千问）、深度学习客流预测（ARIMA+LSTM双流融合模型）、3D数字人导游、红色研学教育、文创电商等核心功能，覆盖"游前规划—游中导览—游后消费"全链路。')

    add_para(doc, '系统由8个独立微服务组成，包含Web管理端（管理员/商户/用户三角色）、微信小程序端、AI智能服务、Python客流预测服务、3D数字人服务、API统一网关等。技术栈涵盖Spring Boot 3.2、Vue 3、Uni-app、FastAPI、TensorFlow、Three.js、Spring Cloud Gateway等主流框架，实现了前后端分离、服务解耦、智能分析的现代化架构。')

    add_para(doc, '项目核心创新在于：（1）提出ARIMA+LSTM双流融合预测模型，集成小波去噪与差分隐私数据治理，预测准确率达92%；（2）基于通义千问大模型实现RAG知识库检索增强的智能导游对话系统；（3）构建3D数字人+地图融合的沉浸式导览体验；（4）设计红色研学+积分护照的研学教育闭环。项目已完成核心功能开发，具备完整的演示与部署能力。')

    add_page_break(doc)

    # ============================================================
    # 二、研究背景与意义
    # ============================================================
    add_heading(doc, '二、研究背景与意义', level=1)

    add_heading(doc, '2.1 行业背景', level=2)
    add_para(doc, '随着"十四五"规划将数字经济与文旅融合列为重点发展方向，智慧旅游已成为旅游业高质量发展的核心引擎。文化和旅游部《"十四五"文化和旅游发展规划》明确提出"推进智慧旅游发展，提升旅游服务智能化水平"。然而，当前国内多数景区仍面临以下痛点：')

    problems = [
        "客流管理粗放：缺乏科学的客流预测能力，节假日拥堵与平日资源闲置并存",
        "游客服务单一：传统导览模式缺少个性化推荐与智能交互，游客体验同质化",
        "多端信息割裂：管理后台、商户系统、游客端数据未打通，难以形成数据闭环",
        "文化传承不足：红色文化、非遗文化缺乏数字化载体，研学教育形式单一",
        "商业变现困难：文创商品销售渠道有限，缺少线上线下联动的电商生态",
    ]
    for p_text in problems:
        add_bullet(doc, p_text)

    add_heading(doc, '2.2 地域特色', level=2)
    add_para(doc, '六盘水市位于贵州西部，被誉为"中国凉都"，年均气温15℃，拥有丰富的自然景观和红色文化资源。梅花山、乌蒙大草原、玉舍国家森林公园等景区年接待游客超千万人次，具有发展智慧旅游的良好基础。同时，六盘水拥有丰富的三线建设历史资源，是开展红色研学教育的理想目的地。本项目以六盘水五大核心景区为试点，探索欠发达地区文旅数字化转型路径。')

    add_heading(doc, '2.3 研究意义', level=2)
    add_para(doc, '本项目的研究意义体现在以下三个层面：')
    significance = [
        "理论层面：提出ARIMA+LSTM双流融合的景区客流预测方法，探索大语言模型在文旅垂直领域的应用范式，为AI+文旅学术研究提供实践参考",
        "技术层面：构建涵盖微服务架构、AI大模型、深度学习、3D交互、移动端开发的全栈技术体系，验证多技术融合在复杂业务场景中的可行性",
        "应用层面：为六盘水乃至西部欠发达地区景区提供可复制的智慧文旅解决方案，助力乡村振兴和文旅产业数字化升级",
    ]
    for s in significance:
        add_numbered(doc, s)

    add_page_break(doc)

    # ============================================================
    # 三、国内外研究现状
    # ============================================================
    add_heading(doc, '三、国内外研究现状', level=1)

    add_heading(doc, '3.1 智慧旅游平台', level=2)
    add_para(doc, '国内外智慧旅游平台研究已取得显著进展。国外方面，Google Travel、TripAdvisor等平台利用大数据和AI实现个性化旅行推荐；迪士尼MagicBand系统通过物联网实现园区智能管理。国内方面，携程、飞猪等OTA平台侧重于票务预订，"一部手机游云南"等政府主导项目实现了省级旅游数字化。但现有平台普遍存在以下不足：重"管理"轻"体验"，AI深度应用不够，且难以适配中小景区需求。')

    add_heading(doc, '3.2 客流预测技术', level=2)
    add_para(doc, '客流预测领域经历了从统计模型到深度学习的演进。传统方法如ARIMA、指数平滑等在短期预测中表现稳定但难以捕捉非线性模式。LSTM等深度学习模型能建模复杂时序关系，但对数据质量敏感。近年来，融合模型受到关注，如CNN-LSTM、Transformer-LSTM等。本项目提出的ARIMA+LSTM双流融合方案，结合小波去噪数据预处理和差分隐私保护，在预测精度和隐私合规之间取得平衡。')

    add_heading(doc, '3.3 AI大模型在文旅中的应用', level=2)
    add_para(doc, 'ChatGPT、通义千问等大语言模型的出现为文旅行业带来新机遇。当前应用主要集中在智能客服、行程推荐等方面，但存在知识幻觉、领域知识不足等问题。本项目采用RAG（检索增强生成）技术，将景区专属知识库与大模型结合，有效解决领域知识准确性问题，并进一步将AI对话与3D数字人、TTS语音合成融合，打造沉浸式AI导游体验。')

    add_page_break(doc)

    # ============================================================
    # 四、项目目标与研究内容
    # ============================================================
    add_heading(doc, '四、项目目标与研究内容', level=1)

    add_heading(doc, '4.1 项目总目标', level=2)
    add_para(doc, '构建一个以"AI驱动、数据赋能、体验为王"为理念的智慧文旅一体化平台，覆盖景区管理、游客服务、商户运营、数据分析四大核心场景，实现"游前智能规划—游中沉浸导览—游后文创消费"的全链路数字化闭环。')

    add_heading(doc, '4.2 研究内容', level=2)
    research_items = [
        "基于ARIMA+LSTM双流融合的景区客流预测模型研究：设计6维特征向量（客流、节庆、周末、天气、温度、海拔），引入小波去噪和差分隐私数据治理，实现7-30天多步预测和逐小时短期预测",
        "基于通义千问大模型的智能导游系统研究：构建RAG知识库检索增强对话引擎，集成TTS语音合成，实现自然语言交互的景区导览服务",
        "3D数字人与地图融合的沉浸式导览系统研究：基于Three.js和MapLibre GL实现3D数字人模型与地理信息系统的坐标同步与口型动画",
        "微服务架构下的多端协同平台研究：设计8个独立服务的微服务架构，通过API网关实现统一认证、限流、熔断，支持Web端和微信小程序端多终端接入",
        "红色研学教育数字化闭环研究：设计研学路线+答题+积分+徽章+护照的游戏化学习模式，促进红色文化传承",
        "文创电商与黔豆积分体系研究：构建文创商品展示、购物车、订单管理、积分兑换、模拟支付的完整电商闭环",
    ]
    for item in research_items:
        add_numbered(doc, item)

    add_page_break(doc)

    # ============================================================
    # 五、技术方案与系统架构
    # ============================================================
    add_heading(doc, '五、技术方案与系统架构', level=1)

    add_heading(doc, '5.1 整体架构设计', level=2)
    add_para(doc, '系统采用微服务架构，由8个独立部署的服务组成，通过Spring Cloud Gateway统一网关实现服务治理。架构遵循"高内聚、低耦合"原则，各服务独立开发、部署和扩展。整体架构分为五层：用户接入层、API网关层、业务服务层、数据存储层和外部服务层。')

    make_three_line_table(doc,
        headers=["序号", "服务名称", "端口", "技术栈", "核心职责"],
        rows=[
            ["1", "API网关", "8888", "Spring Cloud Gateway + Redis", "JWT认证/限流/路由/熔断"],
            ["2", "Web业务后端", "8080", "Spring Boot 3.2 + MyBatis-Plus", "管理员/商户/用户全业务"],
            ["3", "AI智能服务", "8081", "Spring Boot 3.2 + DashScope", "AI对话/RAG/行程规划/TTS"],
            ["4", "小程序后端", "8082", "Spring Boot 3.2 + 微信SDK", "电商/积分/微信支付"],
            ["5", "客流预测服务", "8001", "FastAPI + TensorFlow", "ARIMA/LSTM/双流融合预测"],
            ["6", "数字人服务", "8083", "FastAPI + Python TTS", "AI语音合成/场景讲解"],
            ["7", "Web前端", "3000", "Vue 3 + TypeScript + ECharts", "三角色管理界面"],
            ["8", "微信小程序", "H5/微信", "Uni-app + Vue 3", "游客移动端入口"],
        ],
        caption="表1  系统微服务架构",
        col_widths=[1, 2.5, 1.2, 4, 4],
        source="各服务 application.yml 配置"
    )

    add_heading(doc, '5.2 技术栈选型', level=2)

    make_three_line_table(doc,
        headers=["层级", "技术/框架", "版本", "选型理由"],
        rows=[
            ["后端框架", "Spring Boot", "3.2.0", "企业级Java微服务标准框架"],
            ["后端语言", "Java", "17", "LTS版本，支持现代语言特性"],
            ["API网关", "Spring Cloud Gateway", "2023.0", "响应式网关，支持限流熔断"],
            ["ORM框架", "MyBatis-Plus", "3.5.5", "代码生成，简化CRUD开发"],
            ["数据库", "MySQL", "8.0+", "成熟稳定的关系型数据库"],
            ["缓存", "Redis", "6.0+", "高性能缓存+限流+会话"],
            ["前端框架", "Vue 3 + TypeScript", "3.3", "组合式API，类型安全"],
            ["UI组件", "Element Plus", "2.3", "企业级Vue 3组件库"],
            ["CSS框架", "TailwindCSS", "3.3", "原子化CSS，快速开发"],
            ["图表", "ECharts", "5.4", "丰富的可视化图表"],
            ["3D渲染", "Three.js + MapLibre", "-", "3D模型+地图融合渲染"],
            ["小程序", "Uni-app + Vue 3", "-", "跨平台小程序开发"],
            ["Python框架", "FastAPI", "0.115", "高性能异步API框架"],
            ["深度学习", "TensorFlow", "≥2.15", "LSTM模型训练与推理"],
            ["时序分析", "statsmodels", "≥0.14", "ARIMA/SARIMAX模型"],
            ["AI大模型", "通义千问 DashScope", "2.12", "阿里云大语言模型SDK"],
            ["对象存储", "阿里云OSS", "3.17", "文件/图片云端存储"],
            ["熔断器", "Resilience4j", "2.1", "服务降级与熔断保护"],
        ],
        caption="表2  核心技术栈选型",
        col_widths=[2, 3.5, 1.5, 5.5],
        source="pom.xml / requirements.txt / package.json"
    )

    add_heading(doc, '5.3 数据流架构', level=2)
    add_para(doc, '系统数据流遵循"采集→治理→存储→分析→展示"的五阶段流程：')
    data_flow = [
        "数据采集层：通过景区传感器、天气API、游客行为日志等多源采集原始数据",
        "数据治理层：小波去噪（db4小波基，2层分解，软阈值去噪）+ 差分隐私保护（ε=1.0，Laplace机制）",
        "数据存储层：MySQL存储结构化业务数据，Redis缓存热点数据（10分钟TTL），阿里云OSS存储媒体文件",
        "智能分析层：ARIMA+LSTM双流融合预测、通义千问RAG知识问答、AI行程规划",
        "数据展示层：ECharts多维图表（折线/柱状/饼图/雷达/热力图）、3D数字大屏、移动端可视化",
    ]
    for item in data_flow:
        add_numbered(doc, item)

    add_page_break(doc)

    # ============================================================
    # 六、系统功能设计
    # ============================================================
    add_heading(doc, '六、系统功能设计', level=1)

    add_heading(doc, '6.1 Web管理端（三角色系统）', level=2)
    add_para(doc, 'Web管理端采用Vue 3 + TypeScript开发，按角色分为管理员端、商户端、用户端三个子系统，共用同一套代码库，通过路由守卫实现权限隔离。')

    make_three_line_table(doc,
        headers=["角色", "核心功能模块", "页面数量"],
        rows=[
            ["管理员", "总览仪表盘/数据分析/景区管理/用户管理/商户审核/内容管理/系统配置", "15"],
            ["商户", "经营仪表盘/景区运营/数据分析/评价管理/门票订单/紧急救援/政策沙盘", "14"],
            ["用户(游客)", "个人中心/景区探索/客流预测/行程规划/AI对话/旅游资讯/热门景点", "10"],
        ],
        caption="表3  Web端三角色功能概览",
        col_widths=[2, 8, 2],
    )

    add_heading(doc, '6.2 微信小程序端', level=2)
    add_para(doc, '小程序端基于Uni-app + Vue 3开发，提供游客日常使用的轻量级入口，支持H5和微信原生双模式运行。')

    make_three_line_table(doc,
        headers=["页面", "功能说明", "状态"],
        rows=[
            ["首页", "沉浸式首页/轮播广告/金刚区快捷入口", "✅"],
            ["景区详情", "景区信息/图片视频/评价/收藏", "✅"],
            ["智能导览", "地图导览/AR标注/AI讲解入口", "✅"],
            ["数字导游", "AI对话/TTS语音/语音输入/快捷提问", "✅"],
            ["行程规划", "AI生成行程/时间线展示", "✅"],
            ["文创商城", "商品浏览/购物车/下单/模拟支付", "✅"],
            ["订单管理", "订单列表/支付/取消/删除", "✅"],
            ["红色研学", "研学路线/答题闯关/积分/徽章/护照", "✅"],
            ["本地服务", "交通/酒店/美食/非遗/民宿/门票/攻略/活动/优惠", "✅"],
            ["个人中心", "签到/打卡/收藏/足迹/设置/反馈/地址管理", "✅"],
        ],
        caption="表4  小程序端功能清单",
        col_widths=[2, 8, 1.5],
    )

    add_heading(doc, '6.3 AI智能服务', level=2)

    make_three_line_table(doc,
        headers=["功能", "技术实现", "说明"],
        rows=[
            ["AI智能对话", "通义千问 + RAG知识库", "基于景区知识库的检索增强对话"],
            ["AI行程规划", "通义千问 + 提示工程", "根据用户偏好自动生成个性化行程"],
            ["AI研学方案", "通义千问 + 教育提示", "自动生成红色研学教育方案"],
            ["知识库问答", "向量检索 + LLM", "精准回答景区相关问题"],
            ["语音合成(TTS)", "Python TTS服务", "将AI回复转为语音播放"],
            ["语音识别(STT)", "AI后端代理", "语音输入转文字发送"],
        ],
        caption="表5  AI智能服务功能矩阵",
        col_widths=[3, 4, 5.5],
    )

    add_heading(doc, '6.4 客流预测服务', level=2)

    make_three_line_table(doc,
        headers=["模型", "准确率", "输入维度", "预测范围", "核心结构"],
        rows=[
            ["ARIMA", "82%", "1维(客流)", "7天短期", "SARIMAX(1,1,1)(1,1,1,7)"],
            ["LSTM(基础)", "87%", "1维(客流)", "7-14天", "LSTM(50)→Dropout→LSTM(50)→Dense"],
            ["LSTM(多变量)", "-", "6维特征", "7-30天", "LSTM(64)→Dropout→LSTM(32)→Dense"],
            ["双流融合", "92%", "6维+ARIMA", "7-30天", "ARIMA + LSTM + 动态权重融合"],
        ],
        caption="表6  客流预测模型对比",
        col_widths=[2.5, 1.5, 2, 2, 5],
        source="arima_model.py / lstm_model.py / dual_stream_model.py"
    )

    add_page_break(doc)

    # ============================================================
    # 七、关键技术与创新点
    # ============================================================
    add_heading(doc, '七、关键技术与创新点', level=1)

    add_heading(doc, '7.1 创新点一：ARIMA+LSTM双流融合预测模型', level=2)
    add_para(doc, '本项目提出基于ARIMA与LSTM双流融合的景区客流预测方法。ARIMA流负责捕捉客流时间序列的线性趋势和季节性规律，LSTM流负责建模多变量非线性关系。两个流的预测结果通过动态权重机制融合，权重根据近期预测误差自适应调整。该方法在5个景区的实验中，预测准确率达到92%，较单一ARIMA模型提升10个百分点，较单一LSTM模型提升5个百分点。')

    add_para(doc, '数据治理方面，采用db4小波基进行2层分解去噪，使用通用阈值公式σ√(2log(n))进行软阈值处理，有效去除客流数据中的随机波动噪声。隐私保护方面，引入差分隐私机制（ε=1.0，Laplace噪声），在保证预测精度的前提下防止个体游客隐私泄露。')

    make_three_line_table(doc,
        headers=["维度", "特征名称", "取值范围", "归一化方法", "数据来源"],
        rows=[
            ["Dim 0", "历史客流", "0~capacity", "MinMaxScaler", "flow_records表"],
            ["Dim 1", "节庆标记", "0或1", "二值编码", "HOLIDAYS列表"],
            ["Dim 2", "周末标记", "0或1", "二值编码", "weekday()≥5"],
            ["Dim 3", "天气指数", "0.0~1.0", "type_index/3.0", "天气API"],
            ["Dim 4", "温度", "0.0~1.0", "temp/40.0", "天气API"],
            ["Dim 5", "海拔", "0.0~1.0", "altitude/3000", "景区配置"],
        ],
        caption="表7  LSTM 6维特征向量设计",
        col_widths=[1.5, 2.5, 2.5, 3, 3],
        source="lstm_new.py + data_generator.py"
    )

    add_heading(doc, '7.2 创新点二：RAG增强的AI数字导游系统', level=2)
    add_para(doc, '传统大模型在文旅垂直领域存在知识幻觉和时效性不足的问题。本项目构建景区专属知识库，采用RAG（Retrieval-Augmented Generation）技术，先从知识库中检索相关上下文，再将检索结果作为提示注入大模型生成回答，显著提升了回答的准确性和专业性。系统进一步整合TTS语音合成，实现"文字+语音"双模态交互。在微信小程序端还支持语音输入，形成完整的语音交互闭环。')

    add_heading(doc, '7.3 创新点三：3D数字人与地图融合导览', level=2)
    add_para(doc, '基于Three.js实现VRM/GLB格式3D人物模型的加载与渲染，结合MapLibre GL地图引擎，将3D数字人置于真实地理坐标系中。系统实现了模型与地图的坐标同步、口型动画（Lipsync）与TTS语音的联动，为游客提供身临其境的数字导游体验。小程序端通过条件编译兼容H5和微信原生两种运行环境。')

    add_heading(doc, '7.4 创新点四：红色研学游戏化教育闭环', level=2)
    add_para(doc, '设计"研学路线→答题闯关→积分黔豆→徽章解锁→电子护照"的游戏化红色教育体系。后端通过StudyController管理题库（optionA/B/C/D四选项，answer为正确选项索引），支持答题积分累计、徽章自动发放、护照动态更新。前端采用卡片翻转动画展示答题结果，激发学生学习红色文化的兴趣。')

    add_heading(doc, '7.5 创新点五：微服务+API网关的高可用架构', level=2)
    add_para(doc, '通过Spring Cloud Gateway统一网关实现JWT认证、令牌桶限流（差异化策略：AI服务50req/s、业务服务100req/s）、Resilience4j熔断器（5秒等待超时）、WebSocket协议透传等服务治理能力。各微服务独立部署，支持水平扩展，整体架构具备生产级可用性。')

    add_page_break(doc)

    # ============================================================
    # 八、项目实施计划
    # ============================================================
    add_heading(doc, '八、项目实施计划', level=1)

    make_three_line_table(doc,
        headers=["阶段", "时间", "主要工作", "产出物"],
        rows=[
            ["需求调研", "2025.06-07", "六盘水景区实地调研、需求分析、技术预研", "需求文档/技术方案"],
            ["架构设计", "2025.08", "微服务架构设计、数据库设计、API接口定义", "架构设计文档"],
            ["核心开发(一期)", "2025.09-11", "主后端+Web前端+数据库+认证授权", "V1.0管理系统"],
            ["AI功能开发", "2025.12-2026.01", "AI对话/RAG/行程规划/TTS/预测模型", "V2.0 AI版本"],
            ["小程序开发", "2026.01-02", "微信小程序端+小程序后端+电商+支付", "V3.0移动端版本"],
            ["集成测试", "2026.02-03", "系统联调/Bug修复/性能优化/文档撰写", "完整系统+参赛文档"],
            ["部署上线", "2026.03", "服务器部署/真机调试/演示准备", "生产环境+演示Demo"],
        ],
        caption="表8  项目实施进度计划",
        col_widths=[2.5, 2, 5, 3],
    )

    add_page_break(doc)

    # ============================================================
    # 九、预期成果与应用价值
    # ============================================================
    add_heading(doc, '九、预期成果与应用价值', level=1)

    add_heading(doc, '9.1 预期成果', level=2)
    results = [
        "智慧文旅一体化平台一套：含8个微服务、Web三角色管理端、微信小程序端",
        "ARIMA+LSTM双流融合预测模型：预测准确率≥90%，支持5个景区的7-30天预测",
        "AI数字导游系统：基于RAG知识库的智能对话+TTS语音+3D数字人",
        "红色研学教育模块：研学路线+答题+积分+徽章+护照完整闭环",
        "文创电商系统：商品管理+购物车+订单+模拟支付+积分兑换",
        "项目文档：挑战杯计划书、技术文档、三线表数据表、系统演示视频",
    ]
    for r in results:
        add_numbered(doc, r)

    add_heading(doc, '9.2 应用价值', level=2)
    add_para(doc, '经济价值：通过精准客流预测优化景区资源配置，预计可提升景区接待效率20%以上，降低人力调度成本。文创电商模块打通线上销售渠道，拓展景区收入来源。')
    add_para(doc, '社会价值：红色研学教育模块创新传承方式，服务"三线精神"文化传承。AI导游降低导游人力成本，提升游客服务覆盖面。系统可复制推广至其他欠发达地区景区。')
    add_para(doc, '学术价值：双流融合预测模型和RAG增强AI导游方案具有理论创新性，可形成学术论文和专利申请。')

    add_heading(doc, '9.3 代码规模统计', level=2)

    make_three_line_table(doc,
        headers=["子系统", "语言", "文件数", "主要内容"],
        rows=[
            ["Web业务后端", "Java", "301", "56个Controller + Service + Mapper + Entity"],
            ["AI智能后端", "Java", "46", "AI对话/RAG/TTS代理/WebSocket"],
            ["小程序后端", "Java", "67", "13个Controller + 电商/支付/研学"],
            ["API网关", "Java", "20+", "Gateway路由/限流/认证/熔断"],
            ["Web前端", "Vue/TS", "74", "55+个页面(含组件) + 三角色系统"],
            ["微信小程序", "Vue 3", "28+", "16个页面 + 12个API模块"],
            ["客流预测", "Python", "8", "4种预测模型 + 数据治理"],
            ["数字人", "JS/Python", "15+", "Three.js + MapLibre + TTS"],
        ],
        caption="表9  项目代码规模统计",
        col_widths=[3, 2, 1.5, 6],
    )

    add_page_break(doc)

    # ============================================================
    # 十、项目可行性分析
    # ============================================================
    add_heading(doc, '十、项目可行性分析', level=1)

    add_heading(doc, '10.1 技术可行性', level=2)
    add_para(doc, '本项目所采用的技术栈均为业界成熟且广泛使用的开源框架。Spring Boot 3.2、Vue 3、FastAPI等均有完善的社区支持和文档。通义千问DashScope SDK提供稳定的AI能力接入。TensorFlow和statsmodels为预测模型提供可靠的计算框架。团队成员具备Java、Python、Vue全栈开发能力，技术储备充分。')

    add_heading(doc, '10.2 经济可行性', level=2)
    add_para(doc, '项目开发使用开源技术栈，无软件授权费用。部署环境可使用云服务器（阿里云ECS 2核4G约100元/月），数据库、Redis等依赖可在同一服务器运行。阿里云OSS存储费用极低（约0.12元/GB/月）。通义千问API有免费额度，开发阶段费用可忽略。总体硬件与服务成本可控。')

    add_heading(doc, '10.3 应用可行性', level=2)
    add_para(doc, '六盘水市拥有丰富的旅游资源和明确的数字化转型需求，是本项目的理想试点区域。系统采用微服务架构，可按模块渐进式部署，降低落地风险。小程序端采用Uni-app跨平台开发，可同时输出H5和微信原生版本，覆盖最广泛的用户群体。')

    add_page_break(doc)

    # ============================================================
    # 附录：核心数据三线表
    # ============================================================
    add_heading(doc, '附录：核心数据表', level=1)

    add_heading(doc, 'A.1 景区基础数据表', level=2)
    make_three_line_table(doc,
        headers=["景区ID", "景区名称", "日均基础客流", "最大承载量", "海拔(m)"],
        rows=[
            ["1", "梅花山风景区", "3000", "8000", "2400"],
            ["2", "玉舍国家森林公园", "2500", "6000", "2300"],
            ["3", "乌蒙大草原", "3500", "10000", "2857"],
            ["4", "水城古镇", "2000", "5000", "1800"],
            ["5", "明湖国家湿地公园", "1800", "4000", "1750"],
        ],
        caption="附表1  五大景区基础参数",
        col_widths=[1.5, 4, 2.5, 2.5, 2],
        source="data_generator.py SCENIC_CONFIG"
    )

    add_heading(doc, 'A.2 客流影响因子表', level=2)
    make_three_line_table(doc,
        headers=["因子类型", "因子名称", "影响系数", "说明"],
        rows=[
            ["天气", "晴天", "×1.2 (P=0.40)", "客流增加20%"],
            ["天气", "多云", "×1.0 (P=0.30)", "基准无影响"],
            ["天气", "阴天", "×0.9 (P=0.15)", "客流减少10%"],
            ["天气", "雨天", "×0.6 (P=0.15)", "客流减少40%"],
            ["季节", "夏季(6-8月)", "×1.3", "凉都旅游旺季"],
            ["季节", "春秋", "×1.1", "温和季节"],
            ["季节", "冬季", "×0.8", "淡季"],
            ["日期", "周末", "×1.5", "weekday()≥5"],
            ["日期", "节假日", "×2.0", "HOLIDAYS列表"],
            ["随机", "随机波动", "U(0.9,1.1)", "模拟真实波动"],
        ],
        caption="附表2  客流量影响因子",
        col_widths=[2, 3, 3, 4.5],
        source="data_generator.py WEATHER_TYPES + calculate_daily_flow()"
    )

    add_heading(doc, 'A.3 数据治理参数表', level=2)
    make_three_line_table(doc,
        headers=["技术", "参数", "取值", "作用"],
        rows=[
            ["小波去噪", "wavelet", "db4", "Daubechies 4小波基"],
            ["小波去噪", "level", "2", "分解层数"],
            ["小波去噪", "threshold", "σ√(2log(n))", "通用阈值公式"],
            ["小波去噪", "mode", "soft", "软阈值处理"],
            ["差分隐私", "epsilon", "1.0", "隐私预算"],
            ["差分隐私", "mechanism", "Laplace", "拉普拉斯噪声机制"],
            ["归一化", "method", "MinMaxScaler", "最小-最大缩放"],
            ["归一化", "range", "(0,1)", "输出范围"],
        ],
        caption="附表3  数据治理参数配置",
        col_widths=[2.5, 3, 3, 4],
        source="preprocessing.py + lstm_new.py"
    )

    add_heading(doc, 'A.4 数据库核心表结构', level=2)
    make_three_line_table(doc,
        headers=["表名", "说明", "关键字段"],
        rows=[
            ["users", "用户表(三角色)", "id/username/role/email/status"],
            ["scenic_spots", "景区主表", "id/name/category/status/location"],
            ["scenic_images", "景区图片", "id/scenic_id/image_url/image_type"],
            ["reviews", "评价表", "id/user_id/scenic_id/rating/content"],
            ["ticket_orders", "门票订单", "id/user_id/scenic_id/status/amount"],
            ["merchant_profiles", "商户资料", "id/user_id/company_name/status"],
            ["banners", "轮播广告", "id/title/image_url/target_url/enabled"],
            ["mp_orders", "小程序订单", "id/user_id/order_no/total_amount/status"],
            ["mp_favorites", "收藏表", "id/user_id/scenic_id/created_at"],
            ["mp_feedback", "用户反馈", "id/user_id/type/content/status"],
            ["study_quiz", "研学题库", "id/question/optionA-D/answer/scenic_name"],
            ["study_badges", "研学徽章", "id/name/icon_char/color/description"],
            ["user_points", "用户积分", "id/user_id/total_points/used_points"],
            ["flow_records", "客流记录", "id/scenic_id/date/flow_count/weather"],
        ],
        caption="附表4  数据库核心表结构",
        col_widths=[3, 3, 6.5],
        source="MySQL travel_prediction 数据库"
    )

    # ============================================================
    # 保存
    # ============================================================
    doc.save(OUTPUT_PATH)
    print(f"[OK] 挑战杯项目计划书已生成: {OUTPUT_PATH}")
    print(f"     含封面 + 目录 + 10个章节 + 附录(4张三线表)")
    print(f"     正文三线表: 9张")
    print(f"     总计: 13张三线表")


if __name__ == '__main__':
    print("=" * 60)
    print("  游韵华章 · 挑战杯项目计划书生成")
    print("=" * 60)
    main()
    print("=" * 60)
