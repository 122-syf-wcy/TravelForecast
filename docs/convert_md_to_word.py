# -*- coding: utf-8 -*-
"""
将项目文档 Markdown 转换为 Word 格式
使用 python-docx 直接构建，确保格式规范
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

INPUT_MD = os.path.join(os.path.dirname(__file__), '项目文档_挑战杯与计算机设计大赛.md')
OUTPUT_DOCX = os.path.join(os.path.dirname(__file__), '项目文档_挑战杯与计算机设计大赛.docx')
CHARTS_DIR = os.path.join(os.path.dirname(__file__), 'charts')


def set_run_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, style=None, font_name='宋体', size=10.5, bold=False, align=None, indent=False):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    if top:
        t = parse_xml(f'<w:top {nsdecls("w")} w:val="{top["val"]}" w:sz="{top["sz"]}" w:color="{top["color"]}" w:space="0"/>')
        tcBorders.append(t)
    if bottom:
        b = parse_xml(f'<w:bottom {nsdecls("w")} w:val="{bottom["val"]}" w:sz="{bottom["sz"]}" w:color="{bottom["color"]}" w:space="0"/>')
        tcBorders.append(b)
    if start:
        s = parse_xml(f'<w:start {nsdecls("w")} w:val="{start["val"]}" w:sz="{start["sz"]}" w:color="{start["color"]}" w:space="0"/>')
        tcBorders.append(s)
    if end:
        e = parse_xml(f'<w:end {nsdecls("w")} w:val="{end["val"]}" w:sz="{end["sz"]}" w:color="{end["color"]}" w:space="0"/>')
        tcBorders.append(e)
    tcPr.append(tcBorders)


def make_three_line_table(doc, headers, rows):
    """创建三线表"""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin  = {"val": "single", "sz": "6",  "color": "000000"}
    none  = {"val": "nil",    "sz": "0",  "color": "000000"}

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, '宋体', 10, True)
        set_cell_border(cell, top=thick, bottom=thin, start=none, end=none)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, '宋体', 10, False)
            if i == len(rows) - 1:
                set_cell_border(cell, top=none, bottom=thick, start=none, end=none)
            else:
                set_cell_border(cell, top=none, bottom=none, start=none, end=none)

    doc.add_paragraph()
    return table


def add_image(doc, filename):
    path = os.path.join(CHARTS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def parse_md_table(lines):
    """解析 MD 表格行，返回 (headers, rows)"""
    headers = [c.strip().replace('**', '') for c in lines[0].strip('|').split('|')]
    rows = []
    for line in lines[2:]:  # skip separator
        cells = [c.strip().replace('**', '') for c in line.strip('|').split('|')]
        rows.append(cells)
    return headers, rows


def main():
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 逐行解析 MD ----
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(md_lines):
        line = md_lines[i].rstrip('\n')

        # 代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.size = Pt(9)
                run.font.name = 'Consolas'
                p.paragraph_format.left_indent = Cm(1)
                # 加浅灰背景
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                p._element.get_or_add_pPr().append(shading)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            i += 1
            continue

        # 标题
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            if level == 1:
                # 一级标题作为文档大标题
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_run_font(run, '黑体', 22, True)
            elif level == 2:
                add_heading_styled(doc, text, 1)
            elif level == 3:
                add_heading_styled(doc, text, 2)
            elif level == 4:
                add_heading_styled(doc, text, 3)
            i += 1
            continue

        # 引用行
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            # 去掉 markdown bold
            text = text.replace('**', '')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(text)
            set_run_font(run, '楷体', 10, False, (100, 100, 100))
            i += 1
            continue

        # 表格
        if '|' in line and i + 1 < len(md_lines) and '---' in md_lines[i + 1]:
            table_lines = []
            j = i
            while j < len(md_lines) and '|' in md_lines[j].strip():
                table_lines.append(md_lines[j].strip())
                j += 1
            if len(table_lines) >= 3:
                headers, rows = parse_md_table(table_lines)
                make_three_line_table(doc, headers, rows)
            i = j
            continue

        # 图片引用
        if '详见图表' in line or 'charts/' in line:
            # 提取图片文件名
            match = re.search(r'charts/([^`\s]+\.png)', line)
            if match:
                fname = match.group(1)
                added = add_image(doc, fname)
                if added:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f'图：{fname.replace(".png", "").split("_", 1)[-1]}')
                    set_run_font(run, '宋体', 9, False, (128, 128, 128))
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            text = m.group(2).replace('**', '')
            p = doc.add_paragraph(style='List Number')
            p.clear()
            run = p.add_run(text)
            set_run_font(run, '宋体', 10.5)
            i += 1
            continue

        # 无序列表
        if line.startswith('- ') or line.startswith('  - '):
            text = line.lstrip(' -').strip().replace('**', '')
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(text)
            set_run_font(run, '宋体', 10.5)
            i += 1
            continue

        # 普通段落
        text = line.strip().replace('**', '')
        if text:
            add_para(doc, text, indent=True)
        i += 1

    # 保存
    doc.save(OUTPUT_DOCX)
    print(f"[OK] Word文档已生成: {OUTPUT_DOCX}")


if __name__ == '__main__':
    print("=" * 60)
    print("  MD → Word 转换")
    print("=" * 60)
    main()
    print("=" * 60)
