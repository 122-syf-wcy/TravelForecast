# -*- coding: utf-8 -*-
"""
智教黔行 · 完整技术文档 Word生成
便于编写参赛文档、PPT和计划书的参考资料
包含：系统概述、架构详解、功能清单、API接口、数据库设计、部署方案、三线表汇总等
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '完整技术文档_智教黔行.docx')


# ============================================================
# 工具函数（与 generate_project_plan.py 一致）
# ============================================================

def set_run_font(run, name='宋体', size=10.5, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(doc, text, size=10.5, bold=False, indent=True, align=None, font_name='宋体', spacing_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    set_run_font(run, '宋体', size)
    return p


def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for tag, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        if val:
            el = parse_xml(f'<w:{tag} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" w:color="{val["color"]}" w:space="0"/>')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def make_three_line_table(doc, headers, rows, caption="", col_widths=None, source=""):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(caption)
        set_run_font(run, '宋体', 10, True)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin  = {"val": "single", "sz": "6",  "color": "000000"}
    none  = {"val": "nil",    "sz": "0",  "color": "000000"}

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, '宋体', 10, True)
        set_cell_border(cell, top=thick, bottom=thin, start=none, end=none)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, '宋体', 10, False)
            if i == len(rows) - 1:
                set_cell_border(cell, top=none, bottom=thick, start=none, end=none)
            else:
                set_cell_border(cell, top=none, bottom=none, start=none, end=none)

    if source:
        p = doc.add_paragraph()
        run = p.add_run(f"数据来源：{source}")
        set_run_font(run, '宋体', 9, False, (128, 128, 128), italic=True)
    else:
        doc.add_paragraph()
    return table


# ============================================================
# 主逻辑
# ============================================================

def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智教黔行 · 智慧文旅一体化平台')
    set_run_font(run, '黑体', 26, True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('完整技术文档')
    set_run_font(run, '黑体', 20, True, (42, 157, 143))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（参赛文档 · PPT素材 · 计划书参考资料）')
    set_run_font(run, '楷体', 14, False, (100, 100, 100))

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('最后更新：2026年3月')
    set_run_font(run, '宋体', 12, False, (128, 128, 128))

    doc.add_page_break()

    # ============================================================
    # 目录
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    set_run_font(run, '黑体', 18, True)
    doc.add_paragraph()

    toc = [
        "第一章  项目概述",
        "  1.1 项目简介",
        "  1.2 项目定位与目标",
        "  1.3 项目亮点速览（PPT用）",
        "第二章  系统架构设计",
        "  2.1 微服务架构总览",
        "  2.2 服务间调用关系",
        "  2.3 技术栈全景",
        "  2.4 安全与基础设施",
        "第三章  功能模块详解",
        "  3.1 Web管理端——管理员功能",
        "  3.2 Web管理端——商户功能",
        "  3.3 Web管理端——用户(游客)功能",
        "  3.4 微信小程序端",
        "  3.5 小程序后端API",
        "第四章  AI智能服务",
        "  4.1 AI对话系统",
        "  4.2 RAG知识库",
        "  4.3 AI行程规划",
        "  4.4 TTS语音合成",
        "第五章  客流预测系统",
        "  5.1 预测模型对比",
        "  5.2 ARIMA模型",
        "  5.3 LSTM模型",
        "  5.4 双流融合模型",
        "  5.5 数据治理",
        "  5.6 特征工程",
        "第六章  3D数字人系统",
        "第七章  数据库设计",
        "第八章  部署架构",
        "第九章  项目数据汇总（三线表）",
        "第十章  PPT制作素材",
    ]
    for item in toc:
        indent = item.startswith("  ")
        text = item.strip()
        add_para(doc, text, size=11 if not indent else 10.5, indent=indent, bold=not indent, spacing_after=4)

    doc.add_page_break()

    # ============================================================
    # 第一章 项目概述
    # ============================================================
    add_heading(doc, '第一章  项目概述', level=1)

    add_heading(doc, '1.1 项目简介', level=2)
    add_para(doc, '"智教黔行·智游六盘水"是一个面向贵州六盘水市（中国凉都）打造的智慧文旅一体化平台，集"AI数字伴游+沉浸式导览+红色研学+客流预测+文创电商"于一体，覆盖景区管理方（管理员）、景区运营方（商户）、游客（用户）三大角色，贯穿"游前规划—游中体验—游后消费"全流程。')
    add_para(doc, '项目以梅花山风景区、玉舍国家森林公园、乌蒙大草原、水城古镇、明湖国家湿地公园五大景区为核心场景，构建了由8个独立微服务组成的完整系统，技术栈横跨Java、Python、JavaScript/TypeScript三大语言生态。')

    add_heading(doc, '1.2 项目定位与目标', level=2)
    targets = [
        "为景区管理者提供数据驱动的决策支持工具（客流预测、数据大屏、政策沙盘）",
        "为商户提供数字化经营管理平台（订单、评价、数据分析、合同管理）",
        "为游客提供AI驱动的沉浸式旅游体验（数字导游、行程规划、研学教育）",
        "探索AI大模型+深度学习在文旅垂直领域的创新应用范式",
        "为欠发达地区景区提供低成本、可复制的智慧文旅解决方案",
    ]
    for t in targets:
        add_bullet(doc, t)

    add_heading(doc, '1.3 项目亮点速览（PPT用）', level=2)
    add_para(doc, '以下内容可直接用于PPT关键页面的要点提炼：', bold=True, indent=False)

    highlights = [
        "【架构】8个微服务 + API网关统一治理，支持水平扩展",
        "【AI】通义千问大模型 + RAG知识库，实现专业级智能导游",
        "【预测】ARIMA+LSTM双流融合，准确率92%，支持5景区×30天预测",
        "【数字人】Three.js 3D模型 + MapLibre地图 + TTS口型同步",
        "【研学】路线→答题→积分→徽章→护照，红色教育游戏化闭环",
        "【电商】商品→购物车→订单→支付(模拟+黔豆)→发货，完整电商链路",
        "【管理】管理员+商户+用户三角色系统，39个Web页面+16个小程序页面",
        "【安全】JWT认证 + 令牌桶限流 + Resilience4j熔断 + 差分隐私",
        "【数据】ECharts多维图表 + 3D数据大屏 + 热力图 + 实时监控",
        "【跨端】Web端(Vue 3) + 微信小程序(Uni-app)，一套后端服务多端接入",
    ]
    for h in highlights:
        add_bullet(doc, h)

    doc.add_page_break()

    # ============================================================
    # 第二章 系统架构设计
    # ============================================================
    add_heading(doc, '第二章  系统架构设计', level=1)

    add_heading(doc, '2.1 微服务架构总览', level=2)

    make_three_line_table(doc,
        headers=["序号", "服务名称", "端口", "技术栈", "核心职责", "状态"],
        rows=[
            ["1", "TravelForecastGateway", "8888", "Spring Cloud Gateway", "JWT认证/限流/路由/熔断", "✅"],
            ["2", "TravelForecastBackend", "8080", "Spring Boot 3.2 + Java 17", "管理员/商户/用户全业务(56 Controller)", "✅"],
            ["3", "TravelForecastingAIBackend", "8081", "Spring Boot 3.2 + DashScope", "AI对话/RAG/行程规划/TTS/研学", "✅"],
            ["4", "TravelForecastMiniProgramBackend", "8082", "Spring Boot 3.2 + MyBatis-Plus", "电商/积分/支付/研学/反馈(13 Controller)", "✅"],
            ["5", "PythonPredictionService", "8001", "FastAPI + TensorFlow", "ARIMA/LSTM/双流融合预测", "✅"],
            ["6", "DigitalHuman", "8083", "FastAPI + Three.js", "3D数字人/TTS语音/场景讲解", "✅"],
            ["7", "TravelForecastFrontend", "3000", "Vue 3 + TS + Element Plus", "三角色Web管理端(74个Vue文件)", "✅"],
            ["8", "TravelForecastMiniProgram", "H5/微信", "Uni-app + Vue 3", "游客小程序端(16页面)", "✅"],
        ],
        caption="表1  系统微服务架构全景",
        col_widths=[0.8, 3.5, 1, 3.5, 3.5, 0.8],
        source="各服务 application.yml / package.json"
    )

    add_heading(doc, '2.2 服务间调用关系', level=2)
    add_para(doc, '系统服务间调用遵循以下规则：所有外部请求经API网关(8888端口)统一接入，网关负责JWT认证、限流和路由分发。各业务后端通过HTTP调用AI服务和预测服务。前端与小程序只与网关通信，不直接访问后端服务。', indent=True)

    make_three_line_table(doc,
        headers=["路由ID", "匹配路径", "转发目标", "限流(req/s)", "熔断等待"],
        rows=[
            ["business-service", "/api/**", "localhost:8080", "100/200", "5s"],
            ["ai-service", "/ai-api/**", "localhost:8081", "50/100", "10s"],
            ["miniprogram-service", "/miniprogram-api/**", "localhost:8082", "80/150", "-"],
            ["prediction-service", "/prediction-api/**", "localhost:8001", "30/60", "15s"],
            ["digital-human-service", "/digital-human-api/**", "localhost:8083", "40/80", "-"],
            ["digital-human-ws", "/ws/**", "ws://localhost:8083", "-", "-"],
        ],
        caption="表2  API网关路由与限流配置",
        col_widths=[3.5, 3, 2.5, 1.8, 1.5],
        source="TravelForecastGateway/application.yml"
    )

    add_heading(doc, '2.3 技术栈全景', level=2)

    make_three_line_table(doc,
        headers=["类别", "技术", "版本", "用途"],
        rows=[
            ["后端框架", "Spring Boot", "3.2.0", "Java微服务框架"],
            ["后端语言", "Java", "17 (LTS)", "主力开发语言"],
            ["API网关", "Spring Cloud Gateway", "2023.0", "统一网关+限流+熔断"],
            ["ORM", "MyBatis-Plus", "3.5.5", "数据访问层"],
            ["数据库", "MySQL", "8.0+", "关系型存储"],
            ["缓存", "Redis", "6.0+", "缓存/限流/会话"],
            ["认证", "JWT (jjwt)", "0.12.3", "Token认证"],
            ["熔断", "Resilience4j", "2.1.0", "服务降级保护"],
            ["前端框架", "Vue 3 + TypeScript", "3.3", "Web管理端"],
            ["UI组件", "Element Plus", "2.3", "企业级组件库"],
            ["CSS", "TailwindCSS", "3.3", "原子化CSS"],
            ["图表", "ECharts + ECharts GL", "5.4", "多维数据可视化"],
            ["地图", "Mapbox GL + MapLibre", "2.15/3.6", "3D地图渲染"],
            ["3D", "Three.js", "0.157", "3D数字人渲染"],
            ["动画", "GSAP + Lottie", "3.12", "交互动画"],
            ["状态管理", "Pinia", "2.1", "Vue状态管理"],
            ["小程序", "Uni-app + Vue 3", "-", "跨平台小程序"],
            ["Python框架", "FastAPI", "0.115", "预测/数字人API"],
            ["深度学习", "TensorFlow", "≥2.15", "LSTM模型"],
            ["时序分析", "statsmodels", "≥0.14", "ARIMA模型"],
            ["小波变换", "PyWavelets", "-", "数据去噪"],
            ["AI大模型", "通义千问 DashScope", "2.12", "AI对话/RAG"],
            ["对象存储", "阿里云OSS", "3.17.4", "文件/图片存储"],
            ["HTTP工具", "Hutool", "5.8.24", "Java工具库"],
        ],
        caption="表3  技术栈依赖版本全景",
        col_widths=[2, 3.5, 1.8, 5],
        source="pom.xml / requirements.txt / package.json"
    )

    add_heading(doc, '2.4 安全与基础设施', level=2)

    make_three_line_table(doc,
        headers=["功能", "实现方式", "说明"],
        rows=[
            ["JWT认证", "JwtInterceptor", "Token验证+用户信息注入+白名单"],
            ["角色权限", "RoleInterceptor", "三角色(user/merchant/admin)访问控制"],
            ["接口限流", "RateLimitInterceptor", "基于Redis令牌桶算法的差异化限流"],
            ["服务熔断", "Resilience4j CircuitBreaker", "服务不可用时自动降级"],
            ["操作日志", "OperationLogAspect", "AOP注解式操作审计日志"],
            ["请求统计", "RequestStatsInterceptor", "API请求耗时统计"],
            ["CORS跨域", "Gateway CORS配置", "多域名白名单"],
            ["数据备份", "BackupController", "MySQL数据备份与恢复"],
            ["差分隐私", "preprocessing.py", "Laplace噪声(ε=1.0)保护客流数据"],
        ],
        caption="表4  安全与基础设施",
        col_widths=[2.5, 4, 6],
        source="各服务拦截器/AOP/配置文件"
    )

    doc.add_page_break()

    # ============================================================
    # 第三章 功能模块详解
    # ============================================================
    add_heading(doc, '第三章  功能模块详解', level=1)

    add_heading(doc, '3.1 Web管理端——管理员功能 (15个页面)', level=2)

    make_three_line_table(doc,
        headers=["页面", "路由", "功能描述"],
        rows=[
            ["总览仪表盘", "/admin/dashboard/overview", "全局KPI/趋势图/实时数据"],
            ["数据分析", "/admin/dashboard/analytics", "深度多维数据分析"],
            ["系统监控", "/admin/dashboard/monitor", "系统性能/API监控"],
            ["商户列表", "/admin/business/list", "全部商户管理"],
            ["审核管理", "/admin/business/audit", "商户入驻审核"],
            ["审核日志", "/admin/business/audit-logs", "审核操作记录"],
            ["合同管理", "/admin/business/contract", "商户合同管理"],
            ["景区分配", "/admin/business/merchant-scenics", "商户-景区关联"],
            ["用户列表", "/admin/user/list", "用户CRUD/导出CSV"],
            ["行为分析", "/admin/user/behavior", "用户行为统计"],
            ["首页配置", "/admin/content/landing", "Landing页内容管理"],
            ["景区管理", "/admin/content/scenic", "景区CRUD/媒体管理"],
            ["轮播广告", "/admin/content/banner", "Web+小程序轮播"],
            ["文创商品", "/admin/content/mp-product", "文创商品管理"],
            ["系统配置", "/admin/system/settings", "全局配置/验证码/邮件"],
        ],
        caption="表5  管理员端页面清单",
        col_widths=[2.5, 4.5, 5.5],
    )

    add_heading(doc, '3.2 Web管理端——商户功能 (14个页面)', level=2)

    make_three_line_table(doc,
        headers=["页面", "功能描述"],
        rows=[
            ["经营仪表盘", "KPI指标/客流趋势/经营数据图表"],
            ["实时监测", "实时客流/设备状态"],
            ["数据分析", "多维经营分析"],
            ["资源管理", "景区/子景点/设施/活动管理"],
            ["评价系统", "评价管理/回复/分析"],
            ["新闻管理", "新闻发布/编辑"],
            ["门票订单", "订单管理/核销"],
            ["紧急救援", "SOS接警/处理"],
            ["商家资料", "商家信息编辑"],
            ["统计报表", "统计图表"],
            ["客流预测", "预测数据查看"],
            ["政策沙盘", "政策模拟"],
            ["数据导出", "CSV数据导出"],
            ["账号设置", "密码修改/安全设置"],
        ],
        caption="表6  商户端页面清单",
        col_widths=[3, 9.5],
    )

    add_heading(doc, '3.3 Web管理端——用户(游客)功能 (10个页面)', level=2)

    make_three_line_table(doc,
        headers=["页面", "功能描述"],
        rows=[
            ["首页(Landing)", "沉浸式首页/轮播/实景预览/公告"],
            ["用户仪表盘", "个人中心/快捷入口/收藏列表"],
            ["景区探索", "3D地图探索/景区详情/评价"],
            ["客流预测", "多模型预测图表/最佳时段推荐"],
            ["行程规划", "AI智能行程生成/路线优化"],
            ["实时服务", "实时天气/交通/景区状态"],
            ["AI服务", "AI智能对话入口"],
            ["个人中心", "个人信息/头像/密码修改"],
            ["旅游资讯", "新闻浏览"],
            ["热门景点", "热门排行"],
        ],
        caption="表7  用户端页面清单",
        col_widths=[3, 9.5],
    )

    add_heading(doc, '3.4 微信小程序端 (16个页面)', level=2)

    make_three_line_table(doc,
        headers=["页面", "路径", "核心功能", "状态"],
        rows=[
            ["首页", "index/index", "沉浸式首页/金刚区入口/轮播", "✅"],
            ["景区详情", "spot/detail", "景区信息/图片/评价/收藏", "✅"],
            ["搜索", "search/index", "景区关键词搜索", "✅"],
            ["智能导览", "guide/index", "地图导览/AR标注/答题", "✅"],
            ["数字导游", "digital-human/index", "AI对话/TTS/语音输入", "✅"],
            ["行程列表", "itinerary/list", "AI行程/时间线展示", "✅"],
            ["文创商城", "shop/index", "商品列表/分类/黔豆", "✅"],
            ["商品详情", "shop/detail", "商品信息/加入购物车/分享", "✅"],
            ["订单管理", "order/list", "订单/支付/取消/删除", "✅"],
            ["红色研学", "red-study/index", "路线/答题/徽章/护照", "✅"],
            ["本地服务", "life-service/index", "9类本地服务信息", "✅"],
            ["个人中心", "profile/index", "签到/打卡/收藏/订单", "✅"],
            ["设置", "profile/settings", "缓存/隐私/关于/退出", "✅"],
            ["地址管理", "profile/address", "收货地址CRUD", "✅"],
            ["意见反馈", "profile/feedback", "反馈提交(对接后端)", "✅"],
            ["我的收藏", "profile/favorites", "收藏列表/取消收藏", "✅"],
        ],
        caption="表8  小程序端页面清单",
        col_widths=[2, 3, 5, 1],
    )

    add_heading(doc, '3.5 小程序后端API (13个Controller)', level=2)

    make_three_line_table(doc,
        headers=["Controller", "路由前缀", "功能"],
        rows=[
            ["WechatLoginController", "/api/wechat", "微信OAuth登录"],
            ["UserController", "/api/user", "用户信息/收藏"],
            ["HomeController", "/api/home", "首页配置数据"],
            ["ScenicSpotController", "/api/spots", "景区列表/详情/搜索"],
            ["ShopController", "/api/shop", "商品/购物车/订单"],
            ["PaymentController", "/api/pay", "微信支付/黔豆支付/开发模拟"],
            ["ItineraryController", "/api/itinerary", "行程CRUD"],
            ["StudyController", "/api/study", "题库/答题/徽章/积分/护照"],
            ["LifeServiceController", "/api/life-service", "10类本地服务(交通/酒店/美食等)"],
            ["FeedbackController", "/api/feedback", "用户反馈提交"],
            ["OssProxyController", "/api/oss-proxy", "OSS图片代理"],
            ["AdminBannerController", "/api/admin/banners", "小程序轮播管理"],
            ["AdminProductController", "/api/admin/products", "文创商品管理"],
        ],
        caption="表9  小程序后端Controller清单",
        col_widths=[4, 3, 5.5],
        source="TravelForecastMiniProgramBackend/src/main/java/com/travel/controller/"
    )

    doc.add_page_break()

    # ============================================================
    # 第四章 AI智能服务
    # ============================================================
    add_heading(doc, '第四章  AI智能服务', level=1)

    add_heading(doc, '4.1 AI对话系统', level=2)
    add_para(doc, '系统集成阿里云通义千问大模型(DashScope SDK 2.12)，通过Spring Boot封装为RESTful API。支持多轮会话管理（基于Redis存储会话上下文），单次对话超时30秒。AI后端同时服务Web端和小程序端。')

    add_heading(doc, '4.2 RAG知识库', level=2)
    add_para(doc, '采用RAG(Retrieval-Augmented Generation)技术增强大模型在文旅垂直领域的回答质量。系统维护六盘水景区专属知识库，对话时先通过向量检索定位相关知识文档，再将检索结果作为上下文注入提示词，引导大模型生成准确、专业的回答。有效解决大模型"知识幻觉"问题。')

    add_heading(doc, '4.3 AI行程规划', level=2)
    add_para(doc, '用户输入旅行偏好（天数、兴趣、预算等），系统通过精心设计的提示工程(Prompt Engineering)引导通义千问生成个性化行程方案，包含每日景点安排、时间分配、交通建议等。生成的行程自动保存到用户行程列表。')

    add_heading(doc, '4.4 TTS语音合成', level=2)
    add_para(doc, 'AI后端集成TTS(Text-to-Speech)服务，将AI回复文本转换为MP3格式语音。小程序端通过条件编译实现双模式兼容：微信原生模式写入临时文件播放，H5模式转换为Blob URL播放。支持点击任意AI消息重新播放语音。')

    make_three_line_table(doc,
        headers=["接口", "方法", "路径", "功能"],
        rows=[
            ["AI对话", "POST", "/ai-api/chat/message", "发送消息获取AI回复"],
            ["会话历史", "GET", "/ai-api/chat/history/{id}", "获取会话聊天记录"],
            ["清除会话", "DELETE", "/ai-api/chat/history/{id}", "清除指定会话"],
            ["知识检索", "POST", "/ai-api/knowledge/search", "RAG知识库检索"],
            ["知识问答", "GET", "/ai-api/knowledge/answer", "基于知识库回答"],
            ["行程规划", "POST", "/ai-api/ai-planning/generate", "AI生成个性化行程"],
            ["研学路线", "GET", "/ai-api/education/routes", "红色研学路线"],
            ["研学方案", "POST", "/ai-api/education/generate", "AI生成研学方案"],
            ["语音合成", "POST", "/ai-api/speech/tts", "文本转语音"],
            ["WebSocket", "WS", "/ai-api/ws/digital-human", "数字人实时通信"],
        ],
        caption="表10  AI智能服务API接口",
        col_widths=[2, 1.5, 4.5, 4.5],
        source="TravelForecastingAIBackend API"
    )

    doc.add_page_break()

    # ============================================================
    # 第五章 客流预测系统
    # ============================================================
    add_heading(doc, '第五章  客流预测系统', level=1)

    add_heading(doc, '5.1 预测模型对比', level=2)

    make_three_line_table(doc,
        headers=["模型", "准确率", "置信度", "输入维度", "序列长度", "网络结构", "训练数据"],
        rows=[
            ["ARIMA", "82%", "80%", "1维(客流)", "-", "SARIMAX(1,1,1)(1,1,1,7)", "90天"],
            ["LSTM(基础)", "87%", "85%", "1维(客流)", "14天", "LSTM(50)→DO→LSTM(50)→Dense", "90天"],
            ["LSTM(多变量)", "-", "-", "6维特征", "14天", "LSTM(64)→DO→LSTM(32)→Dense", "365天"],
            ["双流融合", "92%", "92%", "6维+ARIMA", "14天", "ARIMA+LSTM+动态权重", "365天"],
        ],
        caption="表11  四种预测模型性能对比",
        col_widths=[2, 1.2, 1.2, 1.8, 1.5, 4, 1.5],
        source="arima_model.py / lstm_model.py / dual_stream_model.py"
    )

    add_heading(doc, '5.2 ARIMA模型', level=2)
    add_para(doc, 'ARIMA(自回归积分滑动平均)模型采用SARIMAX(1,1,1)(1,1,1,7)结构，其中季节性周期为7天（一周），适合捕捉客流的周循环规律。模型在90天历史数据上训练，输出7天短期预测，准确率82%。')

    add_heading(doc, '5.3 LSTM模型', level=2)
    add_para(doc, '长短时记忆网络(LSTM)模型分为基础版(单变量)和增强版(6维多变量)。基础版使用14天滑动窗口输入单一客流序列，两层LSTM(各50单元)+Dropout(0.2)+全连接层。多变量版扩展为6维特征输入(客流/节庆/周末/天气/温度/海拔)，采用LSTM(64)+LSTM(32)+Dense(16)+Dense(1)结构，使用MinMaxScaler归一化。')

    add_heading(doc, '5.4 双流融合模型', level=2)
    add_para(doc, '本项目提出的创新模型。ARIMA流捕捉线性趋势，LSTM流建模非线性关系，两个流的预测结果通过动态权重机制融合：权重根据近期预测误差（MAE）自适应调整，误差较小的模型获得更高权重。融合后预测准确率达92%，较单一模型显著提升。')

    add_heading(doc, '5.5 数据治理', level=2)

    make_three_line_table(doc,
        headers=["技术", "参数", "取值", "作用", "代码位置"],
        rows=[
            ["小波去噪", "wavelet", "db4", "Daubechies 4小波基", "preprocessing.py:27"],
            ["小波去噪", "level", "2", "分解层数", "preprocessing.py:27"],
            ["小波去噪", "threshold", "σ√(2log(n))", "通用阈值公式", "preprocessing.py:46"],
            ["小波去噪", "mode", "soft", "软阈值处理", "preprocessing.py:54"],
            ["差分隐私", "epsilon", "1.0", "隐私预算", "preprocessing.py:73"],
            ["差分隐私", "sensitivity", "1.0", "数据敏感度", "preprocessing.py:73"],
            ["差分隐私", "mechanism", "Laplace", "拉普拉斯噪声", "preprocessing.py:82"],
            ["归一化", "method", "MinMaxScaler", "最小-最大缩放", "lstm_new.py:63"],
            ["归一化", "range", "(0, 1)", "输出范围", "lstm_new.py:63"],
        ],
        caption="表12  数据治理技术参数",
        col_widths=[2, 2.5, 2.5, 3, 3],
        source="utils/preprocessing.py + models/lstm_new.py"
    )

    add_heading(doc, '5.6 特征工程', level=2)

    make_three_line_table(doc,
        headers=["维度", "特征名称", "取值范围", "归一化", "来源"],
        rows=[
            ["Dim 0", "历史客流(flow)", "0~capacity", "MinMaxScaler", "flow_records表"],
            ["Dim 1", "节庆标记(is_holiday)", "0或1", "二值编码", "HOLIDAYS列表"],
            ["Dim 2", "周末标记(is_weekend)", "0或1", "二值编码", "weekday()≥5"],
            ["Dim 3", "天气指数(weather_idx)", "0.0~1.0", "type_index/3.0", "WEATHER_TYPES"],
            ["Dim 4", "温度(temp_avg)", "0.0~1.0", "temp_avg/40.0", "天气API"],
            ["Dim 5", "海拔(altitude)", "0.0~1.0", "altitude/3000", "SCENIC_CONFIG"],
        ],
        caption="表13  LSTM 6维特征向量",
        col_widths=[1.5, 3, 2.5, 2.5, 3],
        source="lstm_new.py + data_generator.py"
    )

    make_three_line_table(doc,
        headers=["因子类型", "因子名称", "系数", "概率/条件"],
        rows=[
            ["天气", "晴天(sunny)", "×1.2", "P=0.40"],
            ["天气", "多云(cloudy)", "×1.0", "P=0.30"],
            ["天气", "阴天(overcast)", "×0.9", "P=0.15"],
            ["天气", "雨天(rainy)", "×0.6", "P=0.15"],
            ["季节", "夏季(6-8月)", "×1.3", "凉都旅游旺季"],
            ["季节", "春秋(4-5,9-10月)", "×1.1", "温和季节"],
            ["季节", "冬季", "×0.8", "淡季"],
            ["日期", "周末", "×1.5", "weekday()≥5"],
            ["日期", "节假日", "×2.0", "HOLIDAYS列表"],
            ["随机", "随机波动", "U(0.9,1.1)", "random.uniform()"],
        ],
        caption="表14  客流量影响因子",
        col_widths=[2, 3, 2.5, 5],
        source="data_generator.py"
    )

    make_three_line_table(doc,
        headers=["时段", "8:00", "9:00", "10:00", "11:00", "12:00", "13:00",
                 "14:00", "15:00", "16:00", "17:00", "18:00", "19:00"],
        rows=[
            ["分布系数", "0.2", "0.4", "0.7", "0.9", "1.0", "0.9", "1.0", "0.95", "0.8", "0.6", "0.4", "0.2"],
        ],
        caption="表15  小时级客流分布系数",
        col_widths=[2] + [1]*12,
        source="data_generator.py HOURLY_DISTRIBUTION"
    )

    doc.add_page_break()

    # ============================================================
    # 第六章 3D数字人系统
    # ============================================================
    add_heading(doc, '第六章  3D数字人系统', level=1)

    add_para(doc, '3D数字人系统是本项目的特色功能之一，基于Three.js和MapLibre GL构建，实现了3D虚拟导游与地理信息系统的深度融合。')

    make_three_line_table(doc,
        headers=["功能模块", "技术实现", "说明"],
        rows=[
            ["3D模型加载", "Three.js GLTFLoader", "支持VRM/GLB格式人物模型"],
            ["地图集成", "MapLibre GL JS", "矢量地图+3D地形渲染"],
            ["坐标同步", "自定义坐标转换", "3D模型与地图经纬度同步"],
            ["口型动画", "Lipsync算法", "语音播放时嘴型同步"],
            ["TTS语音", "Python TTS服务", "AI回复实时转语音"],
            ["AI对话", "通义千问+RAG", "景区知识库增强对话"],
            ["H5兼容", "条件编译(#ifdef)", "微信/H5双模式音频播放"],
        ],
        caption="表16  3D数字人系统功能",
        col_widths=[2.5, 3.5, 6.5],
    )

    doc.add_page_break()

    # ============================================================
    # 第七章 数据库设计
    # ============================================================
    add_heading(doc, '第七章  数据库设计', level=1)

    add_para(doc, '系统使用MySQL 8.0关系型数据库，数据库名称为travel_prediction。ORM层采用MyBatis-Plus 3.5.5，实体字段采用驼峰命名，自动映射为数据库下划线命名。')

    add_heading(doc, '7.1 主业务数据库核心表', level=2)

    make_three_line_table(doc,
        headers=["表名", "说明", "关键字段"],
        rows=[
            ["users", "用户表(三角色)", "id/username/password/role/email/phone/status/avatar"],
            ["scenic_spots", "景区主表", "id/name/description/category/status/location/capacity"],
            ["scenic_images", "景区图片", "id/scenic_id/image_url/image_type(COVER/GALLERY)"],
            ["scenic_videos", "景区视频", "id/scenic_id/video_url/title"],
            ["scenic_sub_spots", "子景点", "id/scenic_id/name/description/location"],
            ["scenic_facilities", "景区设施", "id/scenic_id/name/type/status"],
            ["scenic_activities", "景区活动", "id/scenic_id/title/start_date/end_date"],
            ["reviews", "评价表", "id/user_id/scenic_id/rating/content/created_at"],
            ["ticket_orders", "门票订单", "id/user_id/scenic_id/status/total_amount"],
            ["merchant_profiles", "商户资料", "id/user_id/company_name/license_no/status"],
            ["merchant_audit_logs", "审核日志", "id/merchant_id/action/admin_id/remark"],
            ["merchant_contracts", "商户合同", "id/merchant_id/contract_no/start_date/end_date"],
            ["banners", "轮播广告", "id/title/image_url/target_url/sort_order/enabled"],
            ["showcases", "实景预览", "id/title/image_url/enabled"],
            ["news", "新闻资讯", "id/title/content/author/published_at"],
            ["announcements", "公告", "id/title/content/type/published_at"],
            ["roles", "角色表", "id/name/description"],
            ["permissions", "权限表", "id/name/resource/action"],
            ["operation_logs", "操作日志", "id/user_id/action/module/ip/created_at"],
            ["system_settings", "系统配置", "id/setting_key/setting_value/description"],
            ["user_favorites", "用户收藏", "id/user_id/scenic_id/created_at"],
            ["notifications", "通知消息", "id/user_id/title/content/is_read"],
            ["emergency_reports", "紧急救援", "id/reporter_id/scenic_id/type/status"],
        ],
        caption="表17  主业务数据库核心表",
        col_widths=[3, 2.5, 7],
        source="MySQL travel_prediction 数据库"
    )

    add_heading(doc, '7.2 小程序专属表', level=2)

    make_three_line_table(doc,
        headers=["表名", "说明", "关键字段"],
        rows=[
            ["mp_orders", "小程序订单", "id/user_id/order_no/product_name/total_amount/status/payment_time"],
            ["mp_favorites", "小程序收藏", "id/user_id/scenic_id/created_at"],
            ["mp_feedback", "用户反馈", "id/user_id/type/content/contact/status/created_at"],
            ["mp_banners", "小程序轮播", "id/title/image_url/link_type/link_value/sort/enabled"],
            ["mp_products", "文创商品", "id/name/price/stock/category/description/image_url"],
            ["mp_cart", "购物车", "id/user_id/product_id/quantity"],
            ["study_quiz", "研学题库", "id/question/optionA/optionB/optionC/optionD/answer/scenic_name"],
            ["study_badges", "研学徽章", "id/name/icon_char/color/description/sort_order"],
            ["user_badges", "用户徽章", "id/user_id/badge_id/unlocked_at"],
            ["user_points", "用户积分", "id/user_id/total_points/used_points"],
            ["study_answer_log", "答题记录", "id/user_id/quiz_id/user_answer/is_correct/points"],
            ["flow_records", "客流记录", "id/scenic_id/record_date/flow_count/weather/temperature"],
        ],
        caption="表18  小程序与预测专属表",
        col_widths=[3, 2.5, 7],
        source="SQL迁移脚本 V1-V3"
    )

    doc.add_page_break()

    # ============================================================
    # 第八章 部署架构
    # ============================================================
    add_heading(doc, '第八章  部署架构', level=1)

    make_three_line_table(doc,
        headers=["服务", "端口", "启动命令", "依赖"],
        rows=[
            ["MySQL", "3306", "系统服务", "-"],
            ["Redis", "6379", "系统服务", "-"],
            ["API网关", "8888", "java -jar gateway.jar", "Redis"],
            ["主业务后端", "8080", "java -jar backend.jar", "MySQL/Redis/OSS"],
            ["AI智能后端", "8081", "java -jar ai-backend.jar", "Redis/DashScope"],
            ["小程序后端", "8082", "java -jar mp-backend.jar", "MySQL"],
            ["预测服务", "8001", "uvicorn main:app --port 8001", "MySQL/TensorFlow"],
            ["数字人服务", "8083", "uvicorn main:app --port 8083", "TTS引擎"],
            ["Web前端", "3000", "npm run dev", "-"],
            ["小程序", "H5/微信", "npm run dev:h5", "-"],
        ],
        caption="表19  服务部署清单",
        col_widths=[2.5, 1.2, 4.5, 4],
        source="deploy/server-env.sh + 各服务README"
    )

    add_heading(doc, '8.1 环境要求', level=2)
    env_items = [
        "Java 17 (LTS) — Spring Boot 3.2 要求最低 Java 17",
        "Python 3.10+ — FastAPI + TensorFlow 运行环境",
        "Node.js 18+ — Vue 3 + Vite 构建环境",
        "MySQL 8.0+ — 关系型数据存储",
        "Redis 6.0+ — 缓存/限流/会话",
        "阿里云OSS — 文件/图片云存储（需配置AccessKey）",
        "通义千问DashScope API Key — AI功能必须",
    ]
    for item in env_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # ============================================================
    # 第九章 项目数据汇总（三线表）
    # ============================================================
    add_heading(doc, '第九章  项目数据汇总', level=1)

    add_heading(doc, '9.1 景区基础数据', level=2)
    make_three_line_table(doc,
        headers=["景区ID", "景区名称", "日均基础客流", "最大承载量", "海拔(m)", "拥挤度阈值"],
        rows=[
            ["1", "梅花山风景区", "3,000", "8,000", "2,400", ">6000极端/>4500高/>3000中"],
            ["2", "玉舍国家森林公园", "2,500", "6,000", "2,300", "同上"],
            ["3", "乌蒙大草原", "3,500", "10,000", "2,857", "同上"],
            ["4", "水城古镇", "2,000", "5,000", "1,800", "同上"],
            ["5", "明湖国家湿地公园", "1,800", "4,000", "1,750", "同上"],
        ],
        caption="表20  五大景区基础参数",
        col_widths=[1.2, 3.5, 2, 2, 1.5, 4],
        source="data_generator.py SCENIC_CONFIG"
    )

    add_heading(doc, '9.2 代码规模统计', level=2)
    make_three_line_table(doc,
        headers=["子系统", "语言", "Controller数", "文件数(约)", "核心内容"],
        rows=[
            ["主业务后端", "Java 17", "56", "301", "管理员/商户/用户/认证/景区/评价/订单/通知/系统"],
            ["AI智能后端", "Java 17", "~8", "46", "AI对话/RAG/行程规划/TTS代理/WebSocket"],
            ["小程序后端", "Java 17", "13", "67", "电商/支付/研学/反馈/轮播/商品管理"],
            ["API网关", "Java 17", "-", "20+", "路由/限流/认证/熔断/CORS/WebSocket代理"],
            ["Web前端", "Vue/TS", "-", "74", "55+个页面(含组件)+三角色+ECharts+3D地图"],
            ["微信小程序", "Vue 3", "-", "28+", "16个页面+12个API模块"],
            ["客流预测", "Python", "-", "8", "4种模型+数据治理+特征工程"],
            ["数字人", "JS/Python", "-", "15+", "Three.js+MapLibre+TTS"],
        ],
        caption="表21  项目代码规模统计",
        col_widths=[2.5, 1.5, 2, 1.8, 5],
    )

    add_heading(doc, '9.3 预测服务API', level=2)
    make_three_line_table(doc,
        headers=["方法", "接口路径", "参数", "功能"],
        rows=[
            ["GET", "/api/prediction/flow/{scenic_id}", "model,days,factors", "单景区N天客流预测"],
            ["GET", "/api/prediction/total", "days,model", "全区聚合客流预测"],
            ["GET", "/api/prediction/hourly/{scenic_id}", "date,model", "单景区小时级预测"],
            ["GET", "/api/prediction/hourly/total", "date,model", "全区小时级聚合"],
            ["POST", "/api/prediction/train/{scenic_id}", "model", "触发模型训练"],
            ["GET", "/api/prediction/models/info", "-", "所有模型信息"],
        ],
        caption="表22  Python预测服务API",
        col_widths=[1.5, 4.5, 3, 3.5],
        source="PythonPredictionService/main.py"
    )

    doc.add_page_break()

    # ============================================================
    # 第十章 PPT制作素材
    # ============================================================
    add_heading(doc, '第十章  PPT制作素材', level=1)

    add_para(doc, '以下内容可直接复制到PPT幻灯片中使用。', bold=True, indent=False)

    add_heading(doc, '10.1 一句话项目介绍', level=2)
    add_para(doc, '"智教黔行"是面向六盘水"中国凉都"的智慧文旅一体化平台，融合AI大模型、深度学习预测、3D数字人导游、红色研学教育、文创电商五大核心能力，覆盖管理者-商户-游客全角色、游前-游中-游后全流程。', indent=False)

    add_heading(doc, '10.2 核心数据（适合PPT数字页）', level=2)
    data_points = [
        "8 个微服务 —— 独立部署、弹性扩展(API网关:8888)",
        "92% 预测准确率 —— ARIMA+LSTM双流融合模型",
        "56 个后端Controller —— 覆盖全业务场景",
        "74+16 个前端Vue文件 —— Web三角色端+微信小程序端",
        "5 大景区 —— 梅花山/玉舍/乌蒙/水城/明湖",
        "3 种角色 —— 管理员/商户/游客",
        "6 维特征 —— 客流/节庆/周末/天气/温度/海拔",
        "4 种预测模型 —— ARIMA/LSTM/多变量LSTM/双流融合",
        "10 类本地服务 —— 交通/酒店/美食/非遗/民宿/门票/攻略/活动/优惠/研学",
        "30 天 —— 最长预测周期",
    ]
    for d in data_points:
        add_bullet(doc, d)

    add_heading(doc, '10.3 技术关键词（适合PPT标签云）', level=2)
    add_para(doc, 'Spring Boot 3.2 · Vue 3 · TypeScript · Uni-app · FastAPI · TensorFlow · ARIMA · LSTM · 双流融合 · 通义千问 · RAG · DashScope · Three.js · MapLibre · ECharts · Redis · MySQL · JWT · Spring Cloud Gateway · Resilience4j · 小波去噪 · 差分隐私 · 微信小程序 · TTS · 阿里云OSS · MyBatis-Plus · Pinia · TailwindCSS · Element Plus', indent=False)

    add_heading(doc, '10.4 创新点总结（适合PPT关键页）', level=2)
    innovations = [
        "【模型创新】ARIMA+LSTM双流融合预测，动态权重自适应，准确率92%",
        "【AI创新】RAG增强的大模型导游，知识库检索消除幻觉，专业回答",
        "【交互创新】3D数字人+地图融合+TTS口型同步，沉浸式导览",
        "【教育创新】红色研学游戏化闭环：路线→答题→积分→徽章→护照",
        "【架构创新】8微服务+API网关，JWT+限流+熔断，生产级高可用",
        "【隐私创新】差分隐私(ε=1.0)+小波去噪，数据治理与隐私保护并重",
    ]
    for inn in innovations:
        add_bullet(doc, inn)

    add_heading(doc, '10.5 用户价值（适合PPT对比页）', level=2)
    make_three_line_table(doc,
        headers=["维度", "传统景区", "智教黔行平台"],
        rows=[
            ["客流管理", "人工估算，被动应对", "AI预测(92%准确率)，主动调度"],
            ["游客导览", "纸质地图/人工导游", "3D数字导游+AI对话+TTS语音"],
            ["行程规划", "游客自行搜索攻略", "AI一键生成个性化行程"],
            ["文化教育", "静态展板/单向灌输", "游戏化答题+积分+徽章互动"],
            ["数据决策", "经验判断", "多维数据大屏+政策沙盘模拟"],
            ["商品销售", "线下门店", "线上文创商城+积分兑换"],
            ["多端体验", "仅线下", "Web管理端+微信小程序全覆盖"],
        ],
        caption="表23  传统景区 vs 智教黔行平台对比",
        col_widths=[2, 4.5, 6],
    )

    # ============================================================
    # 保存
    # ============================================================
    doc.save(OUTPUT_PATH)
    print(f"[OK] 完整技术文档已生成: {OUTPUT_PATH}")
    print(f"     含封面 + 目录 + 10章")
    print(f"     总计: 23 张三线表")


if __name__ == '__main__':
    print("=" * 60)
    print("  智教黔行 · 完整技术文档生成")
    print("=" * 60)
    main()
    print("=" * 60)
