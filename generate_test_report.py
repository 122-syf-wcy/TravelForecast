from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path("智教黔行_测试报告.docx")


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, value, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(value))
    set_font(run, bold=header, size=9.5, color="FFFFFF" if header else None)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, "4472C4")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, True)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph("")
    return table


def add_p(doc, text="", indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    set_font(run)
    return para


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "宋体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[style_name].font.name = "黑体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
    return doc


def main():
    doc = setup_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("智教黔行旅游预测与智能服务系统\n测试报告")
    set_font(r, "黑体", 24, True)
    doc.add_paragraph("")
    add_table(doc, ["项目", "内容"], [
        ["项目名称", "智教黔行旅游预测与智能服务系统"],
        ["测试对象", "Web 前端、业务后端、API 网关、AI 服务、预测服务、小程序后端、生产数据库"],
        ["测试方式", "自动化单元测试、编译构建验证、接口健康检查、数据库只读校验、生产服务冒烟测试"],
        ["测试时间", "2026-04-29"],
        ["数据来源", "本地 Maven/npm 构建输出、服务器 curl 健康检查、服务器 MySQL 只读查询、Surefire XML 汇总"],
    ], [4, 11])
    doc.add_page_break()

    doc.add_heading("目录", level=1)
    for item in ["1. 测试需求", "2. 测试计划", "3. 测试方案", "4. 真实测试执行记录", "5. 测试缺陷", "6. 测试总结"]:
        add_p(doc, item, False)
    doc.add_page_break()

    doc.add_heading("1. 测试需求", level=1)
    doc.add_heading("1.1 测试内容分析", level=2)
    add_p(doc, "本次测试围绕智教黔行系统核心交付能力展开，覆盖用户访问入口、景区数据浏览、AI 智能服务、客流预测服务、生产数据库、前端构建和后端可编译性。测试结果均来自真实命令输出，未将未执行或被取消的测试记录为通过。")
    add_table(doc, ["编号", "测试内容", "测试类型", "测试要点", "通过标准", "优先级"], [
        ["ST1", "统一网关与健康检查", "接口/可靠性", "网关进程、健康检查、下游路由", "健康接口 HTTP 200 且返回 UP", "高"],
        ["ST2", "景区浏览接口", "功能/接口", "景区列表接口可访问、返回真实景区数据", "HTTP 200 且返回业务 JSON", "高"],
        ["ST3", "AI 智能服务", "单元/接口", "JWT 拦截器、数字人客户端、预测客户端、RAG 检索", "JUnit 失败数为 0，服务健康为 UP", "高"],
        ["ST4", "预测服务", "接口/算法服务", "预测服务健康、模型加载状态", "健康接口返回 arima、lstm、dual_stream 可用", "高"],
        ["ST5", "数据库", "数据层/只读校验", "核心表存在、字段数量、表数据量", "可连接 travel_prediction 且核心表可查询", "高"],
        ["ST6", "Web 前端", "构建/兼容性", "Vite 构建产物、静态资源输出", "npm run build 成功生成 dist", "中"],
        ["ST7", "后端工程", "编译/集成准备", "业务后端和小程序后端 Java 17 编译", "mvn compile 退出码为 0", "高"],
        ["ST8", "安全鉴权", "安全/接口", "无 Token 访问受保护路径的响应", "返回 401 或按白名单放行", "高"],
    ])

    doc.add_heading("1.2 难点和重点", level=2)
    add_table(doc, ["编号", "难点项", "困难性说明", "处理方式"], [
        ["1", "多服务联调", "系统由网关、业务后端、AI 后端、预测服务、数字人服务、小程序后端组成，单点测试无法代表整体可用性。", "同时执行网关路由、服务直连和健康检查。"],
        ["2", "生产数据库真实性", "报告需要反映当前线上 travel_prediction 库状态。", "通过服务器 MySQL 只读查询 information_schema 和核心表行数。"],
        ["3", "测试环境差异", "本机 Maven 默认 Java 25，但项目要求 Java 17。", "Maven 测试与编译显式指定 JDK 17。"],
        ["4", "外部依赖限制", "部分完整上下文测试需要 MySQL、Redis、外部模型服务。", "离线单元测试与生产健康检查结合，未执行项单独列明。"],
    ])

    doc.add_heading("1.3 测试通过标准", level=2)
    add_p(doc, "自动化单元测试以失败数和错误数为 0 作为通过标准；编译构建以命令退出码 0 作为通过标准；接口测试以 HTTP 状态码、响应体结构和关键字段作为通过标准；数据库测试以只读查询成功、核心表元数据存在作为通过标准。已发现的问题不隐瞒，统一记录在测试缺陷章节。")

    doc.add_heading("2. 测试计划", level=1)
    doc.add_heading("2.1 测试环境", level=2)
    add_table(doc, ["类别", "环境/工具", "实测配置"], [
        ["本地操作系统", "macOS", "mac os x 26.3.1，aarch64"],
        ["本地 Java", "OpenJDK", "17.0.2；Maven 测试显式使用 JDK 17"],
        ["Maven", "Apache Maven", "3.9.12"],
        ["Node.js / npm", "前端构建环境", "Node v25.6.1，npm 11.9.0"],
        ["Python", "预测服务脚本环境", "Python 3.10.11"],
        ["服务器", "ECS 主机", "hostname: iZ2zed9w1py6ejrgxndtzwZ"],
        ["服务器 Java", "OpenJDK", "17.0.18 LTS"],
        ["数据库", "MySQL", "MySQL 8.0.44，travel_prediction，localhost:3306"],
        ["接口工具", "curl", "curl 7.61.1"],
    ])

    doc.add_heading("2.2 测试范围", level=2)
    add_table(doc, ["模块", "测试方式", "是否纳入", "说明"], [
        ["TravelForecastGateway", "JUnit + 生产健康检查", "是", "验证 fallback JSON、网关健康、预测/AI/数字人路由。"],
        ["TravelForecastingAIBackend", "JUnit + 生产健康检查", "是", "验证 JWT、预测客户端、数字人客户端、RAG 检索和 AI 健康接口。"],
        ["TravelForecastBackend", "编译 + 生产接口", "是", "本仓库未发现 src/test 用例，本次执行 Java 编译和景区接口冒烟测试。"],
        ["TravelForecastMiniProgramBackend", "编译 + 健康检查", "是", "该模块无 src/test 目录，本次执行 Java 编译和 actuator 健康检查。"],
        ["TravelForecastFrontend/web", "npm run build", "是", "验证 Web 端可构建并生成 dist/index.html。"],
        ["TravelForecast-PythonPredictionService", "生产健康检查", "部分", "本机缺少 pytest，未将 pytest 记录为通过；生产服务健康接口通过。"],
        ["travel_prediction 数据库", "MySQL 只读查询", "是", "统计 10 张核心表行数和字段数量。"],
    ])

    doc.add_heading("3. 测试方案", level=1)
    doc.add_heading("3.1 功能测试", level=2)
    add_table(doc, ["编号", "测试要点", "测试方法类型", "测试方法详述"], [
        ["FUNC-01", "景区列表浏览", "接口冒烟/等价类", "直连业务后端 /api/spots/list，确认返回 5A 景区等真实业务数据。"],
        ["FUNC-02", "AI 服务可用性", "接口冒烟", "访问 /ai-api/health，确认服务状态 UP 且依赖 prediction-service、digital-human 均 UP。"],
        ["FUNC-03", "预测服务可用性", "接口冒烟", "访问 /prediction-api/health，确认模型 arima、lstm、dual_stream 加载成功。"],
        ["FUNC-04", "数字人服务可用性", "接口冒烟", "访问 /digital-human-api/health，确认返回 healthy。"],
        ["FUNC-05", "小程序后端健康", "接口冒烟", "直连 8082 /actuator/health，确认小程序后端运行状态。"],
    ])

    doc.add_heading("3.2 接口测试", level=2)
    add_p(doc, "接口测试采用 curl 在服务器本机执行，避免公网 DNS、浏览器缓存和跨域策略干扰。每个接口记录 URL、HTTP 状态码、响应大小和响应体关键字段。")
    add_table(doc, ["接口", "实际结果", "结论"], [
        ["http://127.0.0.1:8888/health", "HTTP 200，返回 service=travel-gateway，status=UP", "通过"],
        ["http://127.0.0.1:8080/api/spots/list", "HTTP 200，响应 7997 字节，返回 scenic_spots 业务数据", "通过"],
        ["http://127.0.0.1:8888/api/spots/list", "HTTP 401，返回未提供认证Token", "未通过，见缺陷 D-001"],
        ["http://127.0.0.1:8888/prediction-api/health", "HTTP 200，返回模型 arima/lstm/dual_stream=true", "通过"],
        ["http://127.0.0.1:8888/ai-api/health", "HTTP 200，AI 服务 UP，依赖 digital-human 和 prediction-service 均 UP", "通过"],
        ["http://127.0.0.1:8888/digital-human-api/health", "HTTP 200，返回 status=healthy", "通过"],
        ["http://127.0.0.1:8082/actuator/health", "HTTP 200，返回 status=UP", "通过"],
        ["http://127.0.0.1:8888/miniprogram-api/actuator/health", "HTTP 401，返回未提供认证Token", "未通过，见缺陷 D-002"],
    ])

    doc.add_heading("3.3 自动化单元测试", level=2)
    add_table(doc, ["模块", "命令", "测试数", "失败", "错误", "跳过", "结论"], [
        ["TravelForecastGateway", "JAVA_HOME=JDK17 mvn -q test", "3", "0", "0", "0", "通过"],
        ["TravelForecastingAIBackend", "JAVA_HOME=JDK17 mvn -q test", "22", "0", "0", "1", "通过；1 个完整 Spring 上下文测试因需真实 MySQL/Redis 被跳过"],
    ])
    add_table(doc, ["测试类/用例方向", "覆盖内容", "结果"], [
        ["FallbackControllerTest", "网关业务、AI、预测服务 fallback 响应体和 503 状态码", "通过"],
        ["JwtInterceptorTest", "OPTIONS 放行、公开健康接口放行、有效 Token 解析、无效 Token 返回 401", "通过"],
        ["PredictionClientTest", "预测服务 health、景区流量预测、总量预测、小时预测、非法参数短路", "通过"],
        ["DigitalHumanClientTest", "数字人 response/reply 字段兼容、禁用状态不发请求", "通过"],
        ["Bm25RerankerTest / ScoreTest", "中文 bigram、英文 token、排序、空查询、评分返回", "通过"],
        ["HybridRetrieverTest", "向量检索、BM25 回退、去重和 topK 限制", "通过"],
    ])

    doc.add_heading("3.4 编译与构建测试", level=2)
    add_table(doc, ["模块", "命令", "实际结果", "结论"], [
        ["TravelForecastBackend", "JAVA_HOME=JDK17 mvn -q -DskipTests compile", "退出码 0，无编译错误", "通过"],
        ["TravelForecastMiniProgramBackend", "JAVA_HOME=JDK17 mvn -q -DskipTests compile", "退出码 0，无编译错误", "通过"],
        ["TravelForecastFrontend/web", "npm run build", "退出码 0，生成 dist/index.html，Vite 构建耗时约 6.79s", "通过，但存在大 chunk 警告"],
        ["TravelForecast-PythonPredictionService", "python3 -m pytest -q tests/...", "本机 Python 环境缺少 pytest，命令返回 No module named pytest", "未执行通过，不计入通过项"],
    ])

    doc.add_heading("3.5 数据库测试", level=2)
    add_p(doc, "数据库测试在服务器上加载 /opt/travel-env.sh 后连接 MySQL localhost:3306 的 travel_prediction 库，仅执行 information_schema 和表统计查询，不进行 INSERT、UPDATE、DELETE。")
    add_table(doc, ["表名", "实测行数", "说明"], [
        ["ai_conversations", "57", "AI 会话主表存在历史会话数据"],
        ["ai_messages", "114", "AI 消息表存在对话明细数据"],
        ["flow_records", "5369", "景区客流历史记录量较充分"],
        ["itineraries", "0", "当前生产库暂无行程规划记录"],
        ["mp_orders", "1", "小程序订单表存在少量数据"],
        ["predictions", "0", "当前生产库暂无持久化预测记录"],
        ["scenic_realtime_data", "0", "当前生产库暂无实时数据落库记录"],
        ["scenic_spots", "5", "景区基础数据存在"],
        ["users", "19", "用户数据存在"],
        ["visitor_predictions", "14", "游客预测记录存在"],
    ])
    add_p(doc, "10 张核心表共检测到 131 个字段，说明核心数据库结构存在且可被当前服务账户读取。部分业务表行数为 0 属于当前业务数据状态，不直接等同于功能失败。")

    doc.add_heading("3.6 压力、性能测试", level=2)
    add_p(doc, "本次未进行大规模并发压测，主要完成构建体积和线上服务健康冒烟验证。前端构建提示多个 chunk 超过 500KB，其中 TripPlanning、index 等资源体积较大，建议后续通过动态导入和 manualChunks 拆分优化首屏加载性能。")
    add_table(doc, ["检查项", "实测结果", "评价"], [
        ["Web 构建耗时", "约 6.79s", "正常"],
        ["静态资源体积", "存在 500KB 以上 chunk，最大约 1123.51KB", "需优化"],
        ["预测服务运行时长", "health 返回 uptimeSeconds 约 371721 秒", "稳定运行"],
        ["网关健康响应", "HTTP 200，JSON 体 86 字节", "正常"],
    ])

    doc.add_heading("4. 真实测试执行记录", level=1)
    add_table(doc, ["编号", "执行时间", "执行位置", "命令/操作", "结果"], [
        ["R-001", "2026-04-29 21:59", "本地 TravelForecastGateway", "JAVA_HOME=JDK17 mvn -q test", "3 个 JUnit 用例通过，失败 0，错误 0"],
        ["R-002", "2026-04-29 21:59", "本地 TravelForecastingAIBackend", "JAVA_HOME=JDK17 mvn -q test", "22 个测试记录，失败 0，错误 0，跳过 1"],
        ["R-003", "2026-04-29 22:10", "服务器 zhanghaodong", "curl 网关、AI、预测、数字人、小程序健康接口", "多数健康接口通过，发现网关白名单相关问题"],
        ["R-004", "2026-04-29 22:10", "服务器 MySQL", "information_schema + 核心表行数只读查询", "10 张核心表可查询，共 131 个字段"],
        ["R-005", "2026-04-29 22:11", "本地 TravelForecastFrontend/web", "npm run build", "构建成功，生成 dist/index.html，存在 chunk 体积警告"],
        ["R-006", "2026-04-29 22:12", "本地 TravelForecastBackend", "JAVA_HOME=JDK17 mvn -q -DskipTests compile", "编译通过"],
        ["R-007", "2026-04-29 22:12", "本地 TravelForecastMiniProgramBackend", "JAVA_HOME=JDK17 mvn -q -DskipTests compile", "编译通过"],
    ])

    doc.add_heading("5. 测试缺陷", level=1)
    add_table(doc, ["缺陷编号", "模块", "严重级别", "现象", "影响", "建议处理"], [
        ["D-001", "API 网关 / 景区接口", "高", "直连业务后端 http://127.0.0.1:8080/api/spots/list 返回 200，但经网关 http://127.0.0.1:8888/api/spots/list 返回 401。", "游客公共景区列表可能无法通过统一网关访问，影响 Web 或小程序公共浏览链路。", "将 /api/spots/** 加入网关白名单，或统一控制器路径与网关白名单命名。修复后复测网关路径应返回 200。"],
        ["D-002", "API 网关 / 小程序健康检查", "中", "直连 http://127.0.0.1:8082/actuator/health 返回 200，但经网关 /miniprogram-api/actuator/health 返回 401。", "如果运维监控通过网关探活，会误判小程序后端不可用。", "按实际网关路径补充 /miniprogram-api/actuator/** 白名单，或改用直连服务健康检查。"],
        ["D-003", "Web 前端构建性能", "中", "npm run build 成功，但提示多个 chunk 大于 500KB，最大约 1123.51KB。", "首屏加载和弱网访问体验可能受影响。", "使用动态 import、路由级懒加载和 rollupOptions.output.manualChunks 拆分 ECharts、Three、地图、AI 相关资源。"],
        ["D-004", "测试环境依赖", "低", "本机 Python 环境缺少 pytest，预测服务 tests 目录下 pytest 用例未能在本机执行。", "影响本地自动化测试覆盖闭环，但生产预测服务健康接口已通过。", "在测试环境安装 requirements.txt 和 pytest，或将 pytest 加入依赖并接入 CI。"],
        ["D-005", "自动化测试覆盖", "中", "TravelForecastBackend 和 TravelForecastMiniProgramBackend 未发现 src/test 自动化测试用例。", "业务后端核心登录、景区、订单、权限等逻辑缺少回归测试保护。", "补充 Controller/Service 层单元测试和接口集成测试，重点覆盖认证、权限、订单、景区公共接口。"],
    ])

    doc.add_heading("6. 测试总结", level=1)
    doc.add_heading("6.1 软件能力", level=2)
    add_p(doc, "从本次真实测试结果看，系统核心服务整体处于可运行状态：网关、AI 服务、预测服务、数字人服务、小程序后端和业务后端均存在运行或编译通过证据；AI 后端和网关已有可执行 JUnit 用例且失败数为 0；生产数据库 travel_prediction 可连接，10 张核心表结构和数据状态可查询；Web 前端可以完成生产构建并生成 dist 产物。")
    doc.add_heading("6.2 缺陷分析", level=2)
    add_p(doc, "主要问题集中在网关鉴权白名单与实际服务路径不一致、前端资源体积偏大、部分模块自动化测试覆盖不足三个方面。其中 D-001 会直接影响公共景区接口经网关访问，建议优先修复；D-002 影响监控探活准确性；D-003 属于性能优化项；D-004 和 D-005 属于测试体系建设问题。")
    doc.add_heading("6.3 建议", level=2)
    add_table(doc, ["优先级", "建议项", "预期效果"], [
        ["高", "修复网关白名单，补充 /api/spots/** 和 /miniprogram-api/actuator/** 相关路径复测。", "恢复公共景区浏览和网关健康探活的一致性。"],
        ["高", "为业务后端补充登录、景区列表、景区详情、订单、权限拦截等 JUnit/MockMvc 测试。", "降低核心业务回归风险。"],
        ["中", "将 Maven、npm build、pytest、接口健康检查纳入 CI，测试失败阻断提交或部署。", "形成稳定自动化测试闭环。"],
        ["中", "前端按路由和大型依赖拆分 chunk，并设置构建体积阈值。", "提升首屏加载速度和弱网体验。"],
        ["低", "为测试报告保留 surefire XML、curl 输出和数据库只读查询记录。", "增强提交材料的可追溯性和可复核性。"],
    ])

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "智教黔行旅游预测与智能服务系统测试报告"

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
