from pathlib import Path
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
META_PATH = ROOT / 'travel_core_db_design_metadata.json'
OUT_PATH = ROOT / '智教黔行_数据库设计报告.docx'

TABLE_DESCRIPTIONS = {
    'users': '系统用户主表，存储游客、商家、管理员等账号基础信息，是认证授权、订单、反馈、收藏、AI对话等业务的用户基础。',
    'scenic_spots': '景区基础信息表，存储景区名称、位置、介绍、开放状态、承载量等信息，是文旅资源展示、地图导览和客流预测的核心实体。',
    'scenic_realtime_data': '景区实时状态数据表，用于记录景区当前游客量、舒适度、告警状态等动态信息。',
    'flow_records': '客流历史记录表，用于保存按时间粒度采集的景区客流数据，为统计分析和预测模型提供历史数据。',
    'visitor_predictions': '游客数量预测结果表，保存指定景区、指定日期的预测游客量、置信度、天气因素和节假日因素。',
    'predictions': '预测记录表，用于保存预测任务或预测结果记录，支撑预测服务的结果追踪和历史查询。',
    'ai_conversations': 'AI会话表，记录智能问答、AI导游、研学助手等场景下的会话基础信息。',
    'ai_messages': 'AI消息表，记录会话中的用户消息、助手消息和系统消息，与AI会话表构成一对多关系。',
    'itineraries': '用户行程表，记录用户创建的旅游计划、路线安排、状态和预算等信息。',
    'mp_orders': '小程序订单表，记录文创商城或小程序端交易订单，支撑下单、支付、发货和售后流程。',
}

MODULE_MAP = {
    'users': '用户与权限管理',
    'scenic_spots': '景区资源管理',
    'scenic_realtime_data': '实时客流监测',
    'flow_records': '客流数据采集',
    'visitor_predictions': '游客量预测分析',
    'predictions': '预测服务管理',
    'ai_conversations': 'AI智能交互',
    'ai_messages': 'AI智能交互',
    'itineraries': '行程规划',
    'mp_orders': '小程序商城订单',
}

RELATIONS = [
    'users 与 ai_conversations、ai_messages、itineraries、mp_orders 等表通过用户编号建立业务关联。',
    'scenic_spots 是景区相关数据的主实体，scenic_realtime_data、flow_records、visitor_predictions、predictions 等表围绕景区编号记录实时、历史和预测数据。',
    'ai_conversations 与 ai_messages 构成一对多关系，一个会话包含多条消息记录。',
    'itineraries 保存用户行程主信息，可与行程景点明细表等扩展表形成主从结构。',
    'mp_orders 记录小程序订单主信息，可与订单明细、用户地址、商品等小程序业务表关联。',
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else '')
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_font(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[style_name]
        style.font.name = '黑体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, bold=True, size=9)
        set_cell_shading(table.cell(i, 0), 'D9EAF7')
        set_cell_text(table.cell(i, 1), v, size=9)
    return table


def add_simple_table(doc, headers, rows, font_size=7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size)
        set_cell_shading(table.rows[0].cells[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val if val not in [None, ''] else '-', size=font_size)
    return table


def key_desc(col):
    parts = []
    if col['key'] == 'PRI':
        parts.append('主键')
    elif col['key'] == 'UNI':
        parts.append('唯一索引')
    elif col['key'] == 'MUL':
        parts.append('普通索引/外键关联')
    if col['extra']:
        parts.append(col['extra'])
    return '；'.join(parts) if parts else '-'


def index_summary(table_meta):
    groups = {}
    for r in table_meta.get('indexes', []):
        if len(r) < 5:
            continue
        name = r[2]
        non_unique = r[1]
        col = r[4]
        groups.setdefault(name, {'unique': non_unique == '0', 'cols': []})['cols'].append(col)
    return [[name, '唯一/主键' if info['unique'] else '普通索引', ', '.join(info['cols'])] for name, info in groups.items()]


def build_report():
    meta = json.loads(META_PATH.read_text(encoding='utf-8'))
    doc = Document()
    set_doc_font(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('数据库设计报告')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智教黔行 - 六盘水山地智慧文旅一体化研学平台')
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    add_kv_table(doc, [
        ('文档名称', '智教黔行数据库设计报告'),
        ('数据库名称', meta['database']),
        ('数据库系统', 'MySQL 8.0.44'),
        ('服务器环境', '39.97.232.141 内部 MySQL localhost:3306'),
        ('文档版本', 'V1.0'),
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M')),
    ])
    doc.add_page_break()

    add_heading(doc, '目录', 1)
    for item in ['0. 文档介绍', '1. 数据库环境说明', '2. 数据库命名规则', '3. 概念结构设计', '4. 逻辑结构设计', '5. 物理结构设计', '6. 安全性设计', '7. 数据库优化', '8. 数据库管理与维护说明']:
        doc.add_paragraph(item)
    doc.add_page_break()

    add_heading(doc, '0. 文档介绍', 1)
    add_heading(doc, '0.1 文档目的', 2)
    add_para(doc, '本文档作为“智教黔行 - 六盘水山地智慧文旅一体化研学平台”的数据库设计说明，主要用于规范数据库对象定义、说明核心业务表结构、梳理数据关系，并为后续系统维护、功能扩展、答辩展示和材料提交提供依据。')
    add_heading(doc, '0.2 文档范围', 2)
    add_para(doc, '本文档覆盖数据库环境、命名规则、概念结构、逻辑结构、物理结构、安全性设计、数据库优化及维护说明。表结构部分选取系统最具代表性的十张核心表进行展开，完整数据库结构可参见服务器导出的 travel_prediction_schema_20260429_211418.sql。')
    add_heading(doc, '0.3 读者对象', 2)
    add_para(doc, '本文档面向项目组成员、评审专家、运维人员、开发人员及后续维护人员。')
    add_heading(doc, '0.4 参考依据', 2)
    for ref in ['服务器 MySQL 数据库 travel_prediction 实际表结构', '项目后端配置文件与数据库迁移脚本', '智教黔行系统设计与开发文档', 'DataGrip 数据库结构视图与服务器 mysqldump 导出文件']:
        doc.add_paragraph(ref)
    add_heading(doc, '0.5 术语与缩写解释', 2)
    add_simple_table(doc, ['术语', '解释'], [
        ['MySQL', '系统使用的关系型数据库管理系统。'],
        ['InnoDB', 'MySQL 默认事务型存储引擎，支持事务、行级锁和外键。'],
        ['ER 模型', '实体关系模型，用于描述实体、属性以及实体之间的联系。'],
        ['主键', '唯一标识表中一条记录的字段或字段组合。'],
        ['外键', '用于描述表间引用关系的约束。'],
        ['索引', '用于提高查询效率的数据结构。'],
    ], font_size=9)

    add_heading(doc, '1. 数据库环境说明', 1)
    add_kv_table(doc, [
        ('数据库系统', 'MySQL 8.0.44'),
        ('数据库名称', 'travel_prediction'),
        ('字符集/排序规则', 'utf8mb4 / utf8mb4_unicode_ci、utf8mb4_0900_ai_ci'),
        ('存储引擎', 'InnoDB'),
        ('部署位置', '生产服务器内部 MySQL localhost:3306'),
        ('管理工具', 'DataGrip、MySQL CLI、mysqldump'),
        ('结构来源', 'information_schema 实时查询与 mysqldump 结构导出'),
    ])
    add_para(doc, '生产环境中，旅游业务数据库运行在服务器本机 MySQL 3306 端口，数据库名为 travel_prediction。服务器上另有 Docker MySQL 3307 实例，用于其他中间件或课程项目，不属于本系统业务数据库。')

    add_heading(doc, '2. 数据库命名规则', 1)
    add_heading(doc, '2.1 命名原则', 2)
    for rule in ['数据库对象名称采用英文单词或英文缩写，多个单词之间使用下划线连接。', '表名以业务含义命名，例如 users、scenic_spots、visitor_predictions。', '字段名采用小写英文与下划线组合，例如 user_id、created_at、prediction_date。', '主键字段通常采用 id 或业务实体_id，时间字段统一使用 created_at、updated_at 等形式。', '索引名称采用 idx_、uk_ 等前缀表达索引含义。']:
        doc.add_paragraph(rule)
    add_heading(doc, '2.2 数据库对象命名规则', 2)
    add_simple_table(doc, ['对象类型', '命名规则', '示例'], [
        ['数据表', '英文单词或业务缩写，多个单词用下划线连接', 'scenic_spots、mp_orders'],
        ['字段', '小写英文单词，下划线分隔', 'prediction_date、user_id'],
        ['主键', '一般使用 id 或业务主键字段', 'id、user_id'],
        ['索引', 'idx_ 或 uk_ 前缀加业务字段', 'idx_scenic_date、uk_conversation_uuid'],
        ['时间字段', '统一使用 created_at、updated_at、deleted_at', 'created_at'],
    ], font_size=9)

    add_heading(doc, '3. 概念结构设计', 1)
    add_para(doc, '系统围绕“用户 - 景区 - 客流 - 预测 - AI服务 - 行程 - 小程序订单”建立数据模型。用户作为系统访问主体，景区作为文旅资源主体，客流记录和预测结果围绕景区展开，AI会话与消息支撑智能问答和导游服务，行程与订单支撑用户侧业务闭环。')
    for rel in RELATIONS:
        doc.add_paragraph(rel)

    add_heading(doc, '4. 逻辑结构设计', 1)
    add_heading(doc, '4.0 表汇总', 2)
    summary_rows = [[t['name'], MODULE_MAP.get(t['name'], '-'), t['comment'] or TABLE_DESCRIPTIONS.get(t['name'], ''), str(t['row_count'])] for t in meta['tables']]
    add_simple_table(doc, ['表名', '所属模块', '功能说明', '当前记录数'], summary_rows, font_size=8)

    for idx, t in enumerate(meta['tables'], 1):
        add_heading(doc, f'4.{idx} 表 {t["name"]}', 2)
        add_kv_table(doc, [
            ('表名', t['name']),
            ('功能说明', TABLE_DESCRIPTIONS.get(t['name'], t.get('comment') or '系统业务表')),
            ('存储引擎', t.get('engine') or 'InnoDB'),
            ('排序规则', t.get('collation') or '-'),
            ('当前记录数', str(t.get('row_count', 0))),
        ])
        rows = []
        for no, c in enumerate(t['columns'], 1):
            rows.append([no, c['name'], c['comment'] or '-', c['type'], '空' if c['nullable'] == 'YES' else '非空', key_desc(c), c['default'] if c['default'] is not None else '-'])
        add_simple_table(doc, ['序号', '列名', '描述', '数据类型', '空/非空', '约束条件', '默认值'], rows, font_size=6)
        idx_rows = index_summary(t)
        if idx_rows:
            doc.add_paragraph('索引说明：')
            add_simple_table(doc, ['索引名称', '索引类型', '字段'], idx_rows, font_size=8)

    add_heading(doc, '5. 物理结构设计', 1)
    add_para(doc, '系统数据库采用 MySQL 8.0.44，核心业务表主要使用 InnoDB 存储引擎。InnoDB 支持事务、行级锁和崩溃恢复，适用于订单、用户、预测记录等对一致性要求较高的业务场景。')
    add_heading(doc, '5.1 数据库存储设计', 2)
    for item in ['字符集采用 utf8mb4，支持中文、表情符号及多语言内容存储。', '主键多采用 bigint 自增或业务唯一标识，便于扩展和跨模块关联。', '时间字段统一记录创建时间、更新时间或业务时间，便于审计和统计。', 'JSON 字段用于存储图片列表、标签等结构灵活的数据。']:
        doc.add_paragraph(item)
    add_heading(doc, '5.2 核心业务场景对应表', 2)
    add_simple_table(doc, ['业务场景', '核心表', '说明'], [
        ['用户登录与权限', 'users', '保存系统用户与角色身份基础信息。'],
        ['景区资源展示', 'scenic_spots', '保存景区名称、地址、介绍、开放状态等。'],
        ['实时客流监测', 'scenic_realtime_data、flow_records', '保存实时与历史客流数据。'],
        ['游客量预测', 'visitor_predictions、predictions', '保存预测任务和预测结果。'],
        ['AI问答与导游', 'ai_conversations、ai_messages', '保存会话和消息上下文。'],
        ['行程规划', 'itineraries', '保存用户行程计划。'],
        ['小程序商城', 'mp_orders', '保存小程序订单信息。'],
    ], font_size=9)

    add_heading(doc, '6. 安全性设计', 1)
    add_heading(doc, '6.1 防止用户直接操作数据库的方法', 2)
    add_para(doc, '系统通过后端服务接口访问数据库，普通用户、商家用户和管理员用户均不直接连接数据库。生产数据库运行在服务器内部网络环境，外部访问应通过 SSH 隧道、堡垒机或受控运维通道进行。')
    add_heading(doc, '6.2 用户账号密码保护', 2)
    add_para(doc, '用户密码不应以明文形式存储，系统后端应使用安全哈希算法和统一认证逻辑进行账号校验。数据库连接密码通过服务器环境变量文件维护，避免硬编码到前端代码或公开材料中。')
    add_heading(doc, '6.3 角色与权限', 2)
    add_para(doc, '系统根据游客、商家、管理员等角色进行权限区分。数据库层面通过用户表、角色权限表及后端拦截器共同实现访问控制，业务接口根据登录态和角色控制可访问的数据范围。')

    add_heading(doc, '7. 数据库优化', 1)
    add_heading(doc, '7.1 索引优化', 2)
    add_para(doc, '核心表已围绕高频查询字段建立索引，例如景区编号、预测日期、用户编号、会话编号、订单编号等字段。对于客流预测和实时监测场景，按 scenic_id、prediction_date、created_at 等字段建立索引可提升查询效率。')
    add_heading(doc, '7.2 查询与存储优化', 2)
    for item in ['对历史客流等增长较快的数据表，应按时间范围进行查询，避免全表扫描。', '对订单、会话、消息等业务表，应优先使用主键、用户编号、业务编号进行定位。', '对景区图片、标签等非强结构化字段，可使用 JSON 字段降低扩展成本。', '对统计类数据可通过汇总表或缓存降低实时计算压力。']:
        doc.add_paragraph(item)
    add_heading(doc, '7.3 容灾优化', 2)
    add_para(doc, '生产环境应定期执行 mysqldump 结构与数据备份，并结合服务器快照、日志归档和权限分离降低误操作风险。对于核心业务库，应保留可恢复的备份版本。')

    add_heading(doc, '8. 数据库管理与维护说明', 1)
    add_para(doc, '数据库维护工作包括表结构版本管理、定期备份、慢查询分析、索引检查、敏感字段保护和账号权限审计。新增表结构或字段时，应同步更新数据库设计报告、迁移脚本和后端实体映射，确保文档与系统实现保持一致。')
    add_simple_table(doc, ['维护事项', '维护说明'], [
        ['结构备份', '使用 mysqldump --no-data 导出数据库结构，保存版本记录。'],
        ['数据备份', '对生产数据进行定期全量或增量备份。'],
        ['权限管理', '控制数据库 root 账号使用范围，应用服务使用专用账号更安全。'],
        ['性能维护', '定期检查慢查询、索引命中率和大表增长情况。'],
        ['文档同步', '数据库表结构调整后同步更新设计文档和提交材料。'],
    ], font_size=9)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == '__main__':
    build_report()
