from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path("智教黔行_AI工具使用说明（2026年版）.docx")
WORK_NAME = "智教黔行—六盘水山地智慧文旅一体化研学平台"


def set_run(run, size=10.5, font="宋体", bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def p(doc, text="", size=10.5, bold=False, align=None, font="宋体"):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, font=font)
    return para


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, header=False, center=False, size=8):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header or center) else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(str(text))
    set_run(run, size=size, bold=header, color="FFFFFF" if header else None)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, "4472C4")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, True, 8)
        table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, False, i == 0, 7.5)
            cells[i].width = Cm(widths[i])
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p(doc, "中国大学生计算机设计大赛", 18, True, WD_ALIGN_PARAGRAPH.CENTER, "黑体")
    p(doc, "AI工具使用说明（2026年版）", 16, True, WD_ALIGN_PARAGRAPH.CENTER, "黑体")
    p(doc, f"作品编号：待填写        作品名称：{WORK_NAME}", 11, False, WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, f"填报日期：{date.today()}        说明：本表按真实使用情况填写，未使用或未采纳的内容不计入作品原创贡献。", 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "参赛合规说明：本表仅填报大赛规定的国产AI工具或项目自研AI服务；未将非规定工具计入正式AI工具使用清单。", 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)

    headers = [
        "序号",
        "AI工具的名称、版本、访问方式（网页、API或客户端），使用时间",
        "使用AI工具的环节与目的（立项构思、文献综述、语言润色、内容生成、图表优化、代码编程、数据分析等）",
        "关键提示词",
        "AI回复的关键内容（在此简要说明，并在附录中给出佐证）",
        "AI回复的人工修改说明",
        "采纳比例与说明",
    ]

    rows = [
        [
            "1",
            "阿里通义系列：DashScope 通义千问 qwen-plus/qwen-turbo，API访问，2026年4月系统开发与联调阶段",
            "系统内置AI导游、智能问答、AI行程规划和研学方案生成。",
            "你是六盘水智慧旅游助手，请结合用户需求回答景区、美食、交通、安全、研学和客流相关问题；行程规划需输出结构化JSON，字段包含title、summary、budgetBreakdown、itinerary、tips。",
            "生成景区咨询回答、行程规划建议、研学活动方案、预算拆分和注意事项。",
            "人工设计系统提示词、预算约束、JSON结构、超时控制、缓存策略和失败兜底；AI输出只作为系统实时交互内容，前端展示前由程序进行结构化解析和格式化。",
            "约50%—70%。系统实时回答根据用户问题动态生成，核心业务规则、数据来源、接口流程和展示逻辑由人工设计。",
        ],
        [
            "2",
            "阿里通义系列：DashScope text-embedding-v2，API与后端服务组合使用，2026年4月",
            "知识库检索增强：提升AI问答对六盘水景区、文化、交通、研学资料的引用准确性。",
            "根据用户问题检索知识库，优先返回与景区、文化、交通、政策或研学主题相关的材料；候选资料需支持BM25召回与向量重排。",
            "返回文本向量、候选知识片段和重排后的相关资料，为AI回答提供参考上下文。",
            "人工实现BM25召回、向量重排、启用/禁用开关、无API Key时降级为纯BM25，并维护知识库数据；不直接把检索结果当最终结论。",
            "约40%。向量化与排序结果作为辅助依据，知识库内容建设、检索策略和最终回答约束由人工完成。",
        ],
        [
            "3",
            "阿里通义系列：DashScope CosyVoice TTS，API/后端服务访问，2026年4月",
            "数字人语音播报：将AI导游回答转换为语音，服务Web端或小程序端数字人交互。",
            "请将以下旅游讲解文本转换为普通话语音，语气自然、适合游客导览，保留景区名称和安全提示。",
            "返回可播放的语音音频数据，用于数字人回答播报。",
            "人工限制输入长度、处理失败降级、封装TTS接口、控制前端播放状态，并避免把语音结果作为新增事实来源。",
            "约30%。采纳语音合成本身，文本内容和播报触发逻辑由系统控制。",
        ],
        [
            "4",
            "DeepSeek，网页访问，2026年4月代码设计与问题拆解阶段",
            "代码编程与系统设计：辅助讨论Spring Boot多服务分层、FastAPI预测服务接口、Vue/UniApp前端状态设计和异常处理思路。",
            "请基于“Spring Cloud Gateway + Spring Boot + FastAPI + Vue3 + UniApp”的智慧文旅系统，设计客流预测调用链路：网关路由、Java客户端、FastAPI接口、异常降级、前端图表展示，并给出伪代码结构。",
            "给出接口分层、DTO设计、服务调用链、错误处理、缓存和前端展示建议。",
            "人工根据现有项目结构改写为实际Controller、Service、Client、API工具函数和页面逻辑；删除无法落地的泛化建议，不直接复制整段代码。",
            "约35%。主要采纳设计思路和伪代码结构，实际代码由人工按项目规范重写、调试和测试。",
        ],
        [
            "5",
            "DeepSeek，网页访问，2026年4月数据库与测试设计阶段",
            "数据库设计、接口测试和边界场景分析：辅助梳理用户、景区、AI会话、预测、订单、研学等数据域和测试用例。",
            "请为智慧文旅系统设计一组真实测试用例，覆盖登录鉴权、景区内容、AI问答、客流预测、数字人语音、小程序订单、紧急救援；每个用例包含前置条件、步骤、预期结果和人工验收点。",
            "给出按模块划分的测试用例清单、边界条件、异常场景和回归检查点。",
            "人工只保留已执行或能由代码/日志支撑的测试项；删除固定性能指标、未经验证的准确率和无法复现的测试结论。",
            "约45%。用例框架和边界提醒部分采纳，最终测试结果、截图和结论由人工执行后填写。",
        ],
        [
            "6",
            "Kimi，网页访问，2026年4月长文档整理与语言润色阶段",
            "设计文档、详细设计报告、测试报告和AI工具使用说明的长文本梳理与语言润色。",
            "请按照中国大学生计算机设计大赛文档口径，将以下系统说明改写为正式参赛材料语言，要求突出人的主导地位，不夸大系统能力，不出现未实现功能。",
            "给出更规范的章节标题、表述顺序、摘要性说明和风险提示语句。",
            "人工逐项核对源代码、数据库和运行结果；删除不真实表述，保留真实存在的AI导游、RAG、预测、数字人语音和政策模拟等内容。",
            "约50%。主要采纳语言组织和格式建议，事实内容、技术边界和最终版本由人工确认。",
        ],
        [
            "7",
            "阿里通义系列：通义千问/通义灵码或同系代码辅助能力，网页或IDE插件访问，2026年4月前端交互设计阶段",
            "前端代码设计：辅助生成Vue3/UniApp页面状态管理、接口调用、错误提示和加载态设计样例。",
            "请给出一个Vue3客流预测页面的代码设计示例，包含景区选择、预测天数、模型类型、loading状态、错误提示、ECharts数据映射和空数据兜底，不要写死真实接口密钥。",
            "给出组件状态划分、API调用流程、图表数据映射和用户提示设计。",
            "人工按项目已有路由、请求封装、组件风格和接口字段重写；只采纳局部交互逻辑和错误处理思路。",
            "约30%。示例代码不直接进入最终作品，最终代码经人工重构、联调和样式适配。",
        ],
        [
            "8",
            "自研AI客流预测服务（ARIMA、LSTM、TensorFlow/Keras、FastAPI），项目自有服务，2026年4月",
            "数据分析与客流预测：根据历史客流、日期、节假日、天气、周末等因素生成景区客流预测。",
            "无自然语言提示词，采用程序参数：景区ID、预测天数、模型类型、日期等；示例：scenicId=1、days=7、model=dual_stream。",
            "输出未来客流人数、小时级分布、拥挤度、峰值时段和模型信息。",
            "人工完成数据字段设计、模型调用、动态权重、异常兜底、接口封装和前端可视化；预测结果作为辅助决策参考，不作为绝对承诺。",
            "约45%。采纳模型计算结果，数据治理、接口设计、可视化解释和风险提示由人工完成。",
        ],
        ["9", "未使用", "未使用", "未使用", "未使用", "未使用", "0%。本序号保留为空，后续如补充AI工具再填写。"],
    ]

    add_table(doc, headers, rows, [0.8, 4.0, 4.2, 4.5, 4.5, 4.8, 3.2])

    doc.add_page_break()
    p(doc, "附录1：作品文件夹示例", 13, True, None, "黑体")
    folder_lines = [
        "作品编号待填写-参赛总文件夹",
        "├── 作品编号待填写-01作品与答辩材料",
        "├── 作品编号待填写-02素材与源码",
        "├── 作品编号待填写-03设计与开发文档",
        "└── 作品编号待填写-04作品演示视频",
    ]
    for line in folder_lines:
        p(doc, line, 10, False, None, "宋体")

    p(doc, "附录2：佐证材料说明", 13, True, None, "黑体")
    evidence = [
        "序号1的佐证材料：AI后端聊天、行程规划、研学方案相关接口运行截图或代码片段；注意提交材料中不得暴露API Key。",
        "序号2的佐证材料：知识库表、知识检索接口、RAG问答页面或调试截图；可展示BM25/向量重排的设计说明。",
        "序号3的佐证材料：数字人页面、TTS接口调用截图、语音播放效果截图或演示视频片段；不需要提交外部平台密钥。",
        "序号4的佐证材料：代码设计提示词与人工重构后的接口调用链说明，可附关键代码截图但不暴露密钥。",
        "序号5的佐证材料：测试用例设计提示词、测试报告、构建/接口/数据库校验的终端输出截图。",
        "序号6的佐证材料：详细设计报告、测试报告、技术创新核对版文档的修订记录和人工核验说明。",
        "序号7的佐证材料：前端页面代码设计提示词、页面截图、重构后的Vue/UniApp代码片段。",
        "序号8的佐证材料：预测服务接口返回、客流预测页面截图、测试报告中关于预测服务健康检查和接口校验的记录。",
        "序号9的佐证材料：未使用。",
    ]
    for line in evidence:
        p(doc, line, 10)

    p(doc, "真实性声明：本说明仅披露项目开发、系统运行和材料整理过程中实际使用或集成的AI工具。AI输出均经过人工核对、改写、测试或系统约束处理；作品核心需求分析、系统设计、代码实现、数据组织、功能验收和最终提交判断由团队完成。", 10, False)

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
