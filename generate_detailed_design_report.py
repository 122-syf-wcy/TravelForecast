from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path("智教黔行_详细设计报告.docx")

PROJECT_NAME = "智教黔行旅游预测与智能服务系统"
TEAM_NAME = "智教黔行项目组"


def set_run(run, font="宋体", size=10.5, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", first_line=True, align=None, size=10.5, bold=False):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level in (1, 2, 3) else "Normal"
    r = p.add_run(text)
    size = 16 if level == 1 else 13 if level == 2 else 11.5
    set_run(r, "黑体", size, True)
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, header=False, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header or align_center) else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    set_run(r, size=9.5, bold=header, color="FFFFFF" if header else None)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, "4472C4")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph("")
    return table


def add_placeholder(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell(cell, f"【系统界面截图占位：{title}】\n说明：此处由用户后续插入真实系统截图，不使用示意图或伪造截图。", False, True)
    shade(cell, "F2F2F2")
    doc.add_paragraph("")


def add_flow(doc, title, steps):
    add_heading(doc, title, 3)
    table = doc.add_table(rows=len(steps) * 2 - 1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, step in enumerate(steps):
        box = table.rows[i * 2].cells[0]
        set_cell(box, step, False, True)
        shade(box, "EAF2F8")
        if i < len(steps) - 1:
            arrow = table.rows[i * 2 + 1].cells[0]
            set_cell(arrow, "↓", False, True)
    doc.add_paragraph("")


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return doc


def build():
    doc = setup_doc()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("详细设计报告")
    set_run(r, "黑体", 24, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"【{PROJECT_NAME}】")
    set_run(r, "黑体", 18, True)
    doc.add_paragraph("")
    for label, value in [
        ("学生姓名", "（按实际参赛成员填写）"),
        ("学    院", "六盘水师范学院"),
        ("团队名称", TEAM_NAME),
        ("组    长", "（按实际负责人填写）"),
        ("日    期", str(date.today())),
    ]:
        add_paragraph(doc, f"{label}：{value}", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()

    add_heading(doc, "变更历史", 1)
    add_table(doc, ["序号", "版本", "变更日期", "变更内容", "变更者"], [
        ["1", "1.0", str(date.today()), "根据智教黔行真实系统代码与数据库结构创建详细设计报告", "项目组"],
        ["2", "", "", "", ""],
        ["3", "", "", "", ""],
    ])

    add_heading(doc, "目 录", 1)
    toc_items = [
        "0. 文档介绍", "1. 模块命名规则", "2. 模块汇总", "3. 系统总体设计",
        "4. 服务器端模块设计", "5. Web 端模块设计", "6. 微信小程序端模块设计",
        "7. AI 与预测服务设计", "8. 数据库设计", "9. 接口设计", "10. 系统流程设计", "11. 界面截图清单"
    ]
    for item in toc_items:
        add_paragraph(doc, item, first_line=False)
    doc.add_page_break()

    add_heading(doc, "0. 文档介绍", 1)
    add_heading(doc, "0.1 文档目的", 2)
    add_paragraph(doc, "本文档用于说明智教黔行旅游预测与智能服务系统的详细设计，包括系统模块划分、总体架构、业务流程、数据库设计、接口设计、Web 端、小程序端、AI 数字人和客流预测服务等内容，为后续开发、测试、部署、验收和答辩材料整理提供依据。")
    add_heading(doc, "0.2 文档范围", 2)
    add_paragraph(doc, "本文档覆盖 Web 管理端与游客端、微信小程序、业务后端、AI 后端、Python 客流预测服务、数字人服务、统一网关、MySQL 数据库、Redis 缓存、对象存储和第三方地图/天气/语音能力的设计说明。系统界面截图由用户后续补充真实截图，本文档只预留截图位置。")
    add_heading(doc, "0.3 读者对象", 2)
    add_paragraph(doc, "本文档面向项目设计人员、开发人员、测试人员、部署运维人员、指导教师和评审人员。")
    add_heading(doc, "0.4 参考资料", 2)
    add_table(doc, ["序号", "资料名称", "说明"], [
        ["1", "项目源代码", "Web 端、微信小程序、业务后端、AI 后端、预测服务、数字人服务和网关工程"],
        ["2", "数据库 Schema", "travel_prediction 业务库和小程序相关表结构"],
        ["3", "系统测试报告", "编译、构建、健康检查、数据库只读校验和单元测试结果"],
        ["4", "技术创新核对文档", "已按真实系统能力修订后的技术创新说明"],
    ])
    add_heading(doc, "0.5 术语与缩写解释", 2)
    add_table(doc, ["术语", "解释"], [
        ["Web 端", "基于 Vue 3 的浏览器端系统，包含游客端、商户端和管理员端"],
        ["小程序端", "基于 UniApp 的微信小程序，面向游客移动端使用"],
        ["业务后端", "基于 Spring Boot 的核心业务服务，处理用户、景区、内容、商户、订单和统计等业务"],
        ["AI 后端", "基于 Spring Boot 的 AI 代理服务，提供聊天、知识库、研学、行程规划和语音代理接口"],
        ["预测服务", "基于 FastAPI 的客流预测服务，提供 ARIMA、LSTM、双流融合预测接口"],
        ["数字人服务", "基于 FastAPI 的数字人聊天、语音和 WebSocket 服务"],
        ["RAG", "检索增强生成，通过知识库检索增强大模型回答"],
        ["JWT", "JSON Web Token，用于登录后的身份认证"],
    ])

    add_heading(doc, "1. 模块命名规则", 1)
    add_paragraph(doc, "后端 Java 类名采用大驼峰命名，例如 AuthController、AdminDashboardController；方法和变量采用小驼峰命名，例如 getUserInfo、simulatePolicyEffect；数据库表名采用小写下划线命名，例如 scenic_spots、visitor_predictions、ai_conversations；前端组件采用大驼峰命名，例如 DigitalHuman、FlowPrediction；路由路径采用小写短横线或业务语义路径。")

    add_heading(doc, "2. 模块汇总", 1)
    add_heading(doc, "2.1 模块汇总表", 2)
    add_table(doc, ["端/服务", "模块名称", "主要功能"], [
        ["统一网关", "服务路由与鉴权", "对业务后端、AI 后端、小程序后端、预测服务、数字人服务进行统一转发，提供跨域、限流、熔断和健康检查"],
        ["Web 游客端", "首页、景区探索、客流预测、行程规划、实时服务、AI 服务、旅游资讯、个人中心", "面向游客提供景区浏览、智能问答、预测查询、行程规划和个人信息管理"],
        ["Web 商户端", "景区管理中心、实时监测、数据分析、资源管理、评价系统、门票订单、紧急救援、政策沙盘", "面向景区/商户提供运营管理、数据分析和救援处理能力"],
        ["Web 管理端", "总览、数据分析、系统监控、商户审核、用户管理、内容管理、系统配置", "面向平台管理员提供平台治理和配置维护能力"],
        ["微信小程序", "首页、智游、文创商城、个人中心、数字人、红色研学、订单、行程、收藏、反馈", "面向移动端游客提供轻量化游览、购物、研学和个人服务"],
        ["业务后端", "用户、权限、景区、内容、商户、订单、统计、救援、政策模拟", "承载核心业务数据和管理流程"],
        ["AI 后端", "AI 聊天、知识库、行程规划、研学教育、语音代理", "负责 AI 能力聚合与业务系统对接"],
        ["预测服务", "客流预测、小时预测、模型信息、训练接口", "基于 ARIMA、LSTM、双流融合模型输出景区客流预测"],
        ["数字人服务", "文本对话、TTS、STT、WebSocket", "提供数字人交互、语音播报和流式会话能力"],
    ])

    add_heading(doc, "2.2 可复用模块列表", 2)
    add_table(doc, ["模块", "可复用能力", "适用范围"], [
        ["JWT 鉴权模块", "登录态校验、角色识别、请求拦截", "Web 端、管理端、商户端、小程序后端接口"],
        ["统一响应模块", "统一 code/message/data 返回结构", "Java 后端业务接口"],
        ["文件上传与 OSS 代理", "图片上传、图片代理访问、对象存储地址处理", "景区内容、轮播图、用户反馈、资源管理"],
        ["AI 客户端模块", "AI 聊天、预测服务、数字人服务调用封装", "AI 后端与业务后端"],
        ["可视化图表模块", "统计卡片、趋势图、热力图、政策模拟图表", "Web 游客端、商户端、管理端"],
        ["小程序商品订单模块", "商品、购物车、订单、支付、地址", "微信小程序文创商城"],
    ])

    add_heading(doc, "2.3 模块关系图", 2)
    add_flow(doc, "系统总体模块关系图", [
        "游客 Web 端 / 商户 Web 端 / 管理 Web 端 / 微信小程序",
        "统一网关：路由转发、跨域、JWT 鉴权、限流、熔断",
        "业务后端 / AI 后端 / 小程序后端 / 预测服务 / 数字人服务",
        "MySQL 业务库、Redis 缓存、对象存储、第三方地图天气语音服务",
        "输出：景区浏览、AI 导游、客流预测、订单管理、运营分析、政策模拟"
    ])

    add_heading(doc, "2.4 系统流程", 2)
    add_flow(doc, "用户访问主流程", [
        "用户进入 Web 端或微信小程序",
        "选择游客模式或登录注册",
        "浏览景区、查看资讯、查询客流预测、使用 AI 导游",
        "需要个人服务时进行收藏、行程规划、下单、反馈或救援请求",
        "系统记录业务数据并为统计分析、预测和运营管理提供数据基础"
    ])

    add_heading(doc, "3. 系统总体设计", 1)
    add_heading(doc, "3.1 技术架构", 2)
    add_table(doc, ["层次", "组成", "设计说明"], [
        ["表现层", "Vue 3 Web 端、UniApp 微信小程序", "提供游客、商户、管理员和移动端游客入口"],
        ["接入层", "Spring Cloud Gateway", "统一入口，负责路由、跨域、限流、熔断和鉴权白名单"],
        ["业务层", "Spring Boot 业务后端、小程序后端、AI 后端", "处理核心业务、移动端业务和 AI 能力代理"],
        ["智能服务层", "FastAPI 预测服务、FastAPI 数字人服务", "提供客流预测、数字人会话、语音合成和 WebSocket 通信"],
        ["数据层", "MySQL、Redis、OSS", "存储业务数据、缓存状态和图片等静态资源"],
        ["外部能力", "高德地图/天气、语音服务、大模型服务", "提供地图定位、天气查询、语音和智能问答能力"],
    ])
    add_heading(doc, "3.2 部署结构", 2)
    add_table(doc, ["服务", "默认端口", "职责"], [
        ["统一网关", "8888", "统一入口，转发 /api、/ai-api、/miniprogram-api、/prediction-api、/digital-human-api 等请求"],
        ["业务后端", "8080", "Web 端核心业务接口"],
        ["AI 后端", "8081", "AI 聊天、知识库、研学、行程规划和语音代理"],
        ["小程序后端", "8082", "微信小程序首页、商品、订单、地址、反馈、研学等接口"],
        ["数字人服务", "8083", "数字人聊天、语音、WebSocket"],
        ["预测服务", "8001", "景区客流预测和模型信息接口"],
        ["MySQL", "3306", "travel_prediction 业务数据库"],
        ["Redis", "6379", "缓存、限流和会话辅助"],
    ])

    add_heading(doc, "4. 服务器端模块设计", 1)
    add_heading(doc, "4.1 用户认证与权限模块", 2)
    add_paragraph(doc, "用户认证模块负责登录、注册、验证码、密码重置、JWT 令牌签发和用户信息查询。前端路由根据用户角色跳转到游客端、商户端或管理员端；网关和后端拦截器对非白名单接口进行身份校验。")
    add_table(doc, ["功能", "输入", "处理", "输出"], [
        ["用户登录", "账号、密码、验证码", "校验用户状态并签发 JWT", "用户信息、角色、访问令牌"],
        ["用户注册", "手机号/用户名、密码等", "创建普通用户或商户用户", "注册结果"],
        ["角色分流", "用户角色", "前端路由守卫判断 user/business/admin", "对应工作台首页"],
        ["权限校验", "请求头 Authorization", "网关和后端拦截器解析 JWT", "放行或返回未授权"],
    ])

    add_heading(doc, "4.2 景区与内容管理模块", 2)
    add_paragraph(doc, "景区与内容管理模块负责景区基础资料、轮播图、公告、新闻资讯、首页配置、实景预览和图片资源维护。游客端和小程序端读取公开内容；管理员端负责内容配置；商户端负责其管理范围内的景区资源维护。")
    add_table(doc, ["子模块", "主要对象", "说明"], [
        ["景区管理", "scenic_spots、sub_spots", "维护景区名称、地址、经纬度、介绍、图片、开放状态等信息"],
        ["首页内容", "banners、landing_config、showcases", "维护轮播图、首页展示内容和实景预览"],
        ["公告资讯", "announcements、news", "发布旅游资讯、景区公告和活动信息"],
        ["资源管理", "business_resources、business_resource_images", "商户维护景区设施、餐饮、住宿、活动等资源"],
    ])
    add_placeholder(doc, "Web 管理端景区管理页面")

    add_heading(doc, "4.3 商户运营管理模块", 2)
    add_paragraph(doc, "商户运营管理模块面向景区或商家角色，提供景区运营看板、实时监测、资源管理、评价回复、门票订单、活动公告、紧急救援处理和数据分析能力。")
    add_table(doc, ["功能", "设计说明"], [
        ["商户资料与审核", "商户提交资料后由管理员审核，审核通过后进入商户工作台"],
        ["实时监测", "展示景区客流、预警、状态和运营指标"],
        ["数据分析", "查看客流趋势、收入分析、游客来源、满意度和热点排行"],
        ["评价系统", "查看游客评价并进行商户回复"],
        ["紧急救援", "接收游客救援请求，处理、完成或取消救援记录"],
        ["政策沙盘", "调整折扣、补贴和容量等参数，观察模拟指标变化"],
    ])
    add_placeholder(doc, "商户端数据分析或政策沙盘页面")

    add_heading(doc, "4.4 管理员平台治理模块", 2)
    add_paragraph(doc, "管理员模块面向平台运营方，提供总览、统计报表、系统监控、商户审核、用户管理、内容管理、角色权限、系统通知、备份导出等功能。")
    add_table(doc, ["功能", "设计说明"], [
        ["平台总览", "展示用户、景区、订单、收入、访问等整体指标"],
        ["商户审核", "对商户入驻、合同、景区申请和审核日志进行管理"],
        ["用户管理", "管理用户列表、行为分析和用户状态"],
        ["内容管理", "配置首页、景区内容、轮播图、实景预览和小程序商品"],
        ["系统监控", "查看服务状态、JVM、数据库连接池、QPS、响应时间和缓存状态"],
        ["角色权限", "维护角色、权限和管理员账号配置"],
    ])
    add_placeholder(doc, "管理员端系统监控页面")

    add_heading(doc, "5. Web 端模块设计", 1)
    add_paragraph(doc, "Web 端采用 Vue 3、Vue Router、Pinia、Element Plus 和 ECharts 实现。系统根据登录角色进入不同端：游客端、商户端和管理员端。未登录用户可访问登录、注册、首页，以及部分游客可浏览页面。")
    add_table(doc, ["端", "路由/页面", "功能说明"], [
        ["游客端", "首页、景区探索、客流预测、行程规划、实时服务、AI 服务、个人中心、旅游资讯、热门景点", "提供游客浏览、预测查询、智能服务和个人资料维护"],
        ["商户端", "仪表盘、实时监测、数据分析、资源管理、评价系统、门票订单、紧急救援、统计报表、客流预测、政策沙盘、数据导出", "提供景区运营管理和数据分析"],
        ["管理员端", "总览、数据分析、系统监控、统计报表、客流预测、政策沙盘、商户管理、用户管理、内容管理、系统设置", "提供平台级运营治理能力"],
    ])
    add_placeholder(doc, "Web 游客端首页或景区探索页面")

    add_heading(doc, "6. 微信小程序端模块设计", 1)
    add_paragraph(doc, "微信小程序端采用 UniApp 实现，面向移动端游客。底部导航包含首页、智游、文创、我的，另包含景区详情、商品详情、搜索、本地服务、数字人、红色研学、订单、行程、地址、反馈、收藏、研学护照、足迹、优惠券和使用凭证等页面。")
    add_table(doc, ["页面", "功能说明"], [
        ["首页", "展示轮播图、推荐景区、热门服务和入口导航"],
        ["智游", "提供景区导览和智能游览入口"],
        ["文创商城", "展示文创商品，支持详情、购物车和订单"],
        ["黔小游数字人", "提供移动端聊天式 AI 导游体验"],
        ["红色研学", "展示研学内容、题库、积分和徽章相关入口"],
        ["我的", "个人资料、订单、收藏、足迹、优惠券、地址和意见反馈"],
    ])
    add_placeholder(doc, "微信小程序首页或黔小游页面")

    add_heading(doc, "7. AI 与预测服务设计", 1)
    add_heading(doc, "7.1 AI 智能服务设计", 2)
    add_paragraph(doc, "AI 后端负责将 Web 端、小程序端和业务后端的智能服务请求与知识库、大模型、数字人服务和预测服务进行整合。主要包括 AI 聊天、会话历史、知识库检索、AI 问答、行程规划、研学方案生成和语音合成代理。")
    add_table(doc, ["能力", "设计说明"], [
        ["AI 聊天", "接收用户问题，结合上下文生成旅游咨询回复"],
        ["知识库检索", "检索景区、文化、交通、研学、安全等知识内容"],
        ["AI 行程规划", "根据天数、人数、预算和偏好生成行程建议"],
        ["研学教育", "提供研学路线和 AI 研学方案生成"],
        ["语音服务", "代理数字人 TTS 能力，将文本转换为语音"],
    ])
    add_heading(doc, "7.2 客流预测服务设计", 2)
    add_paragraph(doc, "预测服务提供景区未来客流、全域聚合客流、小时级客流和模型信息接口。模型包含 ARIMA、LSTM 和双流融合模型，预测因子包括历史客流、天气、节假日、周末和海拔等。")
    add_table(doc, ["接口能力", "说明"], [
        ["景区未来 N 天预测", "按景区 ID、预测天数和模型类型返回未来客流、峰值时段、天气和拥挤度"],
        ["全域聚合预测", "汇总多个景区的预测结果，返回整体趋势"],
        ["小时级预测", "按日期返回景区或全域的小时分布"],
        ["模型训练与信息", "提供模型训练入口和模型元信息查询"],
    ])
    add_flow(doc, "客流预测处理流程图", [
        "前端选择景区、预测天数和模型类型",
        "业务后端或网关转发到预测服务",
        "预测服务读取历史数据或使用降级数据",
        "ARIMA 与 LSTM 模型分别生成预测结果",
        "双流融合模型计算最终预测值和拥挤度",
        "前端以折线图、柱状图或热力图展示结果"
    ])
    add_heading(doc, "7.3 数字人服务设计", 2)
    add_paragraph(doc, "数字人服务提供文本对话、TTS、STT 和 WebSocket 通道。服务端包含简易限流、缓存预热、RAG 检索、工具调用和语音合成逻辑，前端以悬浮数字人或聊天式页面进行展示。")
    add_flow(doc, "数字人问答流程图", [
        "用户在 Web 端或小程序端输入问题",
        "数字人服务接收消息并建立会话上下文",
        "RAG 检索本地景区与研学知识",
        "必要时调用客流预测或天气工具",
        "大模型生成回答并返回文本",
        "TTS 服务生成语音，前端展示文字并播放语音"
    ])

    add_heading(doc, "8. 数据库设计", 1)
    add_paragraph(doc, "系统主要使用 MySQL 数据库 travel_prediction 存储业务数据。数据库覆盖用户、角色权限、景区、内容、商户、订单、客流、AI 会话、研学、小程序商城和反馈等数据对象。Redis 用于缓存、限流和会话辅助。")
    add_heading(doc, "8.1 核心数据表", 2)
    add_table(doc, ["数据域", "代表数据表", "设计说明"], [
        ["用户与权限", "users、roles、permissions、role_permissions、wechat_users", "存储账号、角色、权限和微信用户映射"],
        ["景区与内容", "scenic_spots、sub_spots、banners、announcements、news、showcases、landing_config", "存储景区资料、轮播图、公告、新闻和首页展示内容"],
        ["商户运营", "merchant_profiles、merchant_contracts、merchant_audit_logs、business_resources、business_reviews、business_todos", "存储商户资料、合同、审核、资源、评价和待办"],
        ["客流与统计", "flow_records、daily_flow_summary、visitor_predictions、scenic_statistics、platform_statistics", "存储客流记录、日汇总、预测结果和平台统计"],
        ["AI 服务", "ai_conversations、ai_messages、ai_knowledge、ai_study_routes", "存储 AI 会话、消息、知识库和研学路线"],
        ["小程序商城", "mp_products、mp_cart、mp_orders、mp_order_items、mp_user_address", "存储文创商品、购物车、订单、订单明细和收货地址"],
        ["研学与积分", "mp_study_quiz、mp_study_badge、mp_user_badge、mp_user_points、mp_study_answer_log", "存储题库、徽章、积分和答题记录"],
        ["反馈与救援", "admin_feedback、mp_feedback、emergency_rescue", "存储意见反馈和紧急救援记录"],
    ])
    add_heading(doc, "8.2 主要表关系说明", 2)
    add_table(doc, ["关系", "说明"], [
        ["用户与 AI 会话", "一个用户可以拥有多个 AI 会话，每个会话包含多条消息"],
        ["景区与客流", "一个景区对应多条实时或历史客流记录，并可生成多条预测结果"],
        ["商户与景区", "商户资料、合同和可管理景区通过商户用户 ID 与景区 ID 建立关联"],
        ["商品与订单", "一个订单包含多条订单明细，订单明细关联商品信息和购买数量"],
        ["研学与积分", "用户答题记录关联题库，正确答题后更新积分和徽章"],
    ])

    add_heading(doc, "9. 接口设计", 1)
    add_heading(doc, "9.1 网关路由设计", 2)
    add_table(doc, ["外部路径", "转发服务", "说明"], [
        ["/api/**", "业务后端", "Web 端核心业务接口"],
        ["/ai-api/**", "AI 后端", "AI 聊天、知识库、研学和语音代理接口"],
        ["/miniprogram-api/**", "小程序后端", "微信小程序业务接口"],
        ["/prediction-api/**", "预测服务", "客流预测接口"],
        ["/digital-human-api/**", "数字人服务", "数字人聊天和语音接口"],
        ["/ws/**", "数字人 WebSocket", "数字人实时通信通道"],
    ])
    add_heading(doc, "9.2 业务接口设计", 2)
    add_table(doc, ["接口类别", "代表路径", "主要功能"], [
        ["认证接口", "/api/auth/**、/api/captcha", "登录、注册、验证码、退出和密码相关操作"],
        ["游客内容接口", "/api/content/**、/api/news/**、/api/scenics/**", "首页、景区、新闻、公告和内容展示"],
        ["用户接口", "/api/users、/api/user、/api/favorites", "用户资料、收藏、足迹、优惠券和通知"],
        ["商户接口", "/api/merchant/**", "商户资料、景区资源、活动、公告、评价、订单和数据分析"],
        ["管理员接口", "/api/admin/**", "平台总览、用户、商户、角色、内容、监控、备份和导出"],
        ["紧急救援接口", "/api/emergency-rescue/**", "创建救援、查看救援、处理救援、完成救援和统计"],
        ["政策模拟接口", "/api/admin/policy/simulate", "根据折扣、补贴和容量等参数模拟运营指标"],
    ])
    add_heading(doc, "9.3 AI、预测和小程序接口设计", 2)
    add_table(doc, ["服务", "代表接口", "主要功能"], [
        ["AI 后端", "/ai-api/chat/message、/ai-api/knowledge/search、/ai-api/ai-planning/generate、/ai-api/education/routes、/ai-api/speech/tts", "AI 对话、知识库、行程规划、研学和语音"],
        ["预测服务", "/prediction-api/api/prediction/flow/{scenicId}、/prediction-api/api/prediction/total、/prediction-api/api/prediction/hourly/{scenicId}", "景区预测、全域预测和小时预测"],
        ["数字人服务", "/digital-human-api/api/chat、/digital-human-api/api/tts、/digital-human-api/api/stt、/ws/avatar", "数字人文本对话、语音合成、语音识别和 WebSocket"],
        ["小程序后端", "/miniprogram-api/home/**、/miniprogram-api/scenic/**、/miniprogram-api/shop/**、/miniprogram-api/api/pay/**、/miniprogram-api/api/user/**", "小程序首页、景区、商城、支付、用户服务"],
    ])

    add_heading(doc, "10. 系统流程设计", 1)
    add_flow(doc, "登录与角色分流流程图", [
        "用户提交账号、密码和验证码",
        "业务后端校验账号状态和密码",
        "校验通过后生成 JWT 和用户角色信息",
        "前端保存登录态并根据角色跳转",
        "后续请求携带 JWT，由网关和后端进行权限校验"
    ])
    add_flow(doc, "景区浏览与 AI 导游流程图", [
        "用户进入景区探索或小程序智游页面",
        "系统加载景区基础信息、图片、公告和服务信息",
        "用户选择景区后查看详情、收藏或进入实时服务",
        "用户向 AI 导游提问景区、安全、研学或客流问题",
        "AI 服务结合知识库、天气和预测信息生成回答",
        "前端展示回答并可进行语音播报"
    ])
    add_flow(doc, "商户运营管理流程图", [
        "商户登录并进入景区智慧管理中心",
        "系统展示景区运营概览、实时监测和待办事项",
        "商户维护景区资源、活动、公告和新闻",
        "商户处理游客评价、门票订单和紧急救援",
        "商户查看数据分析、客流预测和政策模拟结果"
    ])
    add_flow(doc, "小程序文创商城订单流程图", [
        "用户进入文创商城并浏览商品",
        "查看商品详情并加入购物车或立即购买",
        "选择收货地址并提交订单",
        "系统创建订单和订单明细",
        "用户选择模拟支付、微信支付或积分支付",
        "支付成功后更新订单状态并生成使用凭证"
    ])
    add_flow(doc, "管理员治理流程图", [
        "管理员登录平台后台",
        "查看总览、统计报表和系统监控",
        "处理商户入驻审核、合同和景区申请",
        "维护用户、角色权限、内容配置和小程序资源",
        "导出数据、查看日志并处理系统通知"
    ])

    add_heading(doc, "11. 界面截图清单", 1)
    add_paragraph(doc, "以下截图由用户后续从真实运行系统中截取并插入，本文档不使用示意截图。")
    add_table(doc, ["序号", "截图名称", "建议截图页面"], [
        ["1", "Web 游客端首页", "首页或游客工作台"],
        ["2", "景区探索页面", "景区探索/景区详情"],
        ["3", "客流预测页面", "折线图、柱状图或热力图展示"],
        ["4", "AI 服务/数字人页面", "Web 端 AI 服务或数字人聊天窗口"],
        ["5", "商户端数据分析页面", "商户数据分析、实时监测或政策沙盘"],
        ["6", "管理员端系统监控页面", "系统监控、服务状态或平台总览"],
        ["7", "微信小程序首页", "首页 Tab"],
        ["8", "微信小程序黔小游页面", "聊天式数字人页面"],
        ["9", "微信小程序文创商城", "商品列表或商品详情"],
        ["10", "微信小程序个人中心", "我的页面、订单或收藏"],
    ])

    add_heading(doc, "12. 设计约束与后续完善", 1)
    add_table(doc, ["约束/风险", "说明", "后续处理建议"], [
        ["第三方能力依赖", "地图、天气、语音、大模型能力依赖外部服务配置", "在部署文档中明确 API Key、网络和降级策略"],
        ["预测数据质量", "预测效果依赖历史客流和天气节假日等数据质量", "持续补充真实客流数据并定期评估模型效果"],
        ["截图需真实补充", "本文档只预留截图占位", "由用户运行系统后插入真实界面截图"],
        ["权限边界", "不同角色页面和接口权限不同", "测试时覆盖游客、普通用户、商户、管理员四类角色"],
        ["移动端兼容", "小程序页面依赖编译产物", "提交前重新构建并使用微信开发者工具预览"],
    ])

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
