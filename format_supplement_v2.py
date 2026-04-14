from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor

doc = Document()

def set_font(run, cn='宋体', en='Times New Roman', size=9, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)

# 说明
p0 = doc.add_paragraph()
r0 = p0.add_run('替换原文"在六盘水真实数据上的测试表明……优势。"（同时引用图3和表3）')
set_font(r0, bold=True)

# 替换内容
p1 = doc.add_paragraph()
text = (
    '在六盘水市乌蒙大草原、梅花山、玉舍森林公园、水城古镇、明湖湿地5个核心景区'
    '2020—2024年的真实数据上（涵盖12类维度共计28万余条多源异构数据，其中历史客流约'
    '9000条、游客轨迹约12万条、气象数据约7300条、社交媒体情感约4.2万条、OTA退订约'
    '8500条等，各维度数据来源与贡献度详见图3及表3）的测试表明，此混合模型的平均绝对误差'
    '（MAE）仅为124.3，显著低于单一LSTM模型（MAE=186.5）和ARIMA模型（MAE=213.1），'
    '误差降低41.7%，且统计检验显示此提升具有显著性（p<0.01）。'
    '这充分证明了混合模型在处理山地旅游复杂时序数据方面的优势。'
)
r1 = p1.add_run(text)
set_font(r1)

# 分隔
doc.add_paragraph()

# 图3标题替换说明
p2 = doc.add_paragraph()
r2 = p2.add_run('图3标题替换：')
set_font(r2, bold=True)

p3 = doc.add_paragraph()
r3 = p3.add_run('原："图3 混合模型收敛曲线对比图"')
set_font(r3)
r3.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

p4 = doc.add_paragraph()
r4 = p4.add_run('改为："图3 系统整合的12类维度数据概览"')
set_font(r4)

# 图3说明文字
p5 = doc.add_paragraph()
r5 = p5.add_run('图3下方新增说明文字：')
set_font(r5, bold=True)

p6 = doc.add_paragraph()
text2 = (
    '如图3所示，系统整合的12类维度数据覆盖环境数据（气温/气象、紫外线指数、空气质量，'
    '贡献度28.7%）、游客行为数据（历史客流量、游客轨迹、OTA退订、社交媒体情感，贡献度32.1%）、'
    '山地特色数据（海拔梯度、坡度/地形、交通到达量，贡献度23.5%）以及节庆事件数据（节假日标签、'
    '民族节庆，贡献度15.7%）四大类别。其中游客行为数据贡献最大，游客轨迹数据量最大（约12万条），'
    '节庆事件数据虽量少但对火把节等突发客流预测具有关键作用。'
)
r6 = p6.add_run(text2)
set_font(r6)

doc.save('/Users/dongsiwei/TravelForecast/补充数据来源文字.docx')
print("补充数据来源文字.docx 已更新（含图3引用）")
