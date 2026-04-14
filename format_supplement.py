from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt

doc = Document('/Users/dongsiwei/TravelForecast/补充数据来源文字.docx')

def set_font(run, cn='宋体', en='Times New Roman', size=9, bold=False):
    """设置中文宋体、西文Times New Roman、小五号"""
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)

# 标题
p0 = doc.add_paragraph()
r0 = p0.add_run('以下文字替换原文"在六盘水真实数据上的测试表明……优势。"')
set_font(r0, bold=True)

# 替换内容
p1 = doc.add_paragraph()
text = (
    '在六盘水市乌蒙大草原、梅花山、玉舍森林公园、水城古镇、明湖湿地5个核心景区'
    '2020—2024年的真实数据上（涵盖12类维度共计28万余条多源异构数据，其中历史客流约'
    '9000条、游客轨迹约12万条、气象数据约7300条、社交媒体情感约4.2万条、OTA退订约'
    '8500条等，详见表3）的测试表明，此混合模型的平均绝对误差（MAE）仅为124.3，显著低于'
    '单一LSTM模型（MAE=186.5）和ARIMA模型（MAE=213.1），误差降低41.7%，且统计检验'
    '显示此提升具有显著性（p<0.01）。这充分证明了混合模型在处理山地旅游复杂时序数据方面的优势。'
)
r1 = p1.add_run(text)
set_font(r1)

doc.save('/Users/dongsiwei/TravelForecast/补充数据来源文字.docx')
print("补充数据来源文字.docx 已生成（宋体+TNR 小五）")
